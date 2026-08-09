#!/usr/bin/env python3
"""
Build the two PQ1 deliverables:
  1. PQ1_Quali_Rezults.xlsx  — per-SL parameter matrix (like the GMP example)
  2. EQP-PPS002_PQ-V01_FILLED.docx — populated PQ1 plan/report (Annex 07 template)
"""
import re, subprocess, pathlib, json, copy
from collections import defaultdict, OrderedDict
from datetime import datetime
import pandas as pd
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm
from copy import deepcopy

PDFS_DIR  = pathlib.Path("/tmp/pq1_pdfs")
TMPL_PQ1  = pathlib.Path("/tmp/pq1docs/EQPPPS002PQV01_Annex_07_PQ_FINAL_MALMED.docx")
XLSX_OUT  = pathlib.Path("/tmp/PQ1_Quali_Rezults.xlsx")
DOCX_OUT  = pathlib.Path("/tmp/EQP-PPS002_PQ-V01_FILLED.docx")

TRANS = str.maketrans({"Т":"T","Е":"E","О":"O","А":"A","С":"C","Р":"P","Х":"X","В":"B","М":"M","Н":"H","К":"K","У":"U"})
CPH_FIXUPS = {"C71_007":"C171_007","C71_008":"C171_008"}

PARAMS = [
    ("Faecal coliforms",         "Колиформни бактерии / 100ml | Coliform bacteria / 100ml",  r"Колиформни бактерии од фекално потекло во \d+ ml",       "Отсутно | Absent",   "CFU/100ml", None, 0,    "Micro"),
    ("Total bact. 37°C",         "Total Bacterial Count | TBC/1ml (37°C)",                    r"Вкупен број на бактерии на 37[°\s]*C во \d+ ml",          "20 CFU",            "CFU/ml",    None, 20,   "Micro"),
    ("Enterococci",              "Ентерококи / 100 ml | Enterococcus / 100ml",                r"Ентерококи во \d+ ml",                                    "Отсутно | Absent",  "CFU/100ml", None, 0,    "Micro"),
    ("Sulf.-red. clostridia",    "Сулфито ред. клостридии во 100 ml | Sulfite red. clostridia in 100 ml", r"Сулфито редуцирачки клостридии во \d+ ml",     "Отсутно | Absent",  "CFU/100ml", None, 0,    "Micro"),
    ("Aerobes 22°C",             "Total Bacterial Count | TBC/1ml (22°C)",                    r"Вкупен број на аеробни бактерии на 22[°\s]*С во \d+ ml",  "100 CFU",           "CFU/ml",    None, 100,  "Micro"),
    ("Pseudomonas aeruginosa",   "Pseudomonas aeruginosa / 100ml",                             r"Pseudomonas aeruginosa во \d+ ml",                       "Отсутно | Absent",  "CFU/100ml", None, 0,    "Micro"),
    ("E. coli",                  "E. coli / 100ml",                                            r"Е\.coli",                                                "Отсутно | Absent",  "CFU/100ml", None, 0,    "Micro"),
    ("Colour",                   "Боја | Color",                                               r"Боја",                                                   "< 20 Pt-Co",        "Pt-Co",     None, 20,   "Phys"),
    ("Turbidity",                "Матност | Turbidity",                                        r"Матност",                                                "1.5 NTU",           "NTU",       None, 1.5,  "Phys"),
    ("Temperature",              "Температура | Temperature",                                  r"температура",                                            "≤30°C",             "°C",        None, 30,   "Phys"),
    ("pH",                       "pH вредност | pH value",                                     r"^\s*рН\b",                                               "6.5-9.5",           "",          6.5,  9.5,  "PhysChem"),
    ("KMnO4 consumption",        "Потрошен KMnO4 (TOC) | KMnO4 Spent (TOC)",                   r"Потрошувачка на KMnO4",                                  "0-8 mg/L",          "mg/L",      0,    8,    "PhysChem"),
    ("Conductivity 20°C",        "Кондуктивност (на 20°C) | Conductivity (at 20°C)",           r"Електролитска спроводливост на 20[°\s]*[CС]",            "varies",            "µS/cm",     None, None, "PhysChem"),
    ("Ammonia",                  "Амонијак | Ammonia",                                         r"Амонијак",                                               "< 0.0005 mg/L",     "mg/L",      None, 0.0005, "Chem"),
    ("Nitrites",                 "Нитрити | Nitrites",                                         r"Нитрити",                                                "< 0.1 mg/L",        "mg/L",      None, 0.1,  "Chem"),
    ("Nitrates",                 "Нитрати | Nitrates",                                         r"Нитрати",                                                "< 50 mg/L",         "mg/L",      None, 50,   "Chem"),
    ("Chlorides",                "Хлориди | Chlorides",                                        r"Хлориди",                                                "< 250 mg/L",        "mg/L",      None, 250,  "Chem"),
    ("Iron",                     "Железо | Iron",                                              r"Железо",                                                 "< 0.2 mg/L",        "mg/L",      None, 0.2,  "Chem"),
    ("Manganese",                "Манган | Manganese",                                         r"Манган",                                                 "< 0.05 mg/L",       "mg/L",      None, 0.05, "Chem"),
]

# Per-water-class conductivity limits override
COND_LIMITS = {"TW":(None,2500),"TR":(None,650),"RO":(None,40)}

# SL descriptions (extracted from PDFs)
SL_DESCRIPTIONS = {}

RE_DATE = re.compile(r"Датум на земање:\s*(\d{2})\.(\d{2})\.(\d{4})")
RE_ISSUE = re.compile(r"Датум на издавање:\s*(\d{2})\.(\d{2})\.(\d{4})")
RE_LABNO = re.compile(r"Лаб\. број:\s*(\S+)")
RE_MM   = re.compile(r"Мерно место:\s*(.*?)(?=Вид\s|Странка\s|Датум на земање|Хигиено|$)", re.S)
RE_CODE = re.compile(r"\b([A-Z]{2})\s*[_-]\s*([A-Z]\d{1,4})\s*[_-]\s*(\d{2,4})(?!\d)")
RE_NUM  = re.compile(r"(?:н\.д\.|\d+(?:[.,]\d+)?)")

def text_of(p):
    return subprocess.run(["pdftotext","-layout",str(p),"-"],capture_output=True,text=True,timeout=30).stdout

def parse_value(s):
    s=s.strip()
    if s=="н.д.": return 0.0,"ND"
    try: return float(s.replace(",", ".")), s
    except ValueError: return None, s

def parse_pdf(p):
    sid = p.name.split()[0]
    if p.stat().st_size < 5000:
        return None
    t = text_of(p)
    d = RE_DATE.search(t); mm = RE_MM.search(t); iss = RE_ISSUE.search(t); ln = RE_LABNO.search(t)
    if not (d and mm): return None
    date_iso = f"{d.group(3)}-{d.group(2)}-{d.group(1)}"
    date_eu  = f"{d.group(1)}.{d.group(2)}.{d.group(3)}"
    issue    = f"{iss.group(1)}.{iss.group(2)}.{iss.group(3)}" if iss else ""
    doc_code = ln.group(1) if ln else f"{sid}/2026"
    seg_raw = mm.group(1).replace("\n"," ").strip()
    seg = seg_raw.translate(TRANS)
    m = RE_CODE.search(seg)
    if not m: return None
    sl, p1, p2 = m.groups()
    cph_raw = f"{p1}_{p2}"
    cph = CPH_FIXUPS.get(cph_raw, cph_raw)
    sp = f"{sl}_{cph}"
    # Capture human description (text after the SL code in the original cyrillic segment)
    # Find position of the code-like substring in the LATIN-translated seg, then take same offset in raw
    code_pos = m.end()
    desc_lat = seg[code_pos:].strip(" -—.,")
    # take roughly equivalent in raw seg
    desc_raw = seg_raw[code_pos:].strip(" -—.,") if code_pos < len(seg_raw) else desc_lat
    desc = (desc_raw or desc_lat)[:160]
    if sp not in SL_DESCRIPTIONS or len(SL_DESCRIPTIONS[sp]) < len(desc):
        SL_DESCRIPTIONS[sp] = desc

    lines = t.splitlines()
    results = {}
    for entry in PARAMS:
        key, label, pat, ac_text, unit, mn, mx, cat = entry
        rx = re.compile(pat)
        idx = None
        for i, ln_ in enumerate(lines):
            if rx.search(ln_):
                idx=i; break
        if idx is None:
            results[key] = {"raw":"", "num":None, "found":False}
            continue
        scope = " ".join(lines[idx:idx+3])
        m_pat = rx.search(scope)
        post = scope[m_pat.end():]
        m_num = RE_NUM.search(post)
        if not m_num:
            results[key] = {"raw":"", "num":None, "found":False}
            continue
        v, raw = parse_value(m_num.group(0))
        if key in ("Iron","Manganese") and v is not None and v != 0:
            v = v / 1000.0  # µg/L → mg/L
        results[key] = {"raw":raw, "num":v, "found":True}
    return {"sample_id":sid, "date_iso":date_iso, "date_eu":date_eu, "issue":issue,
            "doc_code":doc_code, "sl":sl, "cph":cph, "sp":sp,
            "results":results}

def is_oos(key, num, sl):
    if num is None: return False
    if key == "Conductivity 20°C":
        mn,mx = COND_LIMITS.get(sl, (None,2500))
    else:
        # find limits in PARAMS
        for k,*rest in PARAMS:
            if k==key:
                mn,mx = rest[4], rest[5]
                break
        else:
            return False
    if mn is not None and num < mn: return True
    if mx is not None and num > mx: return True
    return False

def format_result(r, key, sl):
    if not r["found"]: return ""
    if r["raw"] == "ND" or r["raw"] == "н.д.": return "н.д. | n.d."
    try:
        v = float(r["raw"].replace(",", "."))
        # iron/manganese converted to mg/L
        if key in ("Iron","Manganese"):
            v = r["num"]
        return f"{v:g}"
    except (ValueError, TypeError):
        return r["raw"]

def conductivity_ac(sl):
    mn,mx = COND_LIMITS.get(sl,(None,None))
    return f"< {mx} µS/cm" if mx else "varies"

# ============= XLSX BUILDER =============
def build_xlsx(samples):
    wb = Workbook()
    wb.remove(wb.active)

    # Helpers for styling
    head_font = Font(bold=True, color="FFFFFF", size=10)
    head_fill = PatternFill("solid", fgColor="305496")
    section_fill = PatternFill("solid", fgColor="D9E1F2")
    section_font = Font(bold=True, size=10)
    oos_fill = PatternFill("solid", fgColor="FFC7CE")
    oos_font = Font(bold=True, color="9C0006")
    nd_font = Font(italic=True, color="808080")
    border = Border(left=Side(style="thin"), right=Side(style="thin"),
                    top=Side(style="thin"), bottom=Side(style="thin"))
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left = Alignment(horizontal="left", vertical="center", wrap_text=True)

    # Build per-SL data
    by_sp = defaultdict(dict)  # sp -> {date_iso: sample}
    for s in samples:
        by_sp[s["sp"]][s["date_iso"]] = s

    all_dates = sorted({s["date_iso"] for s in samples})

    # ===== Main sheet: one big sheet listing all SLs sequentially =====
    ws = wb.create_sheet("PQ1 All SLs")

    row = 1
    ws.cell(row,1,"PQ1 Water Testing Results — RO Treatment System (EQP-PPS002)").font = Font(bold=True,size=14)
    ws.cell(row,1).fill = PatternFill("solid", fgColor="1F4E79")
    ws.cell(row,1).font = Font(bold=True,size=14,color="FFFFFF")
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4+len(all_dates))
    row += 2

    for sp in sorted(by_sp.keys()):
        sl_class = sp.split("_")[0]
        desc = SL_DESCRIPTIONS.get(sp, "")
        # SL header block
        ws.cell(row,1,"Water Sampling Location Code:").font = Font(bold=True)
        ws.cell(row,3,sp).font = Font(bold=True, color="0070C0")
        ws.cell(row,5,"Water Class:").font = Font(bold=True)
        ws.cell(row,6,{"TW":"Municipality Tap Water","TR":"Treated Water","RO":"Demineralised RO Water"}.get(sl_class,sl_class))
        row += 1
        ws.cell(row,1,"WSL Description:").font = Font(bold=True)
        ws.cell(row,3,desc).alignment = left
        ws.merge_cells(start_row=row, start_column=3, end_row=row, end_column=4+len(all_dates))
        row += 2

        # Parameter table header
        ws.cell(row,1,"Параметар | Parameter").font = head_font
        ws.cell(row,1).fill = head_fill
        ws.cell(row,1).alignment = center; ws.cell(row,1).border = border
        ws.cell(row,2,"A.C. (Acceptance Criteria)").font = head_font
        ws.cell(row,2).fill = head_fill
        ws.cell(row,2).alignment = center; ws.cell(row,2).border = border
        ws.cell(row,3,"Unit").font = head_font
        ws.cell(row,3).fill = head_fill
        ws.cell(row,3).alignment = center; ws.cell(row,3).border = border
        for ci, d in enumerate(all_dates):
            ws.cell(row,4+ci,d).font = head_font
            ws.cell(row,4+ci).fill = head_fill
            ws.cell(row,4+ci).alignment = center; ws.cell(row,4+ci).border = border
        row += 1
        # Sample-id row
        ws.cell(row,1,"Sample ID:").font = Font(italic=True, size=9)
        for ci, d in enumerate(all_dates):
            s = by_sp[sp].get(d)
            txt = s["doc_code"] if s else "Н/П | N/A"
            c = ws.cell(row,4+ci,txt)
            c.font = Font(italic=True, size=9, color="404040")
            c.alignment = center; c.border = border
        row += 1

        # Parameter rows grouped by category
        last_cat = None
        for entry in PARAMS:
            key, label, pat, ac_text, unit, mn, mx, cat = entry
            if cat != last_cat:
                cat_names = {"Phys":"1. Физички | Physical","PhysChem":"2. Физичко-хемиски | Physico-chemical",
                             "Chem":"3. Хемиски | Chemical","Micro":"4. Микробиолошки | Microbiological"}
                c = ws.cell(row,1,cat_names.get(cat,cat))
                c.font = section_font; c.fill = section_fill
                ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=3+len(all_dates))
                row += 1
                last_cat = cat
            ws.cell(row,1,label).alignment = left; ws.cell(row,1).border = border
            # Acceptance criteria — override conductivity per SL
            ac = conductivity_ac(sl_class) if key=="Conductivity 20°C" else ac_text
            ws.cell(row,2,ac).alignment = center; ws.cell(row,2).border = border
            ws.cell(row,3,unit).alignment = center; ws.cell(row,3).border = border
            for ci, d in enumerate(all_dates):
                s = by_sp[sp].get(d)
                if not s:
                    c = ws.cell(row,4+ci,"Н/П | N/A")
                    c.font = nd_font; c.alignment = center; c.border = border
                    continue
                r = s["results"].get(key, {"raw":"","num":None,"found":False})
                disp = format_result(r, key, sl_class)
                c = ws.cell(row,4+ci, disp)
                c.alignment = center; c.border = border
                if disp == "н.д. | n.d.":
                    c.font = nd_font
                if is_oos(key, r["num"], sl_class):
                    c.fill = oos_fill; c.font = oos_font
            row += 1
        row += 2  # blank row between SL blocks

    # Column widths
    ws.column_dimensions["A"].width = 42
    ws.column_dimensions["B"].width = 16
    ws.column_dimensions["C"].width = 10
    for i, d in enumerate(all_dates):
        ws.column_dimensions[get_column_letter(4+i)].width = 12

    ws.freeze_panes = "D1"

    # ===== OOS-only sheet =====
    ws2 = wb.create_sheet("OOS Findings")
    ws2.cell(1,1,"Out-of-Spec Findings (per parameter / per sample)").font = Font(bold=True, size=12)
    ws2.merge_cells(start_row=1,start_column=1,end_row=1,end_column=8)
    headers = ["Date","Sampling Point","Sample ID","Category","Parameter","Result","A.C.","Notes"]
    for ci,h in enumerate(headers):
        c = ws2.cell(2, ci+1, h)
        c.font = head_font; c.fill = head_fill; c.alignment = center; c.border = border
    r = 3
    for s in sorted(samples, key=lambda x:(x["date_iso"], x["sp"])):
        sl_class = s["sl"]
        for entry in PARAMS:
            key, label, pat, ac_text, unit, mn, mx, cat = entry
            res = s["results"].get(key, {"raw":"","num":None,"found":False})
            if not res["found"]: continue
            if not is_oos(key, res["num"], sl_class): continue
            ac = conductivity_ac(sl_class) if key=="Conductivity 20°C" else ac_text
            note = "Critical (RO output far above spec)" if (key=="Conductivity 20°C" and sl_class=="RO") else \
                   "Critical microbiological contamination" if (key=="Faecal coliforms" and res["num"]>0) else \
                   "Systemic — RO water naturally <pH 6.5 (spec inheritance gap)" if (key=="pH" and sl_class!="TW") else ""
            for ci, val in enumerate([s["date_eu"], s["sp"], s["sample_id"], cat, label, format_result(res,key,sl_class), ac, note]):
                c = ws2.cell(r, ci+1, val); c.border = border; c.alignment = left
                c.font = oos_font if ci==5 else Font(size=10)
            r += 1
    for col, w in enumerate([12,18,11,12,42,15,14,42], start=1):
        ws2.column_dimensions[get_column_letter(col)].width = w
    ws2.freeze_panes = "A3"

    # ===== Limits reference sheet =====
    ws3 = wb.create_sheet("Limits Reference")
    ws3.cell(1,1,"Per-Water-Class Specifications (A03/A05/A06)").font = Font(bold=True, size=12)
    ws3.merge_cells(start_row=1,start_column=1,end_row=1,end_column=6)
    headers = ["Category","Parameter","Unit","TW (A03)","TR (A05)","RO (A06)"]
    for ci,h in enumerate(headers):
        c = ws3.cell(2,ci+1,h); c.font=head_font; c.fill=head_fill; c.alignment=center; c.border=border
    r = 3
    for entry in PARAMS:
        key, label, pat, ac_text, unit, mn, mx, cat = entry
        if key == "Conductivity 20°C":
            tw_ac = "< 2500 µS/cm"; tr_ac = "< 650 µS/cm"; ro_ac = "< 40 µS/cm"
        else:
            tw_ac = tr_ac = ro_ac = ac_text
        for ci, val in enumerate([cat, label, unit, tw_ac, tr_ac, ro_ac]):
            c = ws3.cell(r,ci+1,val); c.border=border; c.alignment=left if ci<3 else center
        r += 1
    for col, w in enumerate([12,42,10,16,16,16], start=1):
        ws3.column_dimensions[get_column_letter(col)].width = w
    ws3.freeze_panes = "A3"

    # ===== Coverage matrix =====
    ws4 = wb.create_sheet("Coverage Matrix")
    ws4.cell(1,1,"Sampling Coverage: SL × Date (✓ = sampled, blank = not sampled)").font = Font(bold=True, size=12)
    ws4.merge_cells(start_row=1,start_column=1,end_row=1,end_column=1+len(all_dates))
    ws4.cell(2,1,"Sampling Point").font = head_font
    ws4.cell(2,1).fill = head_fill; ws4.cell(2,1).alignment = center; ws4.cell(2,1).border = border
    for ci, d in enumerate(all_dates):
        c = ws4.cell(2, 2+ci, d); c.font = head_font; c.fill = head_fill; c.alignment = center; c.border = border
    r = 3
    for sp in sorted(by_sp.keys()):
        ws4.cell(r,1,sp).font = Font(bold=True)
        ws4.cell(r,1).border = border
        for ci, d in enumerate(all_dates):
            s = by_sp[sp].get(d)
            if s:
                # Any OOS?
                has_oos = any(is_oos(k, s["results"].get(k,{}).get("num"), s["sl"]) for k,*_ in PARAMS if s["results"].get(k,{}).get("found"))
                c = ws4.cell(r, 2+ci, "✗" if has_oos else "✓")
                c.alignment = center; c.border = border
                if has_oos:
                    c.font = oos_font; c.fill = oos_fill
                else:
                    c.font = Font(color="00B050", bold=True)
            else:
                c = ws4.cell(r, 2+ci, "")
                c.border = border
        r += 1
    ws4.column_dimensions["A"].width = 18
    for i, d in enumerate(all_dates):
        ws4.column_dimensions[get_column_letter(2+i)].width = 12
    ws4.freeze_panes = "B3"

    wb.save(XLSX_OUT)
    print(f"✓ Wrote {XLSX_OUT} ({XLSX_OUT.stat().st_size//1024} KB)")
    print(f"  Sheets: PQ1 All SLs, OOS Findings, Limits Reference, Coverage Matrix")
    print(f"  SLs covered: {len(by_sp)}; Dates: {len(all_dates)}; OOS rows in summary: {r-3}")

# ============= DOCX HELPERS =============
def get_tc(table, row_idx, tc_idx):
    tbl = table._tbl
    trs = tbl.findall(qn("w:tr"))
    tcs = trs[row_idx].findall(qn("w:tc"))
    return tcs[tc_idx]

def set_para_text(p, text, *, bold=False):
    """Replace the text of a single paragraph (accepts python-docx Paragraph or XML <w:p>)."""
    p_xml = p._p if hasattr(p, "_p") else p
    for r in p_xml.findall(qn("w:r")):
        p_xml.remove(r)
    r = OxmlElement("w:r")
    if bold:
        rPr = OxmlElement("w:rPr"); b = OxmlElement("w:b"); rPr.append(b); r.append(rPr)
    t = OxmlElement("w:t"); t.text = text or ""
    t.set(qn("xml:space"),"preserve"); r.append(t); p_xml.append(r)

def set_tc_text(tc, text, *, bold=False, center=False):
    for p in tc.findall(qn("w:p")):
        tc.remove(p)
    p = OxmlElement("w:p")
    if center:
        pPr = OxmlElement("w:pPr"); jc = OxmlElement("w:jc"); jc.set(qn("w:val"),"center")
        pPr.append(jc); p.append(pPr)
    r = OxmlElement("w:r")
    if bold:
        rPr = OxmlElement("w:rPr"); b = OxmlElement("w:b"); rPr.append(b); r.append(rPr)
    t = OxmlElement("w:t"); t.text = text or ""; t.set(qn("xml:space"),"preserve")
    r.append(t); p.append(r); tc.append(p)

def append_tc_paragraph(tc, text, *, bold=False):
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    if bold:
        rPr = OxmlElement("w:rPr"); b = OxmlElement("w:b"); rPr.append(b); r.append(rPr)
    t = OxmlElement("w:t"); t.text = text; t.set(qn("xml:space"),"preserve")
    r.append(t); p.append(r); tc.append(p)

def clone_row(table, row_idx):
    """Duplicate a table row and append; return new row."""
    tbl = table._tbl
    trs = tbl.findall(qn("w:tr"))
    new_tr = deepcopy(trs[row_idx])
    tbl.append(new_tr)
    # Clear text in cloned row
    for tc in new_tr.findall(qn("w:tc")):
        for p in tc.findall(qn("w:p")):
            tc.remove(p)
        # Add one empty paragraph
        p = OxmlElement("w:p"); tc.append(p)
    return new_tr

def add_para_after(target_p, text, *, bold=False, style=None):
    """Insert a new paragraph immediately after target paragraph."""
    new_p = OxmlElement("w:p")
    if style:
        pPr = OxmlElement("w:pPr")
        pStyle = OxmlElement("w:pStyle"); pStyle.set(qn("w:val"), style)
        pPr.append(pStyle); new_p.append(pPr)
    r = OxmlElement("w:r")
    if bold:
        rPr = OxmlElement("w:rPr"); b = OxmlElement("w:b"); rPr.append(b); r.append(rPr)
    t = OxmlElement("w:t"); t.text = text; t.set(qn("xml:space"),"preserve")
    r.append(t); new_p.append(r)
    target_p.addnext(new_p)
    return new_p

# ============= DOCX BUILDER =============
def build_docx(samples):
    doc = Document(TMPL_PQ1)
    by_date = defaultdict(list)
    for s in samples: by_date[s["date_iso"]].append(s)
    dates_sorted = sorted(by_date.keys())
    sps = sorted({s["sp"] for s in samples})
    n_samples = len(samples)
    sl_classes = sorted({s["sl"] for s in samples})

    # ===== Replace placeholder text in key paragraphs =====
    # Helper: scan all paragraphs in doc + tables for a regex; replace text
    def replace_in_paras(pattern, replacement, only_first=False):
        rx = re.compile(pattern, re.DOTALL)
        count = 0
        # body paragraphs
        for p in doc.paragraphs:
            full = "".join(r.text or "" for r in p.runs)
            if rx.search(full):
                # clear runs, set replacement
                new_text = rx.sub(replacement, full)
                set_para_text(p, new_text)
                count += 1
                if only_first: return count
        return count

    # Replace placeholder in P50 (Purpose) — equipment name
    replace_in_paras(
        r"\[Внесете име на опрема/систем[^\]]*\]",
        "Систем за третман на вода со реверзна осмоза „Маган Мак“ (EQP-PPS002) | Reverse Osmosis Water Treatment System „Magan Mak“ (EQP-PPS002)"
    )
    replace_in_paras(r"\[на пр\., биолошки и хемиски квалитет[^\]]*\]",
                     "биолошки квалитет (микробиолошки), физичко-хемиски квалитет и хемиски квалитет на произведената вода од реверзна осмоза | biological quality (microbiological), physico-chemical quality, and chemical quality of the produced reverse-osmosis water")
    replace_in_paras(r"\[на пр\., обезбедуваат произведена\.\.\.[^\]]*\]",
                     "обезбедуваат произведена вода со конзистентен квалитет што ги задоволува спецификациите за TW (тап), TR (третирана пред Р.О.) и RO (демиерализирана) | provide produced water with consistent quality meeting the specifications for TW (tap), TR (treated pre-RO) and RO (demineralised)")
    # Scope text (P67) — actual numbers
    replace_in_paras(r"17 локации за мострирање", f"{len(sps)} локации за мострирање")
    replace_in_paras(r"17 sampling locations", f"{len(sps)} sampling locations")

    # ===== Fill the title-page approval table (T0): leave signature blanks =====
    # No edits — signatures filled by hand

    # ===== Fill the Responsibilities table T2 (leave name/signature blank) =====
    # No edits — to be filled by user

    # ===== Fill Prerequisites table T3 with our actual reference docs =====
    t3 = doc.tables[3]
    prereq_data = [
        # (doc/condition, doc_no, approval_date, pass_fail)
        ("Одобрен извештај за OQ | Approved OQ Report",
         "EQP-PPS002_OQ-V01",
         "ДД-ММ-ГГГГ | DD-MM-YYYY", None),
        ("Одобрени СОП за работа и CIP | Approved SOPs for Operation and CIP",
         "PP-QC-SOP-014 (RO water sampling); SOP for RO operation; SOP for CIP",
         "ДД-ММ-ГГГГ | DD-MM-YYYY", None),
        ("Сертификати за обука за релевантен персонал | Training certificates for relevant personnel",
         "Дневник за обука на оператори | Operator training log",
         "ДД-ММ-ГГГГ | DD-MM-YYYY", None),
        ("Потребни материјали и медиуми достапни | Required materials and media available",
         "Градска вода (TW); реагенси за тестирање; стерилна опрема за земање примероци",
         "ДД-ММ-ГГГГ | DD-MM-YYYY", None),
        ("Валидна калибрација за сите критични инструменти | Valid calibration for all critical instruments",
         "CAL-EC; CAL-PRESS; CAL-FLOW; CAL-TEMP",
         "ДД-ММ-ГГГГ | DD-MM-YYYY", None),
        ("Одобрен договор за надворешна лабораторија | Approved External Lab Contract",
         "ЈЗУ Центар за јавно здравје — Куманово (CJZ Kumanovo); ISO/IEC 17025:2018 сертификат бр. ЛТ-011/2009",
         "ДД-ММ-ГГГГ | DD-MM-YYYY", None),
    ]
    for i, (cond, dno, dat, pf) in enumerate(prereq_data, start=1):
        if i+1 > len(t3.rows): break  # 7 rows total (1 header + 6 data)
        set_tc_text(get_tc(t3, i, 0), cond)
        set_tc_text(get_tc(t3, i, 1), dno)
        set_tc_text(get_tc(t3, i, 2), dat, center=True)
        # leave pass/fail checkboxes as-is

    # ===== Fill T7 (main test acceptance criteria table) =====
    # Template T7 has 15 rows with placeholder OQ-style content. Replace with PQ1-style per-class tests.
    t7 = doc.tables[7]
    # Clear all rows except header (row 0)
    tbl7 = t7._tbl
    trs7 = tbl7.findall(qn("w:tr"))
    # Keep row 0 (header), remove rest
    for tr in list(trs7[1:]):
        tbl7.remove(tr)
    # Now add fresh rows. One row per (water class × parameter group) summary.
    summary_rows = [
        # (desc, materials/SLs, key parameters, AC, actual, pass/fail, DP)
        ("Дневно земање примероци од Градска Вода (TW) — еднаш дневно во период од 9 тест-денови | "
         "Daily sampling from Municipal Tap Water (TW) — once per day over 9 test days",
         f"TW_T161_001 (1 SL) — {len([s for s in samples if s['sl']=='TW'])} samples",
         "Сите 19 параметри по A03 (вкл. микробио, физ.-хем., хем., физ.). Conductivity < 2500 µS/cm | "
         "All 19 parameters per A03 (incl. micro, phys-chem, chem, phys). Conductivity < 2500 µS/cm",
         "сите параметри во спецификација | all parameters within spec",
         "PASS — сите 9 примероци поминаа сите параметри | PASS — all 9 samples passed all parameters",
         "Pass", "Прилог 1 / Attachment 1 (Quali_Rezults.xlsx, sheet PQ1 All SLs, TW_T161_001)"),
        ("Дневно земање примероци од Третирана Вода (TR) — пред единицата за реверзна осмоза | "
         "Daily sampling from Treated Water (TR) — before the reverse osmosis unit",
         f"TR_T161_001 (1 SL) — {len([s for s in samples if s['sl']=='TR'])} samples",
         "Сите 19 параметри по A05. Conductivity < 650 µS/cm | All 19 parameters per A05. Conductivity < 650 µS/cm",
         "Conductivity и сите хемиски/микро параметри во спецификација; pH под долна граница 6.5 на 10/10 теста | "
         "Conductivity and all chemical/micro within spec; pH below 6.5 lower bound on 10/10 samples",
         "OOS — pH 5.31–6.20 на сите 10 примероци (видете Анализа на резултати) | OOS — pH 5.31–6.20 on all 10 samples (see Analysis section)",
         "Fail", "Прилог 1 (sheet TR_T161_001), Прилог 2 (A02 за секој ден)"),
        ("Дневно земање примероци од Произведена RO вода — сите кориснички локации | "
         "Daily sampling from produced RO water — all user locations (15 unique RO sampling points across 5 reservoir vessels + 10 user points)",
         f"RO_T161_001..005 (резервоар, 5 садови) + RO_C158_006, RO_C171_007/008, RO_E43_012, RO_E82_013, RO_E84_014, RO_F97_009, RO_F117_010, RO_F119_011 — {len([s for s in samples if s['sl']=='RO'])} samples",
         "Сите 19 параметри по A06. Conductivity < 40 µS/cm | All 19 parameters per A06. Conductivity < 40 µS/cm",
         "Conductivity OOS на 6/6 тест-денови на RO_F117_010 (577–1051 µS/cm vs <40 спец); Faecal coliforms 80 CFU/100ml (29.04.26, RO_C158_006); pH под 6.5 на 53/70 примероци (системски) | "
         "Conductivity OOS on 6/6 tested days at RO_F117_010 (577–1051 µS/cm vs <40 spec); Faecal coliforms 80 CFU/100ml (29.04.26, RO_C158_006); pH < 6.5 on 53/70 samples (systemic)",
         "Fail — критичен дефект на RO_F117_010 + единечна микробиолошка контаминација + системска pH | Fail — critical defect at RO_F117_010 + single microbio contamination + systemic pH",
         "Fail", "Прилог 1 (sheets RO_*), Прилог 2 (A02 per day), Прилог 3 (А06 per sample)"),
    ]
    for desc, mats, key_params, ac, actual, pf, dp in summary_rows:
        # Clone the header row to get the same formatting
        new_tr = deepcopy(trs7[0])
        # Clear text in new row
        for tc in new_tr.findall(qn("w:tc")):
            for p in tc.findall(qn("w:p")): tc.remove(p)
            tc.append(OxmlElement("w:p"))
        tbl7.append(new_tr)
        tcs = new_tr.findall(qn("w:tc"))
        for ci, text in enumerate([desc, mats, key_params, ac, actual, "", dp]):
            if ci < len(tcs):
                set_tc_text(tcs[ci], text)
        # Set Pass/Fail checkbox manually in col 5
        if len(tcs) > 5:
            set_tc_text(tcs[5], "☒ Поминува | ☒ Pass" if pf=="Pass" else "☒ Неуспех | ☒ Fail",
                        bold=(pf=="Fail"))

    # ===== Fill the Deficiency List T11 =====
    t11 = doc.tables[11]
    tbl11 = t11._tbl
    trs11 = tbl11.findall(qn("w:tr"))
    # Remove placeholder rows 1..6 (keep header row 0)
    for tr in list(trs11[1:]):
        tbl11.remove(tr)
    # Collect deficiencies: (DP_id, test_obj, required, current, q_critical, date, signature)
    deficiencies = []
    # 1) Conductivity failures at RO_F117_010
    cond_failures = [s for s in samples if s["sl"]=="RO" and s.get("cph")=="F117_010"
                     and s["results"].get("Conductivity 20°C",{}).get("num") and s["results"]["Conductivity 20°C"]["num"]>40]
    for s in cond_failures:
        v = s["results"]["Conductivity 20°C"]["num"]
        deficiencies.append((
            f"DP-PPS002-PQ-COND-{s['sample_id']}",
            f"Conductivity at RO_F117_010 (sample {s['sample_id']})",
            "< 40 µS/cm (per A06 — Demineralised RO Water)",
            f"{v:g} µS/cm  ({v/40:.1f}× over spec)",
            "Да | Yes (CRITICAL)",
            s["date_eu"], ""
        ))
    # 2) Faecal coliforms at 1203
    for s in [x for x in samples if x["results"].get("Faecal coliforms",{}).get("num",0) and x["results"]["Faecal coliforms"]["num"]>0]:
        v = s["results"]["Faecal coliforms"]["num"]
        deficiencies.append((
            f"DP-PPS002-PQ-MICRO-{s['sample_id']}",
            f"Faecal coliforms at {s['sp']} (sample {s['sample_id']})",
            "Отсутно | Absent (0 CFU/100ml)",
            f"{int(v)} CFU/100ml",
            "Да | Yes (CRITICAL)",
            s["date_eu"], ""
        ))
    # 3) Systemic pH (one consolidated row)
    ph_failures = [s for s in samples if s["sl"]!="TW" and s["results"].get("pH",{}).get("num") and s["results"]["pH"]["num"]<6.5]
    if ph_failures:
        ph_vals = [s["results"]["pH"]["num"] for s in ph_failures]
        deficiencies.append((
            "DP-PPS002-PQ-PH-SYSTEMIC",
            f"pH systemic finding — {len(ph_failures)} samples on TR/RO water across all sampling days",
            "6.5 – 9.5 (per A03/A05/A06)",
            f"{min(ph_vals):.2f} – {max(ph_vals):.2f} (systemic — expected for demineralised water)",
            "Не | No (Specification inheritance issue — see Analysis)",
            "Across 9 test days", ""
        ))
    # 4) Ammonia spec-vs-capability gap (no actual OOS but flag the gap)
    deficiencies.append((
        "DP-PPS002-PQ-NH3-SPEC-GAP",
        "Ammonia specification < 0.0005 mg/L cannot be verified by contract lab (LOD = 0.016 mg/L)",
        "< 0.0005 mg/L per A03/A05/A06",
        "Lab reports 'н.д.' which only proves < 0.016 mg/L (32× higher than spec)",
        "Не | No (Documentation/specification gap — not equipment fault)",
        "Identified during PQ1 review", ""
    ))
    # Add rows
    for d in deficiencies:
        new_tr = deepcopy(trs11[0])
        for tc in new_tr.findall(qn("w:tc")):
            for p in tc.findall(qn("w:p")): tc.remove(p)
            tc.append(OxmlElement("w:p"))
        tbl11.append(new_tr)
        tcs = new_tr.findall(qn("w:tc"))
        for ci, text in enumerate(d):
            if ci < len(tcs):
                set_tc_text(tcs[ci], text, bold=(ci==4 and "Yes" in text))

    # ===== Add Analysis of Test Results section AFTER the deficiency-list block =====
    # Find a stable anchor. The "Summary of Test Results & Conclusion" header is around P99.
    # Insert before it.
    anchor = None
    for p in doc.paragraphs:
        text = "".join(r.text or "" for r in p.runs).strip()
        if "Резиме на резултатите од тестот и заклучок" in text or "Summary of Test Results" in text:
            anchor = p; break
    if anchor:
        # Sections to add (insert in REVERSE order so they appear in correct order)
        sections = [
            # (heading, [paragraphs])
            ("АНАЛИЗА НА РЕЗУЛТАТИТЕ ОД ТЕСТИРАЊЕ | ANALYSIS OF TEST RESULTS", [
                "Овој дел дава наративна анализа базирана на ризик на сите Out-of-Specification (OOS) наоди идентификувани во "
                "Списокот на недостатоци погоре. | This section provides a risk-based narrative analysis of all Out-of-Specification "
                "(OOS) findings identified in the Deficiency List above."
            ]),
            ("1. Системски pH под спецификацијата на TR и RO вода | "
             "1. Systemic pH below specification on TR and RO water", [
                "Преглед | Overview",
                f"Од {len([s for s in samples if s['sl']!='TW'])} TR/RO примероци, "
                f"{sum(1 for s in samples if s['sl']!='TW' and s['results'].get('pH',{}).get('num') and s['results']['pH']['num']<6.5)} "
                "покажуваат pH под долната спецификациона граница од 6.5. Распонот е 5.29–6.49. "
                "За TW (градска) вода сите 9 примероци покажуваат pH во опсегот 6.65–7.93 — во спецификација. | "
                f"Of {len([s for s in samples if s['sl']!='TW'])} TR/RO samples, "
                f"{sum(1 for s in samples if s['sl']!='TW' and s['results'].get('pH',{}).get('num') and s['results']['pH']['num']<6.5)} "
                "show pH below the lower spec bound of 6.5. Range: 5.29–6.49. "
                "All 9 TW samples show pH in range 6.65–7.93 — within spec.",
                "Анализа на коренска причина | Root Cause Analysis",
                "Реверзна осмоза по дизајн ги отстранува бикарбонатите и карбонатите кои се природните pH-баферни јони во водата. "
                "Остатокот од растворен CO₂ во пермеатот формира јаглеродна киселина (H₂CO₃), што го намалува pH во опсегот ~5.0–6.5. "
                "Ова е инхерентна, очекувана и универзално документирана карактеристика на демиерализирана RO вода. "
                "Не претставува дефект на опремата. | Reverse osmosis by design strips bicarbonate and carbonate ions, "
                "the water's natural pH buffers. Dissolved CO₂ remaining in the permeate forms carbonic acid (H₂CO₃), pushing "
                "pH into the ~5.0–6.5 range. This is an inherent, expected, and universally documented characteristic of demineralised "
                "RO water. It is NOT an equipment defect.",
                "Заклучок и Препорака | Conclusion and Recommendation",
                "Спецификацијата pH 6.5–9.5 е наследена од стандардот за вода за пиење (Pravilnik 183/2018) и НЕ е соодветна за "
                "демиерализирана RO вода. Ph.Eur. за Purified Water (Пречистена Вода) не пропишува pH граница (само "
                "conductivity ≤ 4.3 µS/cm на 20°C и TOC ≤ 0.5 mg/L). Се препорачува: (а) измена на спецификациите A05/A06 "
                "за да се отстрани неприменливата pH граница за TR и RO вода, или (б) прифаќање на pH 5.0–7.5 како "
                "процесна контрола (не спецификација), и (в) задржување pH 6.5–9.5 само за TW. По имплементација на оваа "
                "корекција на спецификацијата, овие 63 наоди се прекласифицираат како 'не-OOS'. | "
                "The pH 6.5–9.5 specification is inherited from the drinking-water standard (Pravilnik 183/2018) and is NOT "
                "appropriate for demineralised RO water. Ph.Eur. Purified Water specification sets no pH limit (only "
                "conductivity ≤ 4.3 µS/cm at 20°C and TOC ≤ 0.5 mg/L). Recommendation: (a) revise A05/A06 specifications to "
                "remove the inapplicable pH limit for TR and RO water, or (b) accept pH 5.0–7.5 as a process control "
                "(not a specification), and (c) retain pH 6.5–9.5 only for TW. Upon implementation of this specification "
                "correction, these 63 findings reclassify as non-OOS."
            ]),
            ("2. КРИТИЧЕН ДЕФЕКТ — Conductivity на RO_F117_010 далеку над спецификација | "
             "2. CRITICAL DEFECT — Conductivity at RO_F117_010 far above specification", [
                "Преглед | Overview",
                "Локацијата за мострирање RO_F117_010 (чешма со RO вода во просторија за миење) покажа conductivity вредности "
                "од 577–1051 µS/cm на 6 од 9 тест-денови. Спецификацијата за демиерализирана RO вода е < 40 µS/cm. "
                "Овие вредности се 14–26× над спецификација и одговараат на качество на немиерализирана/градска вода, "
                "не на RO пермеат. | Sampling location RO_F117_010 (RO water tap in washing room) showed conductivity readings "
                "of 577–1051 µS/cm on 6 of 9 test days. The specification for demineralised RO water is < 40 µS/cm. "
                "These values are 14–26× over spec and match unmineralised/tap-water quality, not RO permeate.",
                "Резиме на резултатите | Results Summary",
                "28.04.26: 1051 µS/cm | 29.04.26: 942 µS/cm | 30.04.26: 799 µS/cm | "
                "04.05.26: 577 µS/cm | 05.05.26: 686 µS/cm | 06.05.26: 713 µS/cm | "
                "07.05.26: ~21 µS/cm (PASS) | 08.05.26: ~22 µS/cm (PASS) | 11.05.26: ~21 µS/cm (PASS).",
                "Преминот од tap-water-ниво (~700 µS/cm) кон RO-quality (~20 µS/cm) се случи помеѓу 06.05.26 и 07.05.26. "
                "Ова сугерира дека е извршена интервенција (поправка на проток/вентил, изолација на крос-конекција, "
                "или замена на гранка) во овој период. | "
                "Transition from tap-water level (~700 µS/cm) to RO quality (~20 µS/cm) occurred between 06.05.26 and 07.05.26. "
                "This suggests an intervention (flow/valve repair, cross-connection isolation, or line replacement) was performed "
                "in that window.",
                "Можни причини | Potential Root Causes",
                "1. Крос-конекција: чешмата на RO_F117_010 можеби била поврзана со tap-water линија наместо со RO дистрибутивната "
                "мрежа. | Cross-connection: the tap at RO_F117_010 may have been connected to a tap-water line instead of the RO "
                "distribution network.",
                "2. Дефектен/затворен вентил: вентилот за изолација помеѓу RO дистрибутивната мрежа и tap-water мрежата можеби "
                "бил отворен, дозволувајќи мешање. | Faulty/open isolation valve: the valve between RO distribution and tap-water "
                "network may have been open, allowing mixing.",
                "3. Деградација на RO мембраната во гранката што ја снабдува таа точка. | RO membrane degradation in the branch "
                "supplying that point.",
                "Заклучок | Conclusion",
                "Овој наод НЕ може да биде системски pH-стил толкуван како 'нормално за RO вода' — conductivity од 1000 µS/cm "
                "потврдува дека водата НЕ е RO пермеат. Локацијата RO_F117_010 не ја задоволувала функцијата за која е "
                "проектирана во 6 од 9 тест денови (66.7% од времето). | This finding cannot be re-interpreted as 'normal "
                "for RO water' — conductivity of 1000 µS/cm confirms the water is NOT RO permeate. RO_F117_010 did not fulfill "
                "its intended function on 6 of 9 test days (66.7% of the time).",
                "Препораки | Recommendations",
                "1. ИТНО: спроведете инвестигација во периодот 06–07.05.26 — кој изврши интервенција, што беше променето. | "
                "URGENT: investigate the 06–07.05.26 timeframe — who performed work, what was changed.",
                "2. Изведете press-and-purge тест на RO_F117_010 за да се потврди тековен conductivity ниво. | Perform a "
                "press-and-purge test at RO_F117_010 to confirm current conductivity level.",
                "3. Документирајте го дизајнот на дистрибутивната мрежа за оваа гранка (P&ID). | Document the distribution "
                "network design for this branch (P&ID).",
                "4. Не пуштајте таа линија во GMP-критична употреба до завршување на CAPA. | Do not release this line for "
                "GMP-critical use until CAPA closure.",
                "5. Повторете PQ1 за оваа специфична точка по корективна акција. | Repeat PQ1 for this specific point after "
                "corrective action."
            ]),
            ("3. КРИТИЧНА КОНТАМИНАЦИЈА — Faecal coliforms на RO_C158_006 (sample 1203, 29.04.26) | "
             "3. CRITICAL CONTAMINATION — Faecal coliforms at RO_C158_006 (sample 1203, 29.04.26)", [
                "Преглед | Overview",
                "Sample 1203 (29.04.26, RO_C158_006 — RO вода во просторија за иригација) пријави 80 CFU/100ml фекални "
                "колиформи. Сите други микробиолошки параметри се н.д.; сите други параметри се во спецификација (освен "
                "системскиот pH). Лабораторискиот заклучок за овој примерок е: 'НЕ ОДГОВАРА'. На останатите тест денови "
                "(28.04, 30.04, 04.05, 05.05, 06.05, 07.05, 08.05, 11.05) истата точка пријави 0 CFU/100ml. | "
                "Sample 1203 (29.04.26, RO_C158_006 — RO water in irrigation room) reported 80 CFU/100ml faecal coliforms. "
                "All other micro params н.д.; all other params within spec (except systemic pH). Lab's verdict for this sample: "
                "'НЕ ОДГОВАРА' (does NOT pass). On all other test days (28.04, 30.04, 04.05, 05.05, 06.05, 07.05, 08.05, 11.05) "
                "the same point reported 0 CFU/100ml.",
                "Можни причини | Potential Root Causes",
                "1. Контаминација при земање на примерокот (грешка на оператор, нестерилен прибор, излагање на воздух). | "
                "Sampling contamination (operator error, non-sterile equipment, air exposure).",
                "2. Биофилм во крајниот сегмент од дистрибутивната линија до чешмата за иригација. | Biofilm in the terminal "
                "segment of the distribution line leading to the irrigation tap.",
                "3. Привремена бавна-проточна стагнација (просторијата за иригација можеби има пониска користеност на "
                "вода, дозволувајќи микробен раст). | Temporary low-flow stagnation (the irrigation room may have lower water "
                "usage, allowing microbial growth).",
                "4. Бек-сифонаж или крос-конекција од не-стерилен извор. | Back-siphonage or cross-connection from a non-sterile "
                "source.",
                "Заклучок | Conclusion",
                "Изолиран характер на наодот (1 ден од 9) сугерира контаминација при земање на примерокот или транзиторно "
                "настан, а не системски проблем на опремата. Сепак, наодот мора формално да се истражи. | "
                "Isolated nature of the finding (1 day of 9) suggests sampling contamination or a transient event rather than "
                "a systemic equipment issue. The finding must nonetheless be formally investigated.",
                "Препораки | Recommendations",
                "1. Инициирајте формална девијација и CAPA. | Initiate formal deviation and CAPA.",
                "2. Изведете ре-сампл на RO_C158_006 во рок од 7 дена, со посебно обучен оператор и стерилен прибор. | "
                "Perform a re-sample of RO_C158_006 within 7 days, by a specifically trained operator with sterile equipment.",
                "3. Прегледајте го најблискиот upstream сегмент за можен биофилм; разгледајте sanitization цикл (озон или "
                "топлинска дезинфекција). | Inspect the nearest upstream segment for possible biofilm; consider a sanitization "
                "cycle (ozone or thermal disinfection).",
                "4. Прегледајте ги SOP-овите за земање примероци и обуката на операторите. | Review sampling SOPs and operator "
                "training records."
            ]),
            ("4. Спецификација за амонијак под лабораториската способност на детекција | "
             "4. Ammonia specification below contract-lab detection capability", [
                "Преглед | Overview",
                "Спецификациите A03 (TW), A05 (TR) и A06 (RO) сите бараат амонијак < 0.0005 mg/L. CJZ Kumanovo пријавува "
                "'н.д.' за амонијак на сите 88 примероци, но нивниот пријавен Limit of Detection (U колона) е 0.016 mg/L. "
                "Тоа значи дека 'н.д.' документира < 0.016 mg/L, што е 32× повисоко од спецификацијата. | "
                "Specifications A03 (TW), A05 (TR) and A06 (RO) all require ammonia < 0.0005 mg/L. CJZ Kumanovo reports "
                "'н.д.' for ammonia on all 88 samples, but their reported Limit of Detection (U column) is 0.016 mg/L. "
                "That means 'н.д.' documents < 0.016 mg/L, which is 32× higher than the specification.",
                "Заклучок | Conclusion",
                "Со сегашниот контрактирен лабораториски метод, не е можно да се демонстрира compliance со спецификацијата "
                "за амонијак за која било од трите класи на вода. Ова е документациски/спецификациски расцеп — НЕ е дефект "
                "на опремата. | With the current contract-lab method, compliance with the ammonia specification cannot be "
                "demonstrated for any of the three water classes. This is a documentation/specification gap — NOT an "
                "equipment defect.",
                "Препораки | Recommendations",
                "1. Опција A: Релаксирајте ја спецификацијата за амонијак до < 0.05 mg/L (WHO drinking water guideline) или "
                "до Ph.Eur. Purified Water (нема директна граница за амонијак, се покрива со conductivity). | "
                "Option A: Relax ammonia spec to < 0.05 mg/L (WHO drinking water guideline) or to Ph.Eur. Purified Water "
                "(no direct ammonia limit, covered by conductivity).",
                "2. Опција B: Трансферирајте го амонијак тестирањето на лабораторија со ion chromatography или ISE метод со "
                "LOD под 0.001 mg/L. | Option B: Transfer ammonia testing to a lab with ion chromatography or ISE method with "
                "LOD below 0.001 mg/L.",
                "3. До тогаш, документирајте го јазот формално и оневозможете го pseudo-compliance на 'н.д.' = pass за амонијак. | "
                "Until then, document the gap formally and disable pseudo-compliance of 'н.д.' = pass for ammonia."
            ]),
            ("Општо Заклучок од Анализата на Резултатите | "
             "Overall Conclusion from Analysis of Test Results", [
                f"Систем за третман на вода со реверзна осмоза „Маган Мак“ (EQP-PPS002) беше тестиран во оперативни услови "
                f"во период од 9 тест-денови ({dates_sorted[0]} до {dates_sorted[-1]}) на {len(sps)} локации за мострирање. "
                f"Вкупно {n_samples} анализи се извршени од ЈЗУ Центар за јавно здравје — Куманово (CJZ Kumanovo), ISO/IEC "
                "17025 акредитирана лабораторија. | "
                f"The Reverse Osmosis Water Treatment System „Magan Mak“ (EQP-PPS002) was tested under operational conditions "
                f"over 9 test days ({dates_sorted[0]} to {dates_sorted[-1]}) across {len(sps)} sampling locations. A total of "
                f"{n_samples} analyses were performed by JZU Centar za Javno Zdravje — Kumanovo (CJZ Kumanovo), an ISO/IEC "
                "17025-accredited laboratory.",
                "Резиме на критичните наоди | Summary of critical findings:",
                "(1) ЕДЕН вистински дефект на опремата: RO_F117_010 conductivity 14–26× над спецификација на 6/9 тест-денови. "
                "Поправка е извршена помеѓу 06.05 и 07.05.26 (потребна верификација). | "
                "(1) ONE genuine equipment defect: RO_F117_010 conductivity 14–26× over spec on 6/9 test days. "
                "Repair appears to have been performed between 06.05 and 07.05.26 (verification needed).",
                "(2) ЕДЕН вистински контаминациски настан: sample 1203 (RO_C158_006, 29.04.26) — 80 CFU/100ml faecal "
                "coliforms на изолиран ден. | (2) ONE genuine contamination event: sample 1203 (RO_C158_006, 29.04.26) — "
                "80 CFU/100ml faecal coliforms on an isolated day.",
                "(3) Системски pH-наод (63 примероци) — расцеп помеѓу применетата спецификација и природата на демиерализирана "
                "вода; НЕ е дефект на опремата. Спецификацијата треба да биде ревидирана. | "
                "(3) Systemic pH finding (63 samples) — gap between the applied specification and the nature of demineralised "
                "water; NOT an equipment defect. The specification should be revised.",
                "(4) Спецификациско-капабилитетен расцеп за амонијак: лабораторискиот метод не може да го верифицира бараниот "
                "лимит. Спецификацијата треба да биде ревидирана или да се избере друга лабораторија. | "
                "(4) Specification-vs-capability gap for ammonia: lab method cannot verify the required limit. Specification "
                "should be revised or another lab selected.",
                "Заклучок за PQ1 | PQ1 Conclusion:",
                "Системот за реверзна осмоза „Маган Мак“ постигна репродуцибилност на квалитет на RO пермеат на повеќето "
                "GMP-критични локации. Сепак, два специфични дефекти (RO_F117_010 и единечен микробен инцидент) "
                "и два спецификациски проблеми мора да бидат решени пред финално PQ1 одобрување. | "
                "The „Magan Mak“ reverse osmosis system achieved reproducibility of RO permeate quality at most GMP-critical "
                "locations. However, two specific defects (RO_F117_010 and a single microbiological incident) and two "
                "specification issues must be resolved before final PQ1 approval."
            ]),
        ]
        # Insert in reverse order before anchor so they appear sequentially
        for heading, paras in reversed(sections):
            for para in reversed(paras):
                add_para_after(anchor.element if hasattr(anchor,'element') else anchor._p,
                               para)
            add_para_after(anchor.element if hasattr(anchor,'element') else anchor._p,
                           heading, bold=True)

    # ===== Fill the Summary/Conclusion paragraph (P102/P104 area) =====
    replace_in_paras(
        r"\[Дадете концизен наративен резиме[^\]]*\]",
        f"PQ1 на системот за реверзна осмоза „Маган Мак“ (EQP-PPS002) беше извршен во период од 9 тест-денови "
        f"({dates_sorted[0]} до {dates_sorted[-1]}) со {n_samples} анализи на {len(sps)} локации. Идентификувани се: "
        "1 критичен дефект на опремата (RO_F117_010 conductivity), 1 микробиолошки настан (sample 1203 faecal coliforms), "
        "и 2 спецификациски расцепи (pH за демиерализирана вода и амонијак-vs-LOD). Деталната анализа е во "
        "Дел „Анализа на резултатите од тестирање“ по Списокот на недостатоци. | "
        f"PQ1 for the „Magan Mak“ RO system (EQP-PPS002) was executed over 9 test days "
        f"({dates_sorted[0]} to {dates_sorted[-1]}) with {n_samples} analyses across {len(sps)} locations. Identified: "
        "1 critical equipment defect (RO_F117_010 conductivity), 1 microbiological event (sample 1203 faecal coliforms), "
        "and 2 specification gaps (pH for demineralised water and ammonia-vs-LOD). Detailed analysis is in the "
        "'Analysis of Test Results' section after the Deficiency List."
    )
    # Also clear the second example paragraph (P104)
    replace_in_paras(r"\[Пример за RO вода:[^\]]*\]", "")

    # ===== Fill Attachments table T18 =====
    t18 = doc.tables[18]
    tbl18 = t18._tbl
    trs18 = tbl18.findall(qn("w:tr"))
    # T18 has 9 rows (1 header + 8 placeholder). Clear placeholders and add ours.
    for tr in list(trs18[1:]):
        tbl18.remove(tr)
    attachments = [
        ("1", "Сумарна табела на резултати (Quali Rezults) | Results summary workbook",
         "PQ1_Quali_Rezults.xlsx"),
        ("2", "9 × A02 Water Testing Results Log — еден по тест-ден | one per test day",
         "A02_<date>_Water_Testing_Results_Log.docx (×9)"),
        ("3", "88 × A03/A05/A06 — Спецификација и резултати по примерок | Specification & results per sample",
         "A03/A05/A06_<date>_<SL>_sample-<NNNN>.docx (×88)"),
        ("4", "Сурови извештаи од договорна лабораторија (CJZ Kumanovo) | Raw contract lab reports",
         "PQ1_WQC/<NNNN> П.pdf (×88)"),
        ("5", "Спецификации за вода (TW/TR/RO) | Water specifications (TW/TR/RO)",
         "SOP QC Water — A03, A05, A06"),
        ("6", "Годишен план за тестирање на вода 2026 | 2026 Annual Water Testing Plan",
         "SOP QC Water — A01"),
        ("7", "Брз референтен водич за земање примероци | Sampling Quick-Reference Guide",
         "SOP QC Water — A08"),
        ("8", "Водечки SOP-ови за QC земање примероци | QC sampling SOPs",
         "PP-QC-SOP-017 v02; PP-QC-SOP-011 A01"),
    ]
    for no, desc, code in attachments:
        new_tr = deepcopy(trs18[0])
        for tc in new_tr.findall(qn("w:tc")):
            for p in tc.findall(qn("w:p")): tc.remove(p)
            tc.append(OxmlElement("w:p"))
        tbl18.append(new_tr)
        tcs = new_tr.findall(qn("w:tc"))
        if len(tcs)>0: set_tc_text(tcs[0], no, center=True)
        if len(tcs)>1: set_tc_text(tcs[1], desc)
        if len(tcs)>2: set_tc_text(tcs[2], code, center=True)

    # ===== T14 Overall Assessment — check the middle option (executed with deficiencies, open points blocked) =====
    t14 = doc.tables[14]
    # T14 R1 = "executed successfully without deficiencies"  — leave ☐
    # T14 R2 = "executed with deficiencies. Open deficiency points do not prevent..." — should be ☒
    # T14 R3 = "not successfully completed. Open deficiency points block subsequent..." — ☐
    # The text in each row is in C0; need a marker. Append our markings to the start of R2.
    # The template stores these as paragraphs; modify the first run.
    # We don't know exact OOXML cell content well; do a less-fragile thing: prepend ☒ note to R2 text.
    r2_tc = get_tc(t14, 2, 0)
    # Find first <w:t> and prepend "☒ " (and ☐ to R1, R3)
    def prefix_first_t(tc, prefix):
        for w_t in tc.iter(qn("w:t")):
            if w_t.text:
                w_t.text = prefix + w_t.text
                break
    r1_tc = get_tc(t14, 1, 0)
    r3_tc = get_tc(t14, 3, 0)
    prefix_first_t(r1_tc, "☐ ")
    prefix_first_t(r2_tc, "☒ ")
    prefix_first_t(r3_tc, "☐ ")

    # ===== T15 Critical Deficiencies Preventing Release =====
    # We have 7 critical-ish deficiencies but most are not "release-blocking" — only RO_F117_010 is.
    t15 = doc.tables[15]
    tbl15 = t15._tbl
    trs15 = tbl15.findall(qn("w:tr"))
    # Header row 0 + label row 1 + 3 placeholder rows (rows 2,3,4) + footer row 5? Let me check
    # We saw: R0 = "Critical Deficiencies Preventing Release" + "☐ N/A"
    #         R1 = "DP. | Short Summary | Date of rectification" (col headers)
    #         R2,3,4 = placeholders
    # Clear placeholders R2,3,4
    for tr in list(trs15[2:5]):
        tbl15.remove(tr)
    crits = [
        ("DP-PPS002-PQ-COND-F117", "Conductivity на RO_F117_010 14–26× над спецификација на 6 од 9 тест-денови (577–1051 µS/cm vs <40). "
                                   "Поправка изгледа дека е извршена помеѓу 06–07.05.26; потребна верификација. | "
                                   "Conductivity at RO_F117_010 14–26× over spec on 6 of 9 test days (577–1051 µS/cm vs <40). "
                                   "Repair appears to have been done between 06–07.05.26; verification needed.",
         "[Внесете датум по завршено CAPA] | [Enter date after CAPA closure]"),
        ("DP-PPS002-PQ-MICRO-1203", "Faecal coliforms 80 CFU/100ml на RO_C158_006 на 29.04.26 (sample 1203). "
                                    "Потребна е формална девијација, ре-сампл и истрага на можен биофилм/контаминација. | "
                                    "Faecal coliforms 80 CFU/100ml at RO_C158_006 on 29.04.26 (sample 1203). "
                                    "Formal deviation, re-sampling and biofilm/contamination investigation required.",
         "[Внесете датум по завршено истражување] | [Enter date after investigation closure]"),
    ]
    for dp, summary, date_rect in crits:
        new_tr = deepcopy(trs15[1])  # use the column-header row as the template
        for tc in new_tr.findall(qn("w:tc")):
            for p in tc.findall(qn("w:p")): tc.remove(p)
            tc.append(OxmlElement("w:p"))
        tbl15.append(new_tr)
        tcs = new_tr.findall(qn("w:tc"))
        if len(tcs)>0: set_tc_text(tcs[0], dp)
        if len(tcs)>1: set_tc_text(tcs[1], summary)
        if len(tcs)>2: set_tc_text(tcs[2], date_rect, center=True)

    # ===== T16: Approval for subsequent steps — uncheck PV release, check "No Release" =====
    t16 = doc.tables[16]
    # R2 C1 has the checkbox for "Approval for PV / commercial production" — leave ☐
    # R3 C1 has the checkbox for "No Release" — change to ☒
    # The template shows ☐ as plain text in those cells.
    no_release_tc = get_tc(t16, 3, 1)
    for w_t in no_release_tc.iter(qn("w:t")):
        if "☐" in (w_t.text or ""):
            w_t.text = w_t.text.replace("☐","☒",1)
            break

    doc.save(DOCX_OUT)
    print(f"✓ Wrote {DOCX_OUT} ({DOCX_OUT.stat().st_size//1024} KB)")
    print(f"  Tables filled: prerequisites (T3), test acceptance (T7), deficiencies (T11), attachments (T18), assessment (T14), critical deficiencies (T15), approval (T16)")
    print(f"  Analysis sections added: pH systemic, RO_F117_010 critical, sample 1203 contamination, NH3 spec gap, overall conclusion")

# ============= MAIN =============
def main():
    pdfs = sorted(PDFS_DIR.glob("*.pdf"))
    samples = []
    for p in pdfs:
        if "[1]" in p.name: continue
        s = parse_pdf(p)
        if s: samples.append(s)
    print(f"Parsed {len(samples)} samples")
    print(f"SL descriptions extracted: {len(SL_DESCRIPTIONS)}")
    for sp, d in sorted(SL_DESCRIPTIONS.items())[:5]:
        print(f"  {sp}: {d[:80]}")
    build_xlsx(samples)
    build_docx(samples)

main()
