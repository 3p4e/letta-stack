#!/usr/bin/env python3
"""Extend T1 (Definitions) in EQP-PPS002_PQ-V01_FILLED.docx with water/sampling terms."""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

DOC = "/tmp/EQP-PPS002_PQ-V01_FILLED.docx"

# 20 entries — 10 rows × 2 (abbr | full term) per side
NEW_DEFS = [
    # Row pairs (left-side, right-side)
    (("WSL",   "Water Sampling Location | Локација за земање примероци на вода"),
     ("SL",    "Sampling Location (синоним за WSL) | Sampling Location (synonym for WSL)")),
    (("TW",    "Municipality Tap Water | Градска Водоводна Вода"),
     ("TR",    "Treated Water — pre-RO | Третирана Вода — пред Р.О.")),
    (("RO",    "Reverse Osmosis (Demineralised) Water | Демиерализирана Вода од Реверзна Осмоза"),
     ("CPH",   "Center for Public Health (contract lab CJZ Kumanovo, ISO/IEC 17025) | Центар за Јавно Здравје — Куманово")),
    (("A.C.",  "Acceptance Criteria | Прифатливи критериуми"),
     ("OOS",   "Out of Specification | Надвор од Спецификација")),
    (("DP",    "Deficiency Point | Точка на недостаток"),
     ("LOD",   "Limit of Detection | Граница на детекција")),
    (("TBC",   "Total Bacterial Count | Вкупен број на бактерии"),
     ("CFU",   "Colony Forming Units | Колонии-формирачки единици")),
    (("NTU",   "Nephelometric Turbidity Units | NTU единици на матност"),
     ("Pt-Co", "Platinum-Cobalt colour scale | скала за боја")),
    (("ND",    "Not Detected (below LOD) | Не Детектирано (под границата на детекција, н.д.)"),
     ("RQS",   "Request for Quality Sampling (PP-QC-SOP-001 A01) | Барање за земање примероци")),
    (("SFR",   "Sampling Field Record (PP-QC-SOP-001 A02) | Полски запис за земање примероци"),
     ("STR",   "Sample Transport Record (PP-QC-SOP-001 A03) | Запис за транспорт на примерок")),
    (("A02",   "Water Testing Results Log — one per sampling day | Преглед на резултати — еден по тест-ден"),
     ("A03/A05/A06", "TW/TR/RO Specification & Test Results forms — one per sample | Спецификации TW/TR/RO")),
]

def main():
    d = Document(DOC)
    t = d.tables[1]  # T[1] = Definitions
    tbl = t._tbl
    trs = tbl.findall(qn("w:tr"))
    # Use one of the data rows (e.g. row 1) as the template for new rows
    template_row = trs[1]

    for left, right in NEW_DEFS:
        new_tr = deepcopy(template_row)
        # Clear each tc and add one paragraph with the new text
        tcs = new_tr.findall(qn("w:tc"))
        if len(tcs) < 4: continue
        for ci, val in enumerate([left[0], left[1], right[0], right[1]]):
            tc = tcs[ci]
            for p in tc.findall(qn("w:p")):
                tc.remove(p)
            p = OxmlElement("w:p")
            r = OxmlElement("w:r")
            t_el = OxmlElement("w:t")
            t_el.text = val
            t_el.set(qn("xml:space"), "preserve")
            r.append(t_el); p.append(r); tc.append(p)
        tbl.append(new_tr)

    d.save(DOC)
    # Verify
    d2 = Document(DOC)
    t2 = d2.tables[1]
    print(f"T[1] now has {len(t2.rows)} rows (was 12, added {len(NEW_DEFS)})")
    # Show new entries
    trs2 = t2._tbl.findall(qn("w:tr"))
    print("\nNewly added rows:")
    for tr in trs2[-len(NEW_DEFS):]:
        tcs = tr.findall(qn("w:tc"))
        if len(tcs) >= 4:
            a = "".join((x.text or "") for x in tcs[0].iter(qn("w:t"))).strip()
            b = "".join((x.text or "") for x in tcs[1].iter(qn("w:t"))).strip()
            c = "".join((x.text or "") for x in tcs[2].iter(qn("w:t"))).strip()
            e = "".join((x.text or "") for x in tcs[3].iter(qn("w:t"))).strip()
            print(f"  {a:8s} | {b[:50]:50s}    {c:14s} | {e[:50]}")

main()
