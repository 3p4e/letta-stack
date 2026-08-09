#!/usr/bin/env python3
"""
Generate filled DOCX forms from CPH test result PDFs.

Output:
  /tmp/forms/A02_<date>_Water_Testing_Results_Log.docx      (9 — one per sampling day)
  /tmp/forms/A0X_<date>_<SL>_sample-<NNNN>.docx              (89 — one per OK sample)
  /tmp/forms.zip                                              (bundled for delivery)

A0X = A03 for TW, A05 for TR, A06 for RO.
"""
import re, subprocess, pathlib, shutil, zipfile
from collections import defaultdict
from copy import deepcopy
from datetime import datetime
import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PDFS_DIR = pathlib.Path("/tmp/pq1_pdfs")
TMPL_DIR = pathlib.Path("/tmp/tmpl")
OUT_DIR  = pathlib.Path("/tmp/forms")
ZIP_OUT  = pathlib.Path("/tmp/PQ1_filled_forms.zip")

TRANS = str.maketrans({"Т":"T","Е":"E","О":"O","А":"A","С":"C","Р":"P","Х":"X","В":"B","М":"M","Н":"H","К":"K","У":"U"})
CPH_FIXUPS = {"C71_007":"C171_007","C71_008":"C171_008"}

# === Parameter catalogue (key + anchor pattern + unit) ===
PARAMS = [
    ("Faecal coliforms",         r"Колиформни бактерии од фекално потекло во \d+ ml",       "CFU/100ml"),
    ("Total bact. 37°C",         r"Вкупен број на бактерии на 37[°\s]*C во \d+ ml",           "CFU/ml"),
    ("Enterococci",              r"Ентерококи во \d+ ml",                                    "CFU/100ml"),
    ("Sulf.-red. clostridia",    r"Сулфито редуцирачки клостридии во \d+ ml",                "CFU/100ml"),
    ("Aerobes 22°C",             r"Вкупен број на аеробни бактерии на 22[°\s]*С во \d+ ml",  "CFU/ml"),
    ("Pseudomonas aeruginosa",   r"Pseudomonas aeruginosa во \d+ ml",                        "CFU/100ml"),
    ("E. coli",                  r"Е\.coli",                                                 "CFU/100ml"),
    ("Colour",                   r"Боја",                                                    "Pt-Co"),
    ("Turbidity",                r"Матност",                                                 "NTU"),
    ("Temperature",              r"температура",                                             "°C"),
    ("pH",                       r"^\s*рН\b",                                                ""),
    ("KMnO4 consumption",        r"Потрошувачка на KMnO4",                                   "mg/L"),
    ("Conductivity 20°C",        r"Електролитска спроводливост на 20[°\s]*[CС]",             "µS/cm"),
    ("Ammonia",                  r"Амонијак",                                                "mg/L"),
    ("Nitrites",                 r"Нитрити",                                                 "mg/L"),
    ("Nitrates",                 r"Нитрати",                                                 "mg/L"),
    ("Chlorides",                r"Хлориди",                                                 "mg/L"),
    ("Iron",                     r"Железо",                                                  "mg/L"),     # spec is in mg/L, lab in µg/L — convert!
    ("Manganese",                r"Манган",                                                  "mg/L"),     # same
]

# Spec limits per water class (min, max). None = unbounded.
LIMITS = {
    "Temperature": {"TW":(None,30),"TR":(None,30),"RO":(None,30)},
    "Colour":      {"TW":(None,20),"TR":(None,20),"RO":(None,20)},
    "Turbidity":   {"TW":(None,1.5),"TR":(None,1.5),"RO":(None,1.5)},
    "pH":          {"TW":(6.5,9.5),"TR":(6.5,9.5),"RO":(6.5,9.5)},
    "KMnO4 consumption":{"TW":(0,8),"TR":(0,8),"RO":(0,8)},
    "Conductivity 20°C":{"TW":(None,2500),"TR":(None,650),"RO":(None,40)},
    "Ammonia":     {"TW":(None,0.0005),"TR":(None,0.0005),"RO":(None,0.0005)},
    "Nitrites":    {"TW":(None,0.1),"TR":(None,0.1),"RO":(None,0.1)},
    "Nitrates":    {"TW":(None,50),"TR":(None,50),"RO":(None,50)},
    "Chlorides":   {"TW":(None,250),"TR":(None,250),"RO":(None,250)},
    "Iron":        {"TW":(None,0.2),"TR":(None,0.2),"RO":(None,0.2)},
    "Manganese":   {"TW":(None,0.05),"TR":(None,0.05),"RO":(None,0.05)},
    "Faecal coliforms":     {"TW":(None,0),"TR":(None,0),"RO":(None,0)},
    "Total bact. 37°C":     {"TW":(None,20),"TR":(None,20),"RO":(None,20)},
    "Enterococci":          {"TW":(None,0),"TR":(None,0),"RO":(None,0)},
    "Sulf.-red. clostridia":{"TW":(None,0),"TR":(None,0),"RO":(None,0)},
    "Aerobes 22°C":         {"TW":(None,100),"TR":(None,100),"RO":(None,100)},
    "Pseudomonas aeruginosa":{"TW":(None,0),"TR":(None,0),"RO":(None,0)},
    "E. coli":              {"TW":(None,0),"TR":(None,0),"RO":(None,0)},
}

# Row (in the A03/A05/A06 table) where each parameter lives → fill at C4 (result), C5 (doc), C6 (OOS)
PARAM_ROW = {
    "Temperature":4, "Colour":5, "Turbidity":6,
    "pH":10, "KMnO4 consumption":11, "Conductivity 20°C":12,
    "Ammonia":15, "Nitrites":16, "Nitrates":17, "Chlorides":18, "Iron":19, "Manganese":20,
    "Faecal coliforms":23, "Total bact. 37°C":24, "Enterococci":25,
    "Sulf.-red. clostridia":26, "Aerobes 22°C":27,
    "Pseudomonas aeruginosa":28, "E. coli":29,
}

RE_DATE = re.compile(r"Датум на земање:\s*(\d{2})\.(\d{2})\.(\d{4})")
RE_ISSUE = re.compile(r"Датум на издавање:\s*(\d{2})\.(\d{2})\.(\d{4})")
RE_LABNO = re.compile(r"Лаб\. број:\s*(\S+)")
RE_MM   = re.compile(r"Мерно место:\s*(.*?)(?=Вид\s|Странка\s|Датум на земање|Хигиено|$)", re.S)
RE_CODE = re.compile(r"\b([A-Z]{2})\s*[_-]\s*([A-Z]\d{1,4})\s*[_-]\s*(\d{2,4})(?!\d)")
RE_NUM  = re.compile(r"(?:н\.д\.|\d+(?:[.,]\d+)?)")

def text_of(p):
    return subprocess.run(["pdftotext","-layout",str(p),"-"],capture_output=True,text=True,timeout=30).stdout

def parse_value(s):
    s=s.strip()
    if s=="н.д.": return 0.0,"ND"
    s2=s.replace(",", ".")
    try: return float(s2), s
    except ValueError: return None, s

def extract_sample(p):
    """Returns dict with sample_id, date, sl, cph, water_class, doc_code, results{}."""
    sid = p.name.split()[0]
    if p.stat().st_size < 5000:
        return {"sample_id":sid, "stub":True}
    t = text_of(p)
    d = RE_DATE.search(t); mm = RE_MM.search(t); iss = RE_ISSUE.search(t); ln = RE_LABNO.search(t)
    if not (d and mm): return None
    date_iso = f"{d.group(3)}-{d.group(2)}-{d.group(1)}"
    date_eu = f"{d.group(1)}.{d.group(2)}.{d.group(3)}"
    issue = f"{iss.group(1)}.{iss.group(2)}.{iss.group(3)}" if iss else ""
    doc_code = ln.group(1) if ln else f"{sid}/2026"
    seg = mm.group(1).translate(TRANS).replace("\n"," ")
    m = RE_CODE.search(seg)
    if not m: return None
    sl, p1, p2 = m.groups()
    cph_raw = f"{p1}_{p2}"
    cph = CPH_FIXUPS.get(cph_raw, cph_raw)
    # Extract per-param values
    lines = t.splitlines()
    results = {}
    for key, pat, unit in PARAMS:
        rx = re.compile(pat)
        idx = None
        for i, ln_ in enumerate(lines):
            if rx.search(ln_):
                idx=i; break
        if idx is None:
            results[key] = {"raw":"", "num":None, "found":False, "unit":unit}
            continue
        scope = " ".join(lines[idx:idx+3])
        m_pat = rx.search(scope)
        post = scope[m_pat.end():]
        m_num = RE_NUM.search(post)
        if not m_num:
            results[key] = {"raw":"", "num":None, "found":False, "unit":unit}
            continue
        v, raw = parse_value(m_num.group(0))
        # Lab reports Iron/Manganese in µg/L — spec is in mg/L
        if key in ("Iron","Manganese") and v is not None and v != 0:
            v = v / 1000.0
        results[key] = {"raw":raw if raw=="ND" else f"{v}" if v is not None else raw,
                        "num":v, "found":True, "unit":unit}
    return {"sample_id":sid, "date_iso":date_iso, "date_eu":date_eu, "issue":issue,
            "doc_code":doc_code, "sl":sl, "cph":cph, "sp":f"{sl}_{cph}",
            "water_class":sl,  # TW/TR/RO is the class
            "results":results, "stub":False}

# ============= DOCX helpers =============
def get_tc(table, row_idx, tc_idx):
    """Return the row_idx-th <w:tr>'s tc_idx-th <w:tc>."""
    tbl = table._tbl
    trs = tbl.findall(qn("w:tr"))
    tcs = trs[row_idx].findall(qn("w:tc"))
    return tcs[tc_idx]

def set_tc_text(tc, text, *, bold=False, center=False, append=False):
    """Replace the text content of a <w:tc>. If append=True, add a new paragraph instead."""
    # Remove existing paragraphs unless appending
    if not append:
        for p in tc.findall(qn("w:p")):
            tc.remove(p)
    # Add one paragraph with one run
    p = OxmlElement("w:p")
    if center:
        pPr = OxmlElement("w:pPr")
        jc = OxmlElement("w:jc"); jc.set(qn("w:val"),"center")
        pPr.append(jc); p.append(pPr)
    r = OxmlElement("w:r")
    if bold:
        rPr = OxmlElement("w:rPr")
        b = OxmlElement("w:b"); rPr.append(b); r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text or ""
    t.set(qn("xml:space"),"preserve")
    r.append(t); p.append(r); tc.append(p)

def append_paragraph(tc, text):
    """Add text as a new paragraph in the cell (preserving existing content)."""
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t"); t.text = text; t.set(qn("xml:space"),"preserve")
    r.append(t); p.append(r); tc.append(p)

# ============= Spec doc (A03/A05/A06) filler =============
def fill_spec(sample, tmpl_path, out_path):
    doc = Document(tmpl_path)
    t = doc.tables[0]

    # 1) Header values
    # Row 1 (the header row with date/SL/sampled-by labels): values go in adjacent cells
    set_tc_text(get_tc(t, 1, 8), sample["date_eu"], center=True, bold=True)   # Date of Sampling value
    set_tc_text(get_tc(t, 1, 10), sample["sp"], center=True, bold=True)        # SL value
    # C12 (Sampled by signature) — leave blank

    # 2) Scope of Testing: row 4, C7 — replace placeholder text with one with the "Complete Water QC" box checked
    scope_tc = get_tc(t, 4, 7)
    # Clear existing paragraphs and rewrite with checked "Complete" box
    for p in scope_tc.findall(qn("w:p")):
        scope_tc.remove(p)
    scope_lines = [
        "☐ Физички Параметри | Physical Parameters",
        "☐ Физичко-хемиски Параметри | Physico-chemical Parameters",
        "☐ Хемиски Параметри | Chemical Parameters",
        "☐ Микробиолошки Параметри | Microbiological Parameters",
        "☒ Комплетна KK на Вода | Complete Water QC",
    ]
    for line in scope_lines:
        append_paragraph(scope_tc, line)

    # 3) Per-parameter rows: C4 = result, C5 = doc code, C6 = OOS checkbox
    cls = sample["water_class"]
    any_oos = False
    oos_params = []
    for key, row_idx in PARAM_ROW.items():
        r = sample["results"].get(key, {"raw":"","num":None,"found":False,"unit":""})
        # Determine OOS
        is_oos = False
        if r["found"] and r["num"] is not None:
            mn,mx = LIMITS[key][cls]
            if mn is not None and r["num"] < mn: is_oos = True
            if mx is not None and r["num"] > mx: is_oos = True
        # Build result string (include unit unless ND)
        if not r["found"]:
            disp = "—"
        elif r["raw"] == "ND":
            disp = "ND (н.д.)"
        else:
            # Format number; drop trailing zeros
            try:
                fv = float(r["raw"].replace(",", "."))
                disp = f"{fv:g} {r['unit']}".strip()
            except (ValueError, AttributeError):
                disp = f"{r['raw']} {r['unit']}".strip()
        set_tc_text(get_tc(t, row_idx, 4), disp, center=True,
                    bold=is_oos)
        set_tc_text(get_tc(t, row_idx, 5), sample["doc_code"], center=True)
        set_tc_text(get_tc(t, row_idx, 6), "☒" if is_oos else "☐", center=True,
                    bold=is_oos)
        if is_oos:
            any_oos = True
            oos_params.append((key, disp))

    # 4) Add OOS summary to Remarks (row 7, C8)
    if oos_params:
        notes = "OOS: " + "; ".join(f"{k}={v}" for k,v in oos_params)
        rem_tc = get_tc(t, 7, 8)
        # Clear and add new paragraph after the label
        for p in rem_tc.findall(qn("w:p")):
            rem_tc.remove(p)
        append_paragraph(rem_tc, "Забелешки: | Remarks:")
        append_paragraph(rem_tc, notes)

    doc.save(out_path)
    return any_oos, oos_params

# ============= A02 (per-day log) filler =============
def fill_a02(date_eu, date_iso, samples_for_day, out_path):
    doc = Document(TMPL_DIR / "A02.docx")
    t = doc.tables[0]

    # Header: row 0, C1 (Date of sampling) and C2 (Date of results receipt)
    # Both cells currently contain just the label; append a paragraph with the value
    dos_tc = get_tc(t, 0, 1)
    append_paragraph(dos_tc, date_eu)
    # Results receipt = max issue date across samples (or first if all same)
    receipt = max((s["issue"] for s in samples_for_day if s.get("issue")), default="")
    if receipt:
        dor_tc = get_tc(t, 0, 2)
        append_paragraph(dor_tc, receipt)

    # Row 1, C0 (Planned) — change ☐ to ☒
    planned_tc = get_tc(t, 1, 0)
    # Rewrite text replacing first ☐ with ☒
    for tt in planned_tc.iter(qn("w:t")):
        if "☐" in (tt.text or ""):
            tt.text = tt.text.replace("☐", "☒", 1)
            break

    # Data rows 4..28 — fill with one row per sample
    for i, s in enumerate(samples_for_day, start=1):
        if i > 25: break
        row_idx = 4 + (i - 1)
        set_tc_text(get_tc(t, row_idx, 0), f"{i}.", center=True)
        set_tc_text(get_tc(t, row_idx, 1), s["sp"], center=True)
        # Doc code column: include doc_code + issue date
        doc_text = f"{s['doc_code']} — {s['issue']}" if s.get('issue') else s['doc_code']
        set_tc_text(get_tc(t, row_idx, 2), doc_text, center=True)
        # Deviation checkbox stays ☐
        # OOS checkbox: check if any param OOS for this sample
        cls = s["water_class"]
        any_oos = False
        oos_list = []
        for key, row_in_spec in PARAM_ROW.items():
            r = s["results"].get(key, {})
            if r.get("found") and r.get("num") is not None:
                mn,mx = LIMITS[key][cls]
                num = r["num"]
                if (mn is not None and num < mn) or (mx is not None and num > mx):
                    any_oos = True
                    oos_list.append(key)
        set_tc_text(get_tc(t, row_idx, 4), "☒" if any_oos else "☐", center=True, bold=any_oos)
        # Ref doc no.: blank unless OOS, then note
        if any_oos:
            set_tc_text(get_tc(t, row_idx, 5), f"OOS: {', '.join(oos_list[:3])}", center=True)
        # Remarks
        if s.get("stub"):
            set_tc_text(get_tc(t, row_idx, 6), "EMPTY REPORT — CPH stub", center=True, bold=True)

    doc.save(out_path)

# ============= main =============
def main():
    if OUT_DIR.exists(): shutil.rmtree(OUT_DIR)
    OUT_DIR.mkdir(parents=True)

    # Parse all samples
    pdfs = sorted(PDFS_DIR.glob("*.pdf"))
    samples = []
    for p in pdfs:
        if "[1]" in p.name:  # skip duplicate
            continue
        s = extract_sample(p)
        if s: samples.append(s)
    print(f"Parsed {len(samples)} samples ({sum(1 for s in samples if s.get('stub'))} stubs)")

    # Group by date
    by_date = defaultdict(list)
    for s in samples:
        if s.get("stub"): continue
        by_date[(s["date_iso"], s["date_eu"])].append(s)

    # Render A02 + spec docs
    ok_specs = 0
    ok_oos = 0
    for (date_iso, date_eu), day_samples in sorted(by_date.items()):
        day_samples.sort(key=lambda s: (s["sl"], s["cph"], int(s["sample_id"])))
        # A02
        a02_name = f"A02_{date_iso}_Water_Testing_Results_Log.docx"
        fill_a02(date_eu, date_iso, day_samples, OUT_DIR / a02_name)
        # Per-sample spec
        for s in day_samples:
            cls = s["water_class"]
            tmpl_file = {"TW":"A03.docx","TR":"A05.docx","RO":"A06.docx"}.get(cls)
            if not tmpl_file: continue
            spec_name = f"{tmpl_file[:3]}_{date_iso}_{s['sp']}_sample-{s['sample_id']}.docx"
            try:
                oos, oos_list = fill_spec(s, TMPL_DIR/tmpl_file, OUT_DIR/spec_name)
                ok_specs += 1
                if oos: ok_oos += 1
            except Exception as e:
                print(f"  FAIL {spec_name}: {e!s:.150}")

    print(f"\nGenerated: {len(by_date)} A02s, {ok_specs} spec docs ({ok_oos} with OOS)")

    # Zip the lot
    with zipfile.ZipFile(ZIP_OUT, "w", zipfile.ZIP_DEFLATED) as z:
        for f in sorted(OUT_DIR.iterdir()):
            z.write(f, arcname=f.name)
    print(f"Bundled: {ZIP_OUT} ({ZIP_OUT.stat().st_size//1024} KB, {len(list(OUT_DIR.iterdir()))} files)")

main()
