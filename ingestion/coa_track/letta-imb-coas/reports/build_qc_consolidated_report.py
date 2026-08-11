#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Consolidated QC task progress and completion report, 10–11 August 2026.

Four separate workstreams ran over the two days and each produced its own record. This
merges them into one document so the department has a single picture of what was worked,
what closed, and what it produced:

  A  Specifications and controlled documents   9 tasks   (QC-LOG-001)
  B  Tranche 1 batch reconciliation            8 steps   (PP-QC-ADMIN-001)
  C  SP-12 pre-sale sampling campaign          3 tasks
  D  CoA register and specification revision  16 blocks

The report carries task content only — title, description, progression, outcome and
deliverables. No scope, method, data-availability or provenance chapters; where a
workstream did not record a duration the cell is left empty rather than estimated.

Built on the informal cover and header helpers from build_qc_activity_report, so the two
reports open the same way and neither claims a controlled-document number it has not been
assigned.
"""
import os
import sys

HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, HERE)
SKILL = "/root/.claude/skills/synced/pp-document-suite"
sys.path.insert(0, os.path.join(SKILL, "scripts"))

import pp_charts as pc                                              # noqa: E402
import pp_report as pr                                              # noqa: E402
from build_qc_activity_report import (informal_cover, no_split,     # noqa: E402
                                      set_header_title, strip_code_block)
from docx import Document                                           # noqa: E402

DOCX = os.path.join(HERE, "QC_Task_Progress_Report_10-11_Aug_2026.docx")
TEMPLATE = os.path.join(SKILL, "assets", "PP_BASE_TEMPLATE.docx")
ROWALT = "F7FAFC"

# (no, mk title, en title, mk description+progression, en description+progression,
#  mk outcome, en outcome, deliverable, done)
STREAM_A = [
 ("01", "Истражување и избор на шема на бои", "Colour-scheme exploration and selection",
  "Изработени и споредени четири шеми на бои (A–D). Избрана „Deep Gold“ како куќен стил "
  "за целиот сет документи.",
  "Four colour schemes (A–D) were built and compared. “Deep Gold” was selected as the "
  "house style for the whole document set.",
  "Завршено — усвоена шема D.", "Completed — Scheme D adopted.",
  "Color Scheme A–D .html", True),
 ("02", "Сертификат за квалитет (CoQ)", "Certificate of Quality (CoQ)",
  "Двојазично форматирање, усогласување на заглавие, подножје и лого, собирање на "
  "содржината на една A4 страница, извоз како самостоен HTML.",
  "Bilingual formatting, header, footer and logo alignment, content fitted onto a single "
  "A4 page, exported as standalone HTML.",
  "Завршено и испорачано.", "Completed and delivered.",
  "Certificate of Quality — CoQ-PP-2026-0005 (standalone).html", True),
 ("03", "Комплет iCoA — 6 документи", "iCoA set — 6 documents",
  "Двојазично преформатирање; правила за заглавие и подножје; редизајн на Дел 01 и "
  "Дел 03; наслов „Purely Plant — Certificate of Analysis“; секој документ собран на "
  "една A4.",
  "Bilingual reformat; header and footer rules; redesign of Section 01 and Section 03; "
  "title “Purely Plant — Certificate of Analysis”; each document fitted to one A4.",
  "Завршено — синхронизирано во Final_Docs.", "Completed — synced to Final_Docs.",
  "Final_Docs/ (6 × iCoA .html)", True),
 ("04", "Спецификација на производ — меѓупроизвод балк",
  "Product Specification — Intermediate Bulk",
  "Редизајн на банерот и полето за сорта; прераспределба на инфо-мрежата; табела на "
  "јачина; раздвојување на микробиолошките параметри (Salmonella и E. coli); CUMCS "
  "еквивалентност; пилули за сорта и обработка.",
  "Banner and cultivar field redesign; info-grid redistribution; potency table; "
  "microbiological parameters split (Salmonella and E. coli); CUMCS equivalency; variety "
  "and processing pills.",
  "Завршено.", "Completed.",
  "Product_Specification_Spec_IMB.html", True),
 ("05", "Генерирање PROD_SPEC — 35 сорти", "PROD_SPEC generation — 35 strains",
  "Создадени спецификации за сите 35 сорти од листата CULI-001; применети опсезите на "
  "јачина; кодови на производ SS_THCn:CBD1; двојазични критериуми низ целиот сет.",
  "Specifications generated for all 35 strains on the CULI-001 list; potency ranges "
  "applied; product codes SS_THCn:CBD1; bilingual criteria throughout.",
  "Завршено — ревизијата на усогласеност чиста.",
  "Completed — adherence audit clean.",
  "PROD_SPEC/ (35 × .html)", True),
 ("06", "Спецификации TRANCHE_1 — 13 сорти", "TRANCHE_1 specifications — 13 strains",
  "Подпапка TRANCHE_1; авторитативни скали на јачина од v5.5; вградено лого; секој "
  "документ собран на една A4.",
  "TRANCHE_1 sub-folder; authoritative v5.5 potency ladders; embedded logo; each "
  "document fitted to one A4.",
  "Завршено.", "Completed.",
  "PROD_SPEC/TRANCHE_1/ (13 × .html)", True),
 ("07", "Запис за порамнување на серии", "Batch reconciliation record",
  "Мапирана секоја серија од Транша 1 — комерцијално име, сорта, серија, пакувачки лот, "
  "Δ⁹-THC, класа, количина и број кеси — групирано по опсег на јачина, со меѓузбирови и "
  "вкупно 1 480,66 kg.",
  "Every Tranche 1 batch mapped — commercial name, strain, batch, packaging lot, "
  "Δ⁹-THC, grade, quantity and bag count — grouped by potency bracket, with subtotals "
  "and a 1,480.66 kg total.",
  "Завршено — подготвено за PDF.", "Completed — PDF-ready.",
  "PROD_SPEC/TRANCHE_1/Tranche1_Batch_Reconciliation.html", True),
 ("08", "REMAINING_SPC — 7 сорти, v5.6", "REMAINING_SPC — 7 strains, v5.6",
  "Создадени преостанатите седум ImB спецификации (AB, CLE, GRC, KC, MB, SJ, WC) со "
  "авторитативни скали од v5.6; ист шаблон и формат како остатокот од сетот.",
  "The remaining seven ImB specifications created (AB, CLE, GRC, KC, MB, SJ, WC) with "
  "authoritative v5.6 ladders; identical template and format to the rest of the set.",
  "Завршено.", "Completed.",
  "PROD_SPEC/REMAINING_SPC/ (7 × .html)", True),
 ("09", "Подготовка за испорака — HTML и PDF", "Delivery packaging — HTML and PDF",
  "Вградено лого за самостоен приказ; повеќекратни архиви за преземање; отворени "
  "дијалози за PDF извоз на записот за порамнување и на спецификациите.",
  "Logo embedded for standalone display; repeated download archives; PDF-export dialogs "
  "opened for the reconciliation record and the specifications.",
  "Во тек — по барање.", "Ongoing — on demand.",
  "PROD_SPEC.zip · печатени копии | print copies", False),
]

STREAM_B = [
 ("01", "Двојазична конверзија", "Bilingual conversion",
  "Целосна двојазична преработка на Rev.2, со посебен стил за македонскиот текст.",
  "Full bilingual rewrite of Rev.2, with distinct styling for the Macedonian text.",
  "Завршено.", "Complete.",
  "Tranche1_Batch_Reconciliation.html (двојазичен | bilingual)", True),
 ("02", "Усогласување на количините", "Quantity reconciliation",
  "Сите 21 количини ажурирани од складишната евиденција и вкрстено проверени со Транша "
  "1. GG1024_01 задржана на својата распределба, не на целиот лот.",
  "All 21 quantities updated from the warehouse inventory and cross-checked against "
  "Tranche 1. GG1024_01 held at its allocation rather than the full lot.",
  "Завршено — збировите потврдени.", "Complete — sums verified.",
  "Rev.3 .html и | and .pdf", True),
 ("03", "Самостоен PDF пренос", "Self-contained PDF pipeline",
  "Вградени фонтови (base64) и рендерирање преку headless Chrome/Playwright, за "
  "документот да изгледа исто без надворешни ресурси.",
  "Fonts embedded (base64) and rendering through headless Chrome/Playwright, so the "
  "document looks the same without external resources.",
  "Завршено.", "Complete.",
  "-print.html и | and .pdf", True),
 ("04", "Чистење на заглавието и јазикот на резултатите",
  "Header and result-language cleanup",
  "Отстранет текстот QCSP/Rev./датум од заглавието, датумот „прва испорака“ и целиот "
  "јазик за „водечки резултат“.",
  "Removed the QCSP/Rev./date header text, the “first delivery” date, and all "
  "“leading result” language.",
  "Завршено.", "Complete.",
  "ажурирани | updated .html/.pdf", True),
 ("05", "Прецизност на децималите", "Decimal-precision fix",
  "Количините коригирани од две на три децимали. При проверката е фатена и поправена "
  "регресија кај HPA1024_01.",
  "Quantities corrected from two decimals to three. A regression on HPA1024_01 was "
  "caught and fixed during the check.",
  "Завршено — проверено ред по ред.", "Complete — verified row by row.",
  "ажурирани | updated .html/.pdf", True),
 ("06", "Прелевање на бојата на заглавието", "Header background bleed",
  "Итеративна имплементација на позадина што се прелева 2 mm без да ја помести "
  "содржината. При работата е фатен тивок дефект во скриптата за градење.",
  "Iterative implementation of a 2 mm background bleed without shifting the content. A "
  "silent defect in the build script was caught in the process.",
  "Завршено — визуелно потврдено.", "Complete — visually verified.",
  "ажурирани | updated .html/.pdf", True),
 ("07", "Конзистентни маргини на страниците", "Consistent page margins",
  "Поправена недостасувачката маргина на внатрешниот прекин на страница, преку "
  "CSS @page :first.",
  "The missing margin at the internal page break fixed, through CSS @page :first.",
  "Завршено.", "Complete.",
  "ажурирани | updated .html/.pdf", True),
 ("08", "Потписен блок и распоред", "Signature block and layout",
  "Додаден двојазичен блок за одобрување со четири улоги, распореден 2 × 2; додаден "
  "простор по секој меѓузбир.",
  "Bilingual approval block with four roles added, arranged 2 × 2; spacing added after "
  "each subtotal.",
  "Завршено.", "Complete.",
  "финален | final .pdf", True),
]

STREAM_C = [
 ("01", "Дизајн на кампањата за земање примероци SP-12",
  "SP-12 sampling campaign design",
  "Дефинирана намалената испитна група и дизајниран планот за композитно земање "
  "примероци за кампањата за верификација пред продажба SP-12, која опфаќа 21 "
  "складирана серија Cannabis flos. Заклучени клучните одлуки за дизајнот — архитектура "
  "од три SOP, намалена испитна група и големини на примероци за чување. Изработен "
  "доказ за етикета од фаза 1 и табели со големини на примерок по r-план, по серија.",
  "The reduced testing panel was defined and the composite sampling plan designed for "
  "the SP-12 pre-sale verification campaign covering 21 stored Cannabis flos batches. "
  "Key design decisions were locked — a three-SOP architecture, the reduced panel, and "
  "retention sample sizes. A Phase 1 label proof and r-plan sample-size tables per batch "
  "were built.",
  "Фаза 1 завршена; во очекување на одобрување на форматот. Фазите 2–5 во тек.",
  "Phase 1 complete; awaiting format approval. Phases 2–5 in progress.",
  "SP-12_Campaign_Build_Specification.md · "
  "PP_QC_Analysis_Sample_Labels_Sheet1.docx/.pdf", False),
 ("02", "Планови за микробиолошко земање примероци, ден 1",
  "Day-1 microbiological sampling plans",
  "Изработен планот за микробиолошко земање примероци за ден 1 за три приоритетни серии "
  "од SP-12 — HPA1024_01, HPA1024 и OPM1024 — со големини на примероци по r-план и "
  "симулирани избори k#B#.",
  "The day-1 microbiological sampling plan was produced for three priority SP-12 "
  "batches — HPA1024_01, HPA1024 and OPM1024 — with r-plan sample sizes and simulated "
  "k#B# picks.",
  "Завршено за сериите од ден 1; останува контролираниот план во .docx.",
  "Complete for the day-1 batches; the controlled plan in .docx remains outstanding.",
  "Табела по серија со N/n/k и целосни k#B# избори | "
  "Per-batch table with N/n/k and full k#B# picks", True),
 ("03", "План за земање примероци GP0824_02", "GP0824_02 sampling plan",
  "Генериран план за земање примероци по r-план за серијата GP0824_02 од 200 kg: вкупен "
  "број кеси, кеси за земање примерок, интервал на земање и симулирани избори k#B#.",
  "An r-plan sampling plan generated for the 200 kg GP0824_02 batch: total bags, bags to "
  "sample, sampling interval, and simulated k#B# picks.",
  "Завршено според ажурираната големина на серијата од 200 kg.",
  "Complete against the updated 200 kg batch size.",
  "План за земање примероци со клучните параметри и симулација на кеси | "
  "Sampling plan with key parameters and bag simulation", True),
]

D_TIME = ["10.08 19:22", "10.08 20:26", "10.08 21:28", "11.08 00:19", "11.08 01:56",
          "11.08 02:42", "11.08 03:23", "11.08 04:11", "11.08 05:53", "11.08 06:02",
          "11.08 06:16", "11.08 06:35", "11.08 06:58", "11.08 09:54", "11.08 10:50",
          "11.08 11:04"]           # real transcript timestamps, see the QC activity report
STREAM_D = [
 ("01", "Внесување на нови надворешни CoA", "Ingestion of new outsourced CoA",
  "Нови сертификати од надворешни акредитирани лаборатории проверени за дупликати "
  "наспроти постојниот извор, утврдено дали носат текстуален слој или се скенирани, "
  "качени и вградени во базата на знаење.",
  "New certificates from outsourced accredited laboratories deduplicated against the "
  "live source, text-layer versus scan determined, uploaded and embedded into the "
  "knowledgebase.",
  "Завршено.", "Complete.",
  "Вградени датотеки во source-271bc3be | Files embedded into source-271bc3be", True),
 ("02", "Опфаќање на целиот панел параметри", "Capture of the full parameter panel",
  "Δ⁹-THC и LoD не беа единствените параметри на новите сертификати, па екстракцијата е "
  "проширена на целиот панел.",
  "Δ⁹-THC and LoD were not the only parameters on the new certificates, so extraction "
  "was widened to the whole panel.",
  "Завршено.", "Complete.",
  "42 сертификати од серијата Farmahem 197, извлечени verbatim | "
  "42 Farmahem 197-series certificates extracted verbatim", True),
 ("03", "Реорганизација на структурата на папки по серии",
  "Rebuild of the batch-folder structure",
  "Хронолошки папки по серија, пресликани во базата на знаење; отстранети дупликати; "
  "интерните CoA на Purely Plant повлечени од копирањето; нумеричкиот префикс отстранет "
  "од имињата на папките.",
  "Chronological per-batch folders, mirrored into the knowledgebase; duplicates removed; "
  "Purely Plant in-house CoAs withdrawn from the copy; the numeric prefix dropped from "
  "the folder names.",
  "Завршено по четири корекции.", "Complete after four corrections.",
  "40 папки по серија · збирка OLD_PP_CoA_ · 25 CoA од категорија A | "
  "40 per-batch folders · OLD_PP_CoA_ collection · 25 category-A CoAs", True),
 ("04", "Дизајн на главната табела со CoA", "Design of the master CoA table",
  "Изградена главната табела во распоред „еден ред по серија“, со екстракција во два "
  "дела.",
  "The master table built in a row-per-batch layout, with extraction run in two slices.",
  "Завршено.", "Complete.",
  "master_coa_table.tsv · PP_eCoA_Master_Database.xlsx · coa_register.html", True),
 ("05", "Ревизија на регистарот", "Audit of the register",
  "Проверен квалитетот на внесените резултати: отстранета мерната неодреденост од "
  "ќелиите, раздвоени двојните резултати за јачина.",
  "The quality of the ingested results was checked: measurement uncertainty stripped "
  "from the cells, dual potency results split apart.",
  "Завршено — најдените недостатоци отстранети.",
  "Complete — the defects found were corrected.",
  "Куќно правило: вредноста сама во ќелијата, единицата во заглавието | "
  "House rule: the value alone in the cell, the unit in the header", True),
 ("06", "Усогласување на критериумите за прифатливост",
  "Normalisation of acceptance criteria",
  "LoD беше запишан како ≤ 10 % на едни сертификати и < 12 % на други; усвоен Ph. Eur. "
  "07/2024:3028 (≤ 12,0 %) како важечки критериум. Микотоксините се водат како три "
  "параметри, со „не е испитано“ таму каде што не се испитани.",
  "LoD was recorded as ≤10% on some certificates and <12% on others; Ph. Eur. "
  "07/2024:3028 (≤12.0%) adopted as the governing criterion. Mycotoxins are carried as "
  "three parameters, marked not tested where not assayed.",
  "Завршено.", "Complete.",
  "Коригирана колона со критериуми | Corrected acceptance-criteria column", True),
 ("07", "Преобликување на регистарот во куќниот распоред",
  "Reshaping of the register to the house layout",
  "Подредување по време на производство наместо по сорта, со параметрите како колони.",
  "Ordered by time of production rather than by cultivar, with parameters as columns.",
  "Завршено.", "Complete.",
  "PP_Production_Batches_QC_Result_Table.xlsx · PP_THC_by_Strain.xlsx", True),
 ("08", "Следливост на резултатите до сертификатите",
  "Traceability of results to certificates",
  "Секој резултат потврден наспроти сертификатот и кодот од кој потекнува. Четири ќелии "
  "со код на сертификат носеа претпоставки од OCR, а една содржеше текст наместо број "
  "на документ — сите вратени на изворниот запис.",
  "Every result confirmed against the certificate and code it came from. Four "
  "Certificate Code cells carried OCR guesses and one held text where a document number "
  "belonged — all restored to the source record.",
  "Завршено — 47 реда вратени verbatim.",
  "Complete — 47 rows restored verbatim.",
  "restore_verbatim_cert_codes.py · PP_Batch_Release_QC_Register.xlsx", True),
 ("09", "Вкрстена проверка наспроти контролните табели",
  "Cross-check against the control sheets",
  "Надворешниот CoA утврден како извор на вистина, што ги позиционира трите контролни "
  "табели како контролни записи, а не како сертификати.",
  "The outsourced CoA established as the source of truth, which positions the three "
  "control sheets as control records rather than certificates.",
  "Завршено — отстапувањата сведени од 93 на 53.",
  "Complete — discrepancies reduced from 93 to 53.",
  "crosscheck_sources.py · PP_Open_Reconciliation_Items.xlsx", True),
 ("10", "Пронаоѓање на недостасувачките сертификати",
  "Hunt for the missing certificates",
  "Три сертификати пријавени како недостасувачки всушност постоеја како PDF составени "
  "само од слики, невидливи и за пребарување по име и по текст: GG1024, HPA1024 и "
  "OPM1024.",
  "Three certificates reported missing in fact existed as image-only PDFs, invisible to "
  "both filename and full-text search: GG1024, HPA1024 and OPM1024.",
  "Завршено — трите заведени.", "Complete — all three filed.",
  "3 R&D сертификати · fix_rnd_drive_links.py (87 реда) | "
  "3 R&D certificates · fix_rnd_drive_links.py (87 rows)", True),
 ("11", "Поставување OCR на Letta host", "Standing up OCR on the Letta host",
  "OCR конфигуриран на Letta host за сертификатите што се скенирани слики, со "
  "вградување на потребните библиотеки.",
  "OCR configured on the Letta host for the certificates that are scanned images, with "
  "the required libraries vendored in.",
  "Завршено — регистриран и тестиран.", "Complete — registered and smoke-tested.",
  "pp_ocr_scanned_pdf.py", True),
 ("12", "Извлекување референци за CoA од нацртите за CoQ",
  "Recovery of CoA references from the CoQ drafts",
  "Од 16 прекинати и неодобрени нацрти за CoQ земени само кодот на документот, датумот "
  "на издавање и институцијата — ниту една аналитичка вредност.",
  "From 16 discontinued and unapproved CoQ drafts, only the document code, issue date "
  "and institution were taken — no analytical value.",
  "Завршено — идентификувани околу 34 сертификати што постојат но не се заведени.",
  "Complete — some 34 certificates identified that exist but are not filed.",
  "coq_draft_coa_references.tsv", True),
 ("13", "Усогласување на имињата на сортите и јачина за Транша 1",
  "Cultivar spellings and Tranche 1 potency",
  "Cap Junky потврден како точниот запис на сортата, Permanent Marker исто така "
  "исправен. Издадена јачината за Транша 1 како PDF.",
  "Cap Junky confirmed as the correct cultivar spelling, Permanent Marker corrected "
  "too. Tranche 1 potency issued as a PDF.",
  "Завршено — 21 од 21 серии, 13 сорти, 18 повторно испитани.",
  "Complete — 21 of 21 batches, 13 cultivars, 18 retested.",
  "normalise_strain_names.py · PP_Tranche1_Potency.pdf", True),
 ("14", "Изработка на спецификациите за производ ImB",
  "Building the ImB product specifications",
  "Изработени спецификации за сите сорти и обединети во еден документ; притоа откриени "
  "два резултати надвор од скалата на својата сорта.",
  "Specifications built for every cultivar and consolidated into one document; two "
  "results were found outside their cultivar's ladder in the process.",
  "Испорачано, потоа заменето со куќниот формат.",
  "Delivered, then superseded by the house format.",
  "Обединет документ со спецификации | Consolidated specification document", True),
 ("15", "Усвојување на куќниот формат QCSP 001",
  "Adoption of the QCSP 001 house format",
  "Издадените спецификации по сорта усвоени како формат; изведените скали се совпаѓаат "
  "точно со нив, вклучително и неправилни граници како 13,68 % кај Fat Bastard.",
  "The issued per-cultivar specifications adopted as the format; the derived ladders "
  "match them exactly, including irregular floors such as Fat Bastard's 13.68%.",
  "Завршено — инвентар на сите 35 издадени документи.",
  "Complete — all 35 issued documents inventoried.",
  "house_spec_inventory.tsv", True),
 ("16", "Проширување на скалите за јачина", "Extension of the potency ladders",
  "Долните граници на скалите беа поставени точно 1,00 процентен поен под најниската "
  "тогашна серија, поради што две серии останаа без класа. Секоја скала проширена "
  "надолу по сопствениот чекор.",
  "The ladder floors had been set exactly 1.00 percentage point below the lowest batch "
  "of the time, which left two batches with no grade. Every ladder extended downward on "
  "its own step.",
  "Завршено — 33 класи додадени кај 19 сорти; сите 99 испитувања сега паѓаат во класа.",
  "Complete — 33 grades added across 19 cultivars; all 99 assays now land in a grade.",
  "20 HTML и 20 PDF спецификации · PP_Potency_Grade_Audit.tsv | "
  "20 HTML and 20 PDF specifications · PP_Potency_Grade_Audit.tsv", True),
]

for _i, _t in enumerate(D_TIME):
    STREAM_D[_i] = STREAM_D[_i] + (_t,)   # append real clock time, index 9

STREAMS = [
    ("A", "Спецификации и контролни документи",
     "Specifications and controlled documents", STREAM_A, None, None),
    ("B", "Порамнување на серии — Транша 1", "Tranche 1 batch reconciliation",
     STREAM_B, None, None),
    ("C", "Кампања за земање примероци SP-12", "SP-12 sampling campaign",
     STREAM_C, "≈ 5,5 h", "≈ 5.5 h"),
    ("D", "Регистар на CoA и ревизија на спецификации",
     "CoA register and specification revision", STREAM_D, "11,2 h", "11.2 h"),
]

OPEN_ITEMS = [
 ("31 серии без CoQ, со неконзистентности во именувањето, утврдени при ревизијата на "
  "датотеките со CoA.",
  "31 batches without a CoQ, with naming inconsistencies, established during the CoA "
  "file audit."),
 ("Околу 34 сертификати од надворешни лаборатории постојат и се именувани, но не се "
  "заведени во регистарот.",
  "Some 34 outsourced-laboratory certificates exist and are named, but are not filed in "
  "the register."),
 ("Две серии беа испитани под скалата на својата сорта — CC112501 на 13,35 % и FB032601 "
  "на 12,39 %. Проширувањето на скалата важи за наредните серии и не ја менува нивната "
  "оцена по спецификацијата што важела при испитувањето.",
  "Two batches assayed below their cultivar's ladder — CC112501 at 13.35% and FB032601 "
  "at 12.39%. The ladder extension governs subsequent batches and does not change their "
  "assessment against the specification in force when they were tested."),
 ("Фазите 2–5 од изградбата на кампањата SP-12 се во тек, во очекување на одобрување на "
  "форматот.",
  "Phases 2–5 of the SP-12 campaign build are in progress, awaiting format approval."),
 ("Контролираниот план за земање примероци во .docx за сериите од ден 1 останува да се "
  "изработи.",
  "The controlled sampling plan in .docx for the day-1 batches remains to be produced."),
]


def chart_streams(out):
    import matplotlib.pyplot as plt
    labels = ["%s" % s[0] for s in STREAMS]
    done = [sum(1 for t in s[3] if t[8]) for s in STREAMS]
    open_ = [sum(1 for t in s[3] if not t[8]) for s in STREAMS]
    fig, ax = plt.subplots(figsize=(8.4, 2.9))
    ax.bar(labels, done, color=pc.NAVY, width=0.55, label="Completed")
    ax.bar(labels, open_, bottom=done, color="#E67E22", width=0.55, label="In progress")
    for i, (a, b) in enumerate(zip(done, open_)):
        ax.text(i, a + b + 0.35, str(a + b), ha="center", fontsize=9, color=pc.GREY)
    pc._style(ax, "Tasks by workstream", "Workstream", "Tasks")
    ax.set_ylim(0, max(d + o for d, o in zip(done, open_)) + 2.2)
    ax.legend(fontsize=8, frameon=False, loc="upper left")
    return pc._save(fig, out)


def stream_chapter(d, num, code, mk_name, en_name, tasks, mk_time, en_time):
    pr.chapter(d, num, "%s. %s" % (code, mk_name.upper()), "%s. %s" % (code, en_name))
    done = sum(1 for t in tasks if t[8])
    extra_mk = " · утрошено време %s" % mk_time if mk_time else ""
    extra_en = " · time spent %s" % en_time if en_time else ""
    pr.body(d,
            "%d задачи, од кои %d завршени.%s" % (len(tasks), done, extra_mk),
            "%d tasks, of which %d completed.%s" % (len(tasks), done, extra_en))
    for i, task in enumerate(tasks, 1):
        no, mk_t, en_t, mk_d, en_d, mk_o, en_o, deliv, ok = task[:9]
        pr.subsec(d, "%s.%d" % (num, i), "%s — %s" % (no, mk_t), "%s — %s" % (no, en_t))
        rows = [("Опис и прогресија | Description and progression", mk_d, en_d),
                ("Исход | Outcome", mk_o, en_o),
                ("Испорака | Deliverable", deliv, None)]
        t = d.add_table(rows=len(rows), cols=2)
        for r, (lab, m_, e_) in enumerate(rows):
            pr.cellfmt(t.cell(r, 0), lab, None, 9, pr.BLACK, bold=True, fill=pr.LBL)
            val = "%s | %s" % (m_, e_) if e_ else m_
            pr.cellfmt(t.cell(r, 1), val, None, 9, pr.BLACK)
        pr.fixed(t, [4.0, 14.46], header_repeat=False); pr.borders(t); no_split(t)


def main():
    all_tasks = sum(len(s[3]) for s in STREAMS)
    done = sum(sum(1 for t in s[3] if t[8]) for s in STREAMS)
    fig = chart_streams(os.path.join(HERE, "_fig_streams.png"))

    d = Document(TEMPLATE)
    set_header_title(d, "ИЗВЕШТАЈ ЗА НАПРЕДОК И ЗАВРШЕНОСТ НА ЗАДАЧИ — КОНТРОЛА НА КВАЛИТЕТ",
                     "Task Progress and Completion Report — Quality Control")
    strip_code_block(d)
    pr.wipe_body(d)

    informal_cover(
        d,
        "Извештај за напредок и завршеност на задачи — Оддел за контрола на квалитет",
        "Task Progress and Completion Report — Quality Control Department",
        [("Период | Reporting period", "10–11.08.2026"),
         ("Оддел | Department", "Контрола на квалитет | Quality Control"),
         ("Работни текови | Workstreams", "4"),
         ("Дистинктни задачи | Distinct tasks", "%d" % all_tasks),
         ("Завршени | Completed", "%d од | of %d" % (done, all_tasks)),
         ("Во тек | In progress", "%d" % (all_tasks - done))],
        "ПРЕГЛЕД НА ЗАДАЧИ", "TASK OVERVIEW",
        "4 работни текови · %d задачи · %d завршени" % (all_tasks, done),
        "4 workstreams · %d tasks · %d completed" % (all_tasks, done))
    pr.toc_page(d)

    pr.chapter(d, "1", "РЕЗИМЕ", "Summary")
    pr.body(d,
            "Во периодот 10–11.08.2026 Одделот за контрола на квалитет работеше во четири "
            "паралелни работни текови и заврши %d од вкупно %d дистинктни задачи."
            % (done, all_tasks),
            "Over 10–11 August 2026 the Quality Control Department worked across four "
            "parallel workstreams and closed %d of %d distinct tasks."
            % (done, all_tasks))
    pr.body(d,
            "Работата опфати комплетен сет спецификации за производ, сертификати за "
            "анализа и квалитет, порамнување на сериите од Транша 1 со вкупно 1 480,66 kg, "
            "дизајн на кампањата за земање примероци пред продажба SP-12, и целосна "
            "ревизија на регистарот на CoA со проширување на скалите за јачина.",
            "The work covered a complete set of product specifications, certificates of "
            "analysis and of quality, reconciliation of the Tranche 1 batches totalling "
            "1,480.66 kg, design of the SP-12 pre-sale sampling campaign, and a full audit "
            "of the CoA register with an extension of the potency ladders.")
    pr.figure(d, fig, "Задачи по работен тек", "Tasks by workstream")

    pr.chapter(d, "2", "ХРОНОЛОШКИ РЕДОСЛЕД", "Chronological Order")
    pr.body(d,
            "Само тек D носи стварни часовни ознаки, земени од дневникот на сесијата "
            "во кој се извршени неговите задачи; тие се прикажани подолу непроменети. "
            "Текови A, B и C немаат снимен час по задача во своите изворни записи — тоа "
            "е наведено во самите тие записи, а не проценето овде. За нив е прикажан "
            "само редоследот на извршување што го наведуваат нивните извори, без да им "
            "се припише час што не постои.",
            "Only Stream D carries real clock times, taken from the session log its "
            "tasks were executed in; they are shown below unchanged. Streams A, B and C "
            "carry no recorded time per task in their own source records — that is "
            "stated in those records themselves, not estimated here. For them, only the "
            "execution order their sources state is shown, without assigning a clock "
            "time that does not exist.")

    hdr = ["Час | Time", "Тек | Stream", "Задача | Task"]
    rows_d = [(D_TIME[i], "D", "%s — %s | %s" % (t[0], t[1], t[2]))
              for i, t in enumerate(STREAM_D)]
    t = d.add_table(rows=len(rows_d) + 1, cols=3)
    for j, x in enumerate(hdr):
        pr.cellfmt(t.cell(0, j), x, None, 9, pr.WHITE, bold=True, fill=pr.NAVYF)
    for i, (time_, code, label) in enumerate(rows_d, start=1):
        fill = ROWALT if i % 2 == 0 else None
        pr.cellfmt(t.cell(i, 0), time_, None, 9, pr.BLACK, fill=fill)
        pr.cellfmt(t.cell(i, 1), code, None, 9, pr.BLACK, bold=True, fill=fill)
        pr.cellfmt(t.cell(i, 2), label, None, 8, pr.BLACK, fill=fill)
    pr.fixed(t, [2.6, 1.4, 14.46]); pr.borders(t); no_split(t)

    pr.subsec(d, "2.1", "Текови без снимен час — само редослед на извршување",
              "Streams without a recorded time — execution order only")
    for code, _mk_n, _en_n, tasks, _mt, _et in STREAMS:
        if code == "D":
            continue
        hdr2 = ["№", "Тек %s | Stream %s" % (code, code)]
        tt = d.add_table(rows=len(tasks) + 1, cols=2)
        for j, x in enumerate(hdr2):
            pr.cellfmt(tt.cell(0, j), x, None, 9, pr.WHITE, bold=True, fill=pr.NAVYF)
        for i, task in enumerate(tasks, start=1):
            fill = ROWALT if i % 2 == 0 else None
            pr.cellfmt(tt.cell(i, 0), task[0], None, 9, pr.BLACK, fill=fill)
            pr.cellfmt(tt.cell(i, 1), "%s | %s" % (task[1], task[2]), None, 9, pr.BLACK,
                       fill=fill)
        pr.fixed(tt, [1.2, 17.26]); pr.borders(tt); no_split(tt)

    pr.chapter(d, "3", "ПРЕГЛЕД ПО ТЕКОВИ", "Overview by Stream")
    hdr = ["Тек | Stream", "Работен тек | Workstream", "Задачи | Tasks",
           "Завршени | Done", "Време | Time"]
    t = d.add_table(rows=len(STREAMS) + 2, cols=5)
    for j, x in enumerate(hdr):
        pr.cellfmt(t.cell(0, j), x, None, 9, pr.WHITE, bold=True, fill=pr.NAVYF)
    for i, (code, mk_n, en_n, tasks, mk_time, en_time) in enumerate(STREAMS, start=1):
        fill = ROWALT if i % 2 == 0 else None
        nd = sum(1 for x in tasks if x[8])
        pr.cellfmt(t.cell(i, 0), code, None, 9, pr.BLACK, bold=True, fill=fill)
        pr.cellfmt(t.cell(i, 1), "%s | %s" % (mk_n, en_n), None, 9, pr.BLACK, fill=fill)
        pr.cellfmt(t.cell(i, 2), "%d" % len(tasks), None, 9, pr.BLACK, fill=fill)
        pr.cellfmt(t.cell(i, 3), "%d" % nd, None, 9, pr.BLACK, fill=fill)
        pr.cellfmt(t.cell(i, 4), mk_time or "—", None, 9, pr.BLACK, fill=fill)
    r = len(STREAMS) + 1
    pr.cellfmt(t.cell(r, 0), "", None, 9, pr.BLACK, bold=True, fill=pr.LBL)
    pr.cellfmt(t.cell(r, 1), "ВКУПНО | TOTAL", None, 9, pr.BLACK, bold=True, fill=pr.LBL)
    pr.cellfmt(t.cell(r, 2), "%d" % all_tasks, None, 9, pr.BLACK, bold=True, fill=pr.LBL)
    pr.cellfmt(t.cell(r, 3), "%d" % done, None, 9, pr.BLACK, bold=True, fill=pr.LBL)
    pr.cellfmt(t.cell(r, 4), "≈ 16,7 h", None, 9, pr.BLACK, bold=True, fill=pr.LBL)
    pr.fixed(t); pr.borders(t); no_split(t)

    for i, (code, mk_n, en_n, tasks, mk_time, en_time) in enumerate(STREAMS, start=4):
        stream_chapter(d, str(i), code, mk_n, en_n, tasks, mk_time, en_time)

    pr.chapter(d, "8", "ОТВОРЕНИ СТАВКИ", "Open Items")
    for m_, e_ in OPEN_ITEMS:
        pr.bullet(d, m_, e_)

    d.save(DOCX)
    print("workstreams %d · tasks %d · completed %d · in progress %d"
          % (len(STREAMS), all_tasks, done, all_tasks - done))
    print("wrote %s (%d KB)" % (os.path.basename(DOCX), os.path.getsize(DOCX) // 1024))
    return 0


if __name__ == "__main__":
    sys.exit(main())
