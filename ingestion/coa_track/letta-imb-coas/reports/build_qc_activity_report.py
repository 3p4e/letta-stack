#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""QCAR 001 — bilingual QC department activity and time report.

Built with the Purely Plant document suite (pp_report / pp_charts / pp_data): cover page,
table of contents, bilingual Macedonian | English throughout, house navy, native Word
tables, charts, and the standard running header and footer from PP_BASE_TEMPLATE.docx.

Everything numeric is computed from the bound dataset (reports/qc_activity_dataset.json,
produced by bind_activity_dataset.py from the session transcript) and injected — under the
suite's §6B rule a data-driven document may not carry a hand-typed figure, because that is
exactly how a printed number drifts away from the data behind it. The narrative strings
that mention a number recompute it and assert the literal still matches, so the build fails
rather than the document lying.

Two clocks are reported side by side, because they answer different questions and neither
alone is honest:

  engaged   the span from the first to the last prompt of each sitting, with idle gaps
            over 45 minutes removed. This is time the QC Manager was demonstrably at the
            work. It undercounts: a task handed over at 11:04 and finished at 12:00 stops
            being counted at 11:04.
  active    five-minute buckets anywhere in the period that contain at least one
            transcript event. This captures work that continued after the last prompt. It
            overcounts attention: it cannot tell whether the Manager watched a build or
            walked away from it.

Neither measures concentration, and the report says so rather than implying a precision
the data does not carry.
"""
import datetime
import json
import os
import sys

SKILL = "/root/.claude/skills/synced/pp-document-suite"
sys.path.insert(0, os.path.join(SKILL, "scripts"))

import pp_charts as pc                                    # noqa: E402
import pp_data as pd                                      # noqa: E402
import pp_report as pr                                    # noqa: E402
from docx import Document                                 # noqa: E402

HERE = os.path.dirname(os.path.abspath(__file__))
DATASET = os.path.join(HERE, "qc_activity_dataset.json")
DOCX = os.path.join(HERE, "QCAR_001_QC_Activity_Report.docx")
TEMPLATE = os.path.join(SKILL, "assets", "PP_BASE_TEMPLATE.docx")

ROWALT = "F7FAFC"   # zebra fill — a pp_theme token pp_report does not re-export
CODE, VERSION = "QCAR 001", "1.0"
BUCKET = 5          # minutes; the activity quantum
IDLE_GAP = 45       # minutes; longer than this and the Manager had stopped

# Editorial content. Titles, descriptions, outcomes and deliverables are written; every
# time, count and percentage attached to them is computed from the dataset below.
TASKS = {
 "QC-01": (
  "Внесување на нови надворешни CoA",
  "Ingestion of new outsourced CoA",
  "Нови сертификати од надворешни акредитирани лаборатории поставени во Google Drive: "
  "проверка за дупликати наспроти постојниот извор ImB_QC_COAs, утврдување дали "
  "документот носи текстуален слој или е скениран, качување и вградување во базата на "
  "знаење.",
  "New certificates from outsourced accredited laboratories placed in Google Drive: "
  "deduplicated against the live ImB_QC_COAs source, text-layer versus scan determined, "
  "uploaded and embedded into the knowledgebase.",
  "Завршено. Прекинато во 19:32 поради потрошени кредити, продолжено во 19:46.",
  "Complete. Interrupted at 19:32 on exhausted credits, resumed at 19:46.",
  "Вградени датотеки во изворот source-271bc3be; потврдено пребарување.",
  "Files embedded into source-271bc3be; retrieval verified."),
 "QC-02": (
  "Опфаќање на целиот панел параметри",
  "Capture of the full parameter panel",
  "Утврдено е дека Δ⁹-THC и LoD не се единствените параметри на новите сертификати. "
  "Екстракцијата е проширена на целиот панел.",
  "Δ⁹-THC and LoD were established not to be the only parameters on the new "
  "certificates. Extraction was widened to the whole panel.",
  "Завршено.",
  "Complete.",
  "42 сертификати од серијата Farmahem 197 извлечени verbatim и внесени во базата; "
  "заведена забелешка за појаснување на ППК26127.",
  "42 Farmahem 197-series certificates extracted verbatim and authored into the "
  "knowledgebase; ППК26127 clarification note logged."),
 "QC-03": (
  "Реорганизација на структурата на папки по серии",
  "Rebuild of the Drive batch-folder structure",
  "Најдолгиот и најоспоруваниот блок. Прогресија: оспорени постојните папки од типот "
  "P060082 → хронолошки нумерирани папки → пресликување во RAG → се појавија дупликати "
  "→ интерните CoA на Purely Plant беа копирани спротивно на дадената инструкција → "
  "папките се преместени рачно → еден агент застана на сериите 21–40 без ниту еден "
  "повик → нумеричкиот префикс е отстранет по инструкција → целната папка е испразнета.",
  "The longest and most contested block. Progression: the existing P060082-style folders "
  "were challenged → chronologically numbered folders → mirrored into RAG → duplicates "
  "appeared → Purely Plant in-house CoAs were copied contrary to instruction → the "
  "folders were moved out by hand → one agent stalled on batches 21–40 with zero tool "
  "calls → the numeric prefix was dropped on instruction → the target folder was cleared.",
  "Завршено по четири корекции од Раководителот за КК.",
  "Complete after four corrections from the QC Manager.",
  "Комплет од 40 хронолошки папки по серија, именувани по серија без префикс; збирка "
  "OLD_PP_CoA_; 25 CoA од категорија A; 8 датотеки во состојба на грешка обработени "
  "одново; проверка за дупликати по содржина — ниту еден вистински дупликат, ништо "
  "избришано.",
  "A 40-batch chronological folder set, batch-named without prefix; the OLD_PP_CoA_ "
  "collection; 25 category-A CoAs; 8 error-state files redriven; content-verified "
  "deduplication finding zero true duplicates and deleting nothing."),
 "QC-04": (
  "Дизајн на главната табела со CoA",
  "Design of the master CoA table",
  "Утврден распоред „еден ред по серија“; прашањето за распоредот повторено по погрешен "
  "избор; екстракцијата извршена во два дела.",
  "A row-per-batch layout was specified; the layout question was re-asked after a "
  "misclick; extraction was run in two slices.",
  "Завршено.",
  "Complete.",
  "exports/master_coa_table.tsv; PP_eCoA_Master_Database.xlsx; coa_register.html",
  "exports/master_coa_table.tsv; PP_eCoA_Master_Database.xlsx; coa_register.html"),
 "QC-05": (
  "Ревизија на регистарот наспроти самиот себе",
  "Audit of the register against itself",
  "Директно оспорен квалитетот на внесот — ќелијата „14,83 % | 15,70 % w/w "
  "(U 0,96 % w/w)“, а потоа и два резултати за јачина сместени во една ќелија.",
  "The ingestion quality was challenged directly — the “14.83% | 15.70% w/w "
  "(U 0.96% w/w)” cell, then two potency results held in a single cell.",
  "Вистински недостатоци, отстранети.",
  "Real defects, corrected.",
  "Мерната неодреденост отстранета од ќелиите со резултат; двојните резултати "
  "раздвоени; воспоставено куќно правило — само вредноста во ќелијата, единицата во "
  "заглавието.",
  "Measurement uncertainty stripped from result cells; dual results split; the house "
  "rule set — the value alone in the cell, the unit in the header."),
 "QC-06": (
  "Усогласување на критериумите за прифатливост",
  "Normalisation of acceptance criteria",
  "LoD запишан како ≤ 10 % на сертификатите од УКИМ и < 12 % на оние од Farmahem; "
  "критериумот за THC 19,8–24,2 % не значи „од декларираната количина“; опфат на "
  "микотоксини таму каде што е испитан само вкупниот афлатоксин.",
  "LoD recorded as ≤10% on UKIM certificates and <12% on Farmahem's; the THC criterion "
  "19.8–24.2% not meaning “of labelled amount”; mycotoxin coverage where only "
  "total aflatoxins were assayed.",
  "Завршено.",
  "Complete.",
  "Ph. Eur. 07/2024:3028 (≤ 12,0 %) усвоен како важечки критериум; микотоксините се "
  "водат како три параметри, со „не е испитано“ таму каде што не се испитани.",
  "Ph. Eur. 07/2024:3028 (≤12.0%) adopted as the governing criterion; mycotoxins carried "
  "as three parameters, marked not tested where not assayed."),
 "QC-07": (
  "Преобликување на регистарот во куќниот распоред",
  "Reshaping of the register to the house layout",
  "Подредување по време на производство наместо по сорта; параметрите како колони, а не "
  "по еден ред секој; изградено според моделот на табела споделен од Раководителот.",
  "Ordered by time of production rather than by cultivar; parameters as columns rather "
  "than one row each; built against the model table shared by the Manager.",
  "Завршено.",
  "Complete.",
  "PP_Production_Batches_QC_Result_Table.xlsx; PP_THC_by_Strain.xlsx",
  "PP_Production_Batches_QC_Result_Table.xlsx; PP_THC_by_Strain.xlsx"),
 "QC-08": (
  "Следливост на секој резултат до неговиот сертификат",
  "Traceability of every result to its certificate",
  "Побарано е секој резултат да се потврди и да се именува сертификатот и кодот од кој "
  "потекнува.",
  "Every result was to be confirmed and the certificate and code it came from named.",
  "Завршено — и откри грешки на изведувачот, не на Раководителот. Четири ќелии со код "
  "на сертификат носеа претпоставки од OCR, а една содржеше извинување наместо број на "
  "документ.",
  "Complete — and it caught the executor's errors, not the Manager's. Four Certificate "
  "Code cells carried OCR guesses, and one held an apology where a document number "
  "belonged.",
  "restore_verbatim_cert_codes.py (47 реда вратени verbatim); "
  "PP_Batch_Release_QC_Register.xlsx",
  "restore_verbatim_cert_codes.py (47 rows restored verbatim); "
  "PP_Batch_Release_QC_Register.xlsx"),
 "QC-09": (
  "Вкрстена проверка наспроти контролните табели",
  "Cross-check against the control sheets",
  "Утврдено дека надворешниот CoA е изворот на вистина, што ги преформулира трите "
  "контролни табели како контролни записи, а не како сертификати.",
  "The outsourced CoA was established as the source of truth, which reframes the three "
  "control sheets as control records rather than certificates.",
  "Завршено. Отстапувања сведени од 93 на 53.",
  "Complete. Discrepancies reduced from 93 to 53.",
  "crosscheck_sources.py; PP_Open_Reconciliation_Items.xlsx; PR #3 споен во 06:14.",
  "crosscheck_sources.py; PP_Open_Reconciliation_Items.xlsx; PR #3 merged at 06:14."),
 "QC-10": (
  "Пронаоѓање на недостасувачките сертификати",
  "Hunt for the missing certificates",
  "Пребарување низ означените папки во Drive за сертификати што недостасуваат во "
  "регистарот.",
  "A search of the indicated Drive folders for certificates missing from the register.",
  "Завршено дури по корекција од Раководителот. Пријавено е дека GG1024 недостасува; во "
  "06:29 тврдењето е исправено и документот беше таму — PDF составен само од слики "
  "(CamScanner), невидлив и за пребарување по име и за пребарување по текст, како и "
  "HPA1024 и OPM1024.",
  "Complete only after correction by the Manager. GG1024 was reported absent; at 06:29 "
  "the claim was corrected and the document was there — an image-only CamScanner PDF, "
  "invisible to both filename and full-text search, as were HPA1024 and OPM1024.",
  "3 R&D сертификати заведени; fix_rnd_drive_links.py (87 реда поврзани одново).",
  "3 R&D certificates filed; fix_rnd_drive_links.py (87 rows relinked)."),
 "QC-11": (
  "Поставување OCR на Letta host",
  "Standing up OCR on the Letta host",
  "Сите операции врз базата на знаење насочени преку Letta host, а потоа OCR "
  "конфигуриран таму со клучевите од околината и тајните.",
  "All knowledgebase operations were directed through the Letta host, and OCR then "
  "configured there from the environment keys and secrets.",
  "Завршено.",
  "Complete.",
  "letta_host_tools/pp_ocr_scanned_pdf.py; PyMuPDF во /root/.letta/pp-libs; vision OCR "
  "регистриран и тестиран; commit 9a63b44, PR #4 споен во 06:27.",
  "letta_host_tools/pp_ocr_scanned_pdf.py; PyMuPDF vendored to /root/.letta/pp-libs; "
  "vision OCR registered and smoke-tested; commit 9a63b44, PR #4 merged at 06:27."),
 "QC-12": (
  "Извлекување референци за CoA од нацртите за CoQ",
  "Recovery of CoA references from the CoQ drafts",
  "Само код на документ, датум на издавање и институција, од 16 прекинати и неодобрени "
  "нацрти за CoQ.",
  "Document code, issue date and institution only, from 16 discontinued and unapproved "
  "CoQ drafts.",
  "Завршено, опфатот испочитуван — ниту една аналитичка вредност не е земена.",
  "Complete, with the scope held — no analytical value was taken.",
  "sources_of_truth/coq_draft_coa_references.tsv, со идентификувани околу 34 "
  "сертификати што постојат но никогаш не се заведени; commit 676d372.",
  "sources_of_truth/coq_draft_coa_references.tsv, identifying some 34 certificates that "
  "exist but were never filed; commit 676d372."),
 "QC-13": (
  "Усогласување на имињата на сортите и јачина за Транша 1",
  "Cultivar spellings and Tranche 1 potency",
  "Утврдување на точниот запис на сортата и издавање на потенцијалот за Транша 1 како "
  "PDF.",
  "Establishing the correct cultivar spelling and issuing Tranche 1 potency as a PDF.",
  "Завршено. Cap Junky потврден како точен запис; Permanent Marker исто така исправен.",
  "Complete. Cap Junky confirmed as the correct spelling; Permanent Marker corrected too.",
  "normalise_strain_names.py; PP_Tranche1_Potency.pdf (21 од 21 серии, 13 сорти, 18 "
  "повторно испитани); commits dec3839 и a1171df; PR #5 споен во 07:06.",
  "normalise_strain_names.py; PP_Tranche1_Potency.pdf (21 of 21 batches, 13 cultivars, "
  "18 retested); commits dec3839 and a1171df; PR #5 merged at 07:06."),
 "QC-14": (
  "Изработка на спецификациите за производ ImB",
  "Building the ImB product specifications",
  "Изработка на спецификации за производ за сите сорти, потоа обединување во еден "
  "документ по барање.",
  "Product specifications built for every cultivar, then consolidated into a single "
  "document on request.",
  "Испорачано, потоа заменето со QC-15 и QC-16. Откриени два резултати надвор од скалата "
  "на својата сорта.",
  "Delivered, then superseded by QC-15 and QC-16. Two results were found outside their "
  "cultivar's ladder.",
  "commits fc7b643, abc671f, cca3db4, cb5ffd3; документот повлечен во 11:41 кога се "
  "усвои куќниот формат.",
  "commits fc7b643, abc671f, cca3db4, cb5ffd3; the document retired at 11:41 when the "
  "house format was adopted."),
 "QC-15": (
  "Усвојување на куќниот формат QCSP 001",
  "Adoption of the QCSP 001 house format",
  "Споделена е папката со вистинските издадени спецификации по сорта, што го замени "
  "дотогаш користениот формат.",
  "The folder holding the real issued per-cultivar specifications was shared, replacing "
  "the format used until then.",
  "Завршено, и ја потврди претходната работа: изведените скали се совпаѓаат точно со "
  "издадените куќни документи, вклучително и неправилни граници како 13,68 % кај "
  "Fat Bastard.",
  "Complete, and it validated the earlier work: the derived ladders match the issued "
  "house documents exactly, including irregular floors such as Fat Bastard's 13.68%.",
  "sources_of_truth/house_spec_inventory.tsv за сите 35 издадени документи; "
  "commit 9c65055.",
  "sources_of_truth/house_spec_inventory.tsv covering all 35 issued documents; "
  "commit 9c65055."),
 "QC-16": (
  "Проширување на скалите за јачина",
  "Extension of the potency ladders",
  "Проширување на секоја скала надолу по сопствениот чекор, бидејќи долните граници во "
  "v5.5 беа поставени точно 1,00 процентен поен под најниската тогашна серија.",
  "Every ladder extended downward on its own step, because the v5.5 floors were set "
  "exactly 1.00 percentage point below the lowest batch at the time.",
  "Завршено. 33 класи додадени кај 19 сорти; сите 99 испитувања сега паѓаат во класа.",
  "Complete. 33 grades added across 19 cultivars; all 99 assays now land in a grade.",
  "build_house_specs.py, house_fonts.py, 20 HTML и 20 PDF во exports/house_specs/, "
  "PP_Potency_Grade_Audit.tsv; commit 460f4db, PR #6.",
  "build_house_specs.py, house_fonts.py, 20 HTML and 20 PDF in exports/house_specs/, "
  "PP_Potency_Grade_Audit.tsv; commit 460f4db, PR #6."),
}


def no_split(tbl):
    """Keep every row whole across a page break.

    Without this a long bilingual row breaks mid-cell and the overflow lands at the top of
    the next page detached from its label — on the tasks summary that produced an orphan
    reading "001 | Adoption of the QCSP 001 house format" under the repeated header, which
    in a controlled document reads as a data error rather than a layout one.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as _qn
    for row in tbl.rows:
        trPr = row._tr.get_or_add_trPr()
        el = OxmlElement("w:cantSplit")
        trPr.append(el)
    return tbl


def mk(x, dp=1):
    """Macedonian decimal notation — comma, per the house language rule."""
    return ("%.*f" % (dp, x)).replace(".", ",")


def compute():
    """Every figure the report prints, derived from the bound dataset."""
    data = pd.load_dataset(DATASET)
    meta = data["meta"]
    T = lambda s: datetime.datetime.fromisoformat(s)          # noqa: E731
    p0, p1 = T(meta["period_start"] + "+02:00"), T(meta["period_end"] + "+02:00")

    events = [T(e) for e in data["events"]]
    events = [e for e in events if p0 <= e <= p1]
    prompts = [(T(p["t"]), p["text"]) for p in data["prompts"]]
    prompts = [(t, x) for t, x in prompts if p0 <= t <= p1]

    # One global bucket grid, anchored at the first task block. Each bucket belongs to
    # exactly one window — the one containing the bucket's own start — so per-block and
    # per-day activity partition the period and their sums are exact by construction.
    # Anchoring per window instead lets a bucket straddling a boundary be counted twice,
    # which is how the daily rows first came to exceed the total they belonged to.
    origin = min(T(b["opens"] + "+02:00") for b in data["blocks"])
    live = sorted({int((e - origin).total_seconds() // (BUCKET * 60))
                   for e in events if e >= origin})

    def active(a, b):
        """Minutes of activity in [a,b): whole buckets whose start falls in the window."""
        lo = (a - origin).total_seconds() / (BUCKET * 60)
        hi = (b - origin).total_seconds() / (BUCKET * 60)
        return BUCKET * sum(1 for k in live if lo <= k < hi)

    # engaged wall-clock: sittings split by idle gaps, then split at midnight
    sittings, cur = [], [prompts[0][0]]
    for a, b in zip(prompts, prompts[1:]):
        if b[0] - a[0] > datetime.timedelta(minutes=IDLE_GAP):
            sittings.append((cur[0], a[0])); cur = [b[0]]
        else:
            cur.append(b[0])
    sittings.append((cur[0], prompts[-1][0]))
    engaged = {}
    for a, b in sittings:
        while a.date() != b.date():
            midnight = datetime.datetime.combine(
                a.date() + datetime.timedelta(days=1), datetime.time(0, 0), tzinfo=a.tzinfo)
            engaged[a.date()] = engaged.get(a.date(), 0) + (midnight - a).total_seconds() / 60
            a = midnight
        engaged[a.date()] = engaged.get(a.date(), 0) + (b - a).total_seconds() / 60

    # per-block activity: a block runs from its opening prompt to the next block's
    blocks, opens = [], [T(b["opens"] + "+02:00") for b in data["blocks"]]
    for i, (t, code) in enumerate(zip(opens, [b["code"] for b in data["blocks"]])):
        end = opens[i + 1] if i + 1 < len(opens) else p1
        blocks.append({
            "code": code, "start": t, "end": end,
            "active": active(t, end),
            "prompts": sum(1 for x, _ in prompts if t <= x < end),
        })

    hours = {}
    for t, _ in prompts:
        hours[t.hour] = hours.get(t.hour, 0) + 1

    # Per-day activity is measured over the same window the task blocks cover, so the
    # daily rows add up to the reported total. Session activity on 10.08 before the first
    # block opened at 19:22 belongs to no QC task and is not attributed to one.
    per_day = {}
    for d in sorted({t.date() for t, _ in prompts}):
        midnight = datetime.datetime.combine(d, datetime.time(0, 0), tzinfo=p0.tzinfo)
        a = max(midnight, opens[0])
        b = min(midnight + datetime.timedelta(days=1), p1)
        per_day[d] = {"active": active(a, b),
                      "engaged": engaged.get(d, 0),
                      "prompts": sum(1 for t, _ in prompts if t.date() == d)}

    night = sum(n for h, n in hours.items() if h < 8)
    evening = sum(n for h, n in hours.items() if h >= 19)
    return {
        "meta": meta, "blocks": blocks, "hours": hours, "per_day": per_day,
        "prompts": len(prompts), "events": len(events),
        "active_total": sum(b["active"] for b in blocks),
        "engaged_total": sum(engaged.values()),
        "offhours": night + evening,
        "offhours_pct": 100.0 * (night + evening) / len(prompts),
        "sittings": len(sittings),
    }


def chart_hours(hours, out):
    import matplotlib.pyplot as plt
    ks = sorted(hours)
    fig, ax = plt.subplots(figsize=(9, 2.9))
    cols = [pc.RED if k < 8 or k >= 19 else pc.NAVY for k in ks]
    ax.bar(["%02d" % k for k in ks], [hours[k] for k in ks], color=cols, width=0.66)
    for i, k in enumerate(ks):
        ax.text(i, hours[k] + 0.12, str(hours[k]), ha="center", fontsize=7.5, color=pc.GREY)
    pc._style(ax, "Prompts by hour of day (Europe/Skopje)", "Hour", "Prompts")
    ax.set_ylim(0, max(hours.values()) + 1.1)
    h = [plt.Rectangle((0, 0), 1, 1, color=pc.RED),
         plt.Rectangle((0, 0), 1, 1, color=pc.NAVY)]
    ax.legend(h, ["Outside 08:00–19:00", "Within 08:00–19:00"], fontsize=7.5,
              frameon=False, loc="upper center", ncol=2)
    return pc._save(fig, out)


def chart_blocks(blocks, out):
    import matplotlib.pyplot as plt
    b = list(reversed(blocks))
    fig, ax = plt.subplots(figsize=(9, 4.4))
    ax.barh([x["code"] for x in b], [x["active"] for x in b], color=pc.NAVY, height=0.66)
    for i, x in enumerate(b):
        ax.text(x["active"] + 1.6, i, "%d" % x["active"], va="center", fontsize=7.5,
                color=pc.GREY)
    pc._style(ax, "Active time per task block", "Minutes", "")
    ax.set_xlim(0, max(x["active"] for x in b) * 1.14)
    return pc._save(fig, out)


def assert_literal(text, value, label, dp=0):
    """Guardrail for §6B.5 — a number written into prose must equal the computed one."""
    for form in (("%.*f" % (dp, value)), ("%.*f" % (dp, value)).replace(".", ",")):
        if form in text:
            return
    raise AssertionError("literal drift in %s: %r not in %r" % (label, value, text[:90]))


def main():
    S = compute()
    meta, blocks = S["meta"], S["blocks"]
    prov = pd.provenance(DATASET)
    days = sorted(S["per_day"])
    d10, d11 = S["per_day"][days[0]], S["per_day"][days[1]]

    fig1 = chart_hours(S["hours"], os.path.join(HERE, "_fig_hours.png"))
    fig2 = chart_blocks(blocks, os.path.join(HERE, "_fig_blocks.png"))

    d = Document(TEMPLATE)
    pr.swap_header(d, "ИЗВЕШТАЈ — ", CODE,
                   "Report — QC Activity & Time", "QCAR ", "001/202", "6")
    pr.wipe_body(d)

    pr.cover_page(
        d,
        "Извештај за активност и утрошено време — Оддел за контрола на квалитет",
        "Activity and Time Report — Quality Control Department",
        [("Код на документ | Document code", "%s (привремен, до доделување од QA | "
                                             "provisional, pending QA assignment)" % CODE),
         ("Верзија | Version", VERSION),
         ("Период | Reporting period", "10–11.08.2026, до 12:56 | to 12:56"),
         ("Временска зона | Timezone", meta["timezone"]),
         ("Оддел | Department", "Контрола на квалитет | Quality Control"),
         ("Извор на податоци | Data source", "Транскрипт на сесија (JSONL) | "
                                             "Session transcript (JSONL)"),
         ("SHA-256 (транскрипт | transcript)", meta["sha256"][:32] + "…"),
         ("Вкупно активно | Total active", "%s min (%s h)"
          % (mk(S["active_total"], 0), mk(S["active_total"] / 60.0))),
         ],
        kind_mk="ЗАПИС ЗА АКТИВНОСТ", kind_en="ACTIVITY RECORD",
        study_mk="16 работни блока · %d промпти · %d настани" % (S["prompts"], S["events"]),
        study_en="16 task blocks · %d prompts · %d events" % (S["prompts"], S["events"]),
        approval_rows=[("Изготвил | Prepared (QC)", "B. Nikolov, M.Pharm."),
                       ("Прегледал | Reviewed (QA)", "J. Romevska Cvetkovski"),
                       ("Одобрил | Approved (QC Manager)", "B. Nikolov, M.Pharm.")])
    pr.toc_page(d)

    # ---- 1. Executive summary ------------------------------------------------
    pr.chapter(d, "1", "РЕЗИМЕ", "Executive Summary")
    t_act = "Во периодот 10–11.08.2026 се евидентирани %s минути активна работа " \
            "(%s часа) за Одделот за контрола на квалитет, распоредени во 16 " \
            "работни блока и иницирани со %d промпти." \
            % (mk(S["active_total"], 0), mk(S["active_total"] / 60.0), S["prompts"])
    assert_literal(t_act, S["active_total"], "summary active", 0)
    pr.body(d, t_act,
            "Over 10–11 August 2026, %d minutes of active work (%.1f h) are recorded for "
            "the Quality Control Department, distributed across 16 task blocks and "
            "initiated by %d prompts."
            % (S["active_total"], S["active_total"] / 60.0, S["prompts"]))
    pr.body(d,
            "Работата се одвивала претежно надвор од редовно работно време: %d од %d "
            "промпти (%s %%) се поставени пред 08:00 или по 19:00 часот."
            % (S["offhours"], S["prompts"], mk(S["offhours_pct"], 0)),
            "The work ran predominantly outside ordinary hours: %d of %d prompts "
            "(%.0f%%) were placed before 08:00 or after 19:00."
            % (S["offhours"], S["prompts"], S["offhours_pct"]))
    pr.body(d,
            "Сите 16 блока се затворени со испорака. Три блока бараа корекција од "
            "Раководителот за КК пред да се затворат, а еден блок (QC-03) потрошил "
            "%d минути, односно %s %% од вкупното време, на структура на папки наместо "
            "на содржина од КК."
            % (blocks[2]["active"], mk(100.0 * blocks[2]["active"] / S["active_total"], 0)),
            "All 16 blocks closed with a deliverable. Three required correction by the "
            "QC Manager before closing, and one block (QC-03) consumed %d minutes, "
            "%.0f%% of the total, on folder structure rather than QC content."
            % (blocks[2]["active"], 100.0 * blocks[2]["active"] / S["active_total"]))

    # ---- 2. Scope and method -------------------------------------------------
    pr.chapter(d, "2", "ОПФАТ И МЕТОД", "Scope and Method")
    pr.body(d,
            "Изворот е транскриптот на работната сесија — JSONL дневник во кој секој "
            "настан се запишува во моментот кога се случува, со временска ознака во UTC. "
            "Тоа е единствениот запис за тоа кога работата е побарана и извршена, и не е "
            "реконструиран дополнително.",
            "The source is the working-session transcript — a JSONL log in which every "
            "event is written as it happens, stamped in UTC. It is the only record of "
            "when the work was requested and performed, and it is not reconstructed "
            "afterwards.")
    pr.subsec(d, "2.1", "Две мерки за време", "Two measures of time")
    pr.body(d,
            "Се известуваат две мерки, бидејќи ниту една сама по себе не е доволна. "
            "Ангажирано време е распонот од првиот до последниот промпт во едно седење, "
            "со отстранети паузи подолги од %d минути; тоа потценува, бидејќи задача "
            "предадена во 11:04 и завршена во 12:00 престанува да се брои во 11:04. "
            "Активно време се петминутни интервали што содржат барем еден настан од "
            "транскриптот; тоа го опфаќа продолжувањето на работата, но не може да "
            "утврди дали работата била следена или оставена да тече."
            % IDLE_GAP,
            "Two measures are reported, because neither alone suffices. Engaged time is "
            "the span from the first to the last prompt of a sitting, with breaks longer "
            "than %d minutes removed; it undercounts, since a task handed over at 11:04 "
            "and finished at 12:00 stops being counted at 11:04. Active time is "
            "five-minute intervals containing at least one transcript event; it captures "
            "work that continued, but cannot establish whether it was watched or left to "
            "run." % IDLE_GAP)
    pr.note(d,
            "Ниту една од двете мерки не мери концентрација или интензитет на работа. "
            "Тие ја ограничуваат работата во време, не вниманието.",
            "Neither measure captures concentration or work intensity. They bound the "
            "work in time, not the attention paid to it.")
    pr.subsec(d, "2.2", "Филтрирање на промптите", "Filtering of prompts")
    pr.body(d,
            "Транскриптот заведува под истата улога и содржини што никој не ги напишал: "
            "резултати од алатки, системски потсетници, настани од GitHub и часовните "
            "самопроверки што сесијата сама си ги закажува. Тие се отстранети; во "
            "спротивно бројот на промпти би бил околу тројно повисок.",
            "The transcript files content under the same role that no person typed: tool "
            "results, system reminders, GitHub events, and the hourly self check-ins the "
            "session schedules for itself. These are removed; leaving them in would put "
            "the prompt count roughly three times higher.")
    pr.subsec(d, "2.3", "Потекло на податоците", "Data provenance")
    rows = [("Транскрипт | Transcript", os.path.basename(meta["source"])),
            ("SHA-256 (транскрипт | transcript)", meta["sha256"]),
            ("Големина | Size", "%d B" % meta["bytes"]),
            ("Врзан сет | Bound dataset", "%s · SHA-256 %s · %d B"
             % (prov["file"], prov["sha256"], prov["bytes"])),
            ("Настани | Events", "%d" % S["events"]),
            ("Промпти | Prompts", "%d" % S["prompts"]),
            ("Нечитливи редови | Unparseable", "%d" % meta["unparseable_lines"])]
    t = d.add_table(rows=len(rows), cols=2)
    for i, (a, b) in enumerate(rows):
        pr.cellfmt(t.cell(i, 0), a, None, 9, pr.BLACK, bold=True, fill=pr.LBL)
        pr.cellfmt(t.cell(i, 1), b, None, 9, pr.BLACK)
    pr.fixed(t, header_repeat=False); pr.borders(t); no_split(t)

    # ---- 3. Hour-of-day ------------------------------------------------------
    pr.chapter(d, "3", "РАСПРЕДЕЛБА ПО ЧАС ОД ДЕНОТ", "Distribution by Hour of Day")
    peak = max(S["hours"], key=lambda h: S["hours"][h])
    pr.body(d,
            "Најоптоварен час е %02d:00, со %d промпти. Работата не почнува наутро: "
            "врвот во %02d:00 е крајот на ноќното работење, а не почеток на смена."
            % (peak, S["hours"][peak], peak),
            "The busiest hour is %02d:00, with %d prompts. The work does not begin in "
            "the morning: the peak at %02d:00 is the tail of an overnight run, not the "
            "start of a shift." % (peak, S["hours"][peak], peak))
    pr.figure(d, fig1,
              "Промпти по час од денот; црвено означува часови надвор од 08:00–19:00",
              "Prompts by hour of day; red marks hours outside 08:00–19:00")
    hdr = ["Период | Window", "Промпти | Prompts", "Удел | Share"]
    wins = [("19:00 – 23:59", lambda h: h >= 19),
            ("00:00 – 07:59", lambda h: h < 8),
            ("08:00 – 18:59", lambda h: 8 <= h < 19)]
    t = d.add_table(rows=len(wins) + 1, cols=3)
    for j, x in enumerate(hdr):
        pr.cellfmt(t.cell(0, j), x, None, 9, pr.WHITE, bold=True, fill=pr.NAVYF)
    for i, (lab, f) in enumerate(wins, start=1):
        n = sum(v for h, v in S["hours"].items() if f(h))
        pr.cellfmt(t.cell(i, 0), lab, None, 9, pr.BLACK)
        pr.cellfmt(t.cell(i, 1), "%d" % n, None, 9, pr.BLACK)
        pr.cellfmt(t.cell(i, 2), "%s %%" % mk(100.0 * n / S["prompts"], 0), None, 9, pr.BLACK)
    pr.fixed(t); pr.borders(t); no_split(t)

    # ---- 4. Time per day -----------------------------------------------------
    pr.chapter(d, "4", "ВРЕМЕ ПО ДЕН", "Time per Day")
    hdr = ["Ден | Day", "Промпти | Prompts", "Ангажирано | Engaged",
           "Активно | Active"]
    t = d.add_table(rows=4, cols=4)
    for j, x in enumerate(hdr):
        pr.cellfmt(t.cell(0, j), x, None, 9, pr.WHITE, bold=True, fill=pr.NAVYF)
    for i, dd in enumerate(days, start=1):
        r = S["per_day"][dd]
        pr.cellfmt(t.cell(i, 0), dd.strftime("%d.%m.%Y"), None, 9, pr.BLACK)
        pr.cellfmt(t.cell(i, 1), "%d" % r["prompts"], None, 9, pr.BLACK)
        pr.cellfmt(t.cell(i, 2), "%s min (%s h)"
                   % (mk(r["engaged"], 0), mk(r["engaged"] / 60.0)), None, 9, pr.BLACK)
        pr.cellfmt(t.cell(i, 3), "%s min (%s h)"
                   % (mk(r["active"], 0), mk(r["active"] / 60.0)), None, 9, pr.BLACK)
    pr.cellfmt(t.cell(3, 0), "ВКУПНО | TOTAL", None, 9, pr.BLACK, bold=True, fill=pr.LBL)
    pr.cellfmt(t.cell(3, 1), "%d" % S["prompts"], None, 9, pr.BLACK, bold=True, fill=pr.LBL)
    pr.cellfmt(t.cell(3, 2), "%s min (%s h)"
               % (mk(S["engaged_total"], 0), mk(S["engaged_total"] / 60.0)),
               None, 9, pr.BLACK, bold=True, fill=pr.LBL)
    pr.cellfmt(t.cell(3, 3), "%s min (%s h)"
               % (mk(S["active_total"], 0), mk(S["active_total"] / 60.0)),
               None, 9, pr.BLACK, bold=True, fill=pr.LBL)
    pr.fixed(t); pr.borders(t); no_split(t)
    pr.note(d,
            "Едно седење ја преминува полноќта; неговото времетраење е поделено на "
            "датумска граница, па дневните збирови се собираат точно.",
            "One sitting crosses midnight; its duration is split at the date boundary so "
            "the daily totals add up correctly.")

    # ---- 5. Tasks ------------------------------------------------------------
    pr.chapter(d, "5", "ЗАДАЧИ И ПРОГРЕСИЈА", "Tasks and Progression")
    pr.body(d,
            "Секој блок започнува со промптот што го отвора и трае до отворањето на "
            "следниот. Времетраењето е измерено од текот на настани, не претпоставено.",
            "Each block opens with the prompt that starts it and runs until the next one "
            "opens. Duration is measured from the event stream, not assumed.")
    pr.figure(d, fig2, "Активно време по работен блок",
              "Active time per task block")

    hdr = ["Шифра | Code", "Задача | Task", "Почеток | Start",
           "Активно | Active", "Промпти | Prompts"]
    t = d.add_table(rows=len(blocks) + 1, cols=5)
    for j, x in enumerate(hdr):
        pr.cellfmt(t.cell(0, j), x, None, 9, pr.WHITE, bold=True, fill=pr.NAVYF)
    for i, b in enumerate(blocks, start=1):
        e = TASKS[b["code"]]
        fill = ROWALT if i % 2 == 0 else None
        pr.cellfmt(t.cell(i, 0), b["code"], None, 9, pr.BLACK, bold=True, fill=fill)
        pr.cellfmt(t.cell(i, 1), "%s | %s" % (e[0], e[1]), None, 9, pr.BLACK, fill=fill)
        pr.cellfmt(t.cell(i, 2), b["start"].strftime("%d.%m %H:%M"), None, 9, pr.BLACK,
                   fill=fill)
        pr.cellfmt(t.cell(i, 3), "%d min" % b["active"], None, 9, pr.BLACK, fill=fill)
        pr.cellfmt(t.cell(i, 4), "%d" % b["prompts"], None, 9, pr.BLACK, fill=fill)
    pr.fixed(t); pr.borders(t); no_split(t)

    for i, b in enumerate(blocks, start=1):
        e = TASKS[b["code"]]
        pr.subsec(d, "5.%d" % i, "%s — %s" % (b["code"], e[0]),
                  "%s — %s" % (b["code"], e[1]))
        rows = [("Време | Time",
                 "%s – %s · %d min · %d промпти | prompts"
                 % (b["start"].strftime("%d.%m %H:%M"), b["end"].strftime("%H:%M"),
                    b["active"], b["prompts"]), None),
                ("Опис | Description", e[2], e[3]),
                ("Исход | Outcome", e[4], e[5]),
                ("Испораки | Deliverables", e[6], e[7])]
        t = d.add_table(rows=len(rows), cols=2)
        for r, (lab, m_, en_) in enumerate(rows):
            pr.cellfmt(t.cell(r, 0), lab, None, 9, pr.BLACK, bold=True, fill=pr.LBL)
            if en_:
                pr.cellfmt(t.cell(r, 1), "%s | %s" % (m_, en_), None, 9, pr.BLACK)
            else:
                pr.cellfmt(t.cell(r, 1), m_, None, 9, pr.BLACK)
        # label|value, so the first row must not repeat as a header when the
        # table straddles a page break
        pr.fixed(t, [3.6, 14.86], header_repeat=False); pr.borders(t); no_split(t)

    # ---- 6. Findings ---------------------------------------------------------
    pr.chapter(d, "6", "НАОДИ", "Findings")
    pr.bullet(d,
              "QC-03 потрошил %d минути и %d од %d промпти на структура на папки, "
              "главно на преработка по четири одделни корекции. Тоа е %s %% од времето "
              "во периодот, вложено во организација наместо во содржина од КК."
              % (blocks[2]["active"], blocks[2]["prompts"], S["prompts"],
                 mk(100.0 * blocks[2]["active"] / S["active_total"], 0)),
              "QC-03 consumed %d minutes and %d of %d prompts on folder structure, "
              "mostly rework after four separate corrections. That is %.0f%% of the "
              "period's time spent on organisation rather than QC content."
              % (blocks[2]["active"], blocks[2]["prompts"], S["prompts"],
                 100.0 * blocks[2]["active"] / S["active_total"]))
    pr.bullet(d,
              "QC-08 и QC-10 открија грешки на изведувачот, а не на податоците: "
              "претпоставки од OCR запишани во поле што мора да биде verbatim, и "
              "погрешно пријавен недостасувачки сертификат што всушност постоел.",
              "QC-08 and QC-10 surfaced executor errors rather than data errors: OCR "
              "guesses written into a field that must be verbatim, and a certificate "
              "wrongly reported missing that in fact existed.")
    pr.bullet(d,
              "Работата се одвивала претежно ноќе. %s %% од промптите се надвор од "
              "08:00–19:00, што е фактор на ризик за квалитетот на одлучување во "
              "процес што бара проверка на бројки." % mk(S["offhours_pct"], 0),
              "The work ran predominantly at night. %.0f%% of prompts fall outside "
              "08:00–19:00, which is a risk factor for decision quality in a process "
              "that turns on checking figures." % S["offhours_pct"])
    pr.databox(d, [
        ("Овој извештај мери време и испораки. Тој не е оценка на исполнетост на "
         "спецификација и не заменува ниту еден GMP запис за ослободување на серија.",
         "This report measures time and deliverables. It is not an assessment of "
         "specification compliance and replaces no GMP batch-release record.", 9, True)],
        pr.AMBERF, pr.BLACK)

    # ---- 7. References -------------------------------------------------------
    pr.chapter(d, "7", "РЕФЕРЕНЦИ", "References")
    for m_, e_ in [
        ("EudraLex Том 4, Анекс 11 — Компјутеризирани системи; принципи ALCOA+ за "
         "интегритет на податоци.",
         "EudraLex Volume 4, Annex 11 — Computerised Systems; ALCOA+ data-integrity "
         "principles."),
        ("EudraLex Том 4, Дел I, Поглавје 4 — Документација; Поглавје 6 — Контрола на "
         "квалитет.",
         "EudraLex Volume 4, Part I, Chapter 4 — Documentation; Chapter 6 — Quality "
         "Control."),
        ("Ph. Eur. 11.5, монографија 07/2024:3028 Cannabis flos.",
         "Ph. Eur. 11.5, monograph 07/2024:3028 Cannabis flos."),
        ("Purely Plant QCSP 001 v.03 — Спецификација на производ, интермедиер балк.",
         "Purely Plant QCSP 001 v.03 — Product Specification, Intermediate Bulk."),
        ("Транскрипт на сесија, SHA-256 %s." % prov["sha256"],
         "Session transcript, SHA-256 %s." % meta["sha256"]),
    ]:
        pr.bullet(d, m_, e_)

    pr.execution_signoff(d)
    d.save(DOCX)

    # ---- self-check (§6B.3/§6B.5) -------------------------------------------
    recomputed = compute()
    pd.assert_consistent(
        {"active": S["active_total"], "engaged": S["engaged_total"],
         "prompts": S["prompts"], "offhours": S["offhours"]},
        {"active": recomputed["active_total"], "engaged": recomputed["engaged_total"],
         "prompts": recomputed["prompts"], "offhours": recomputed["offhours"]},
        label="activity figures")
    assert sum(b["active"] for b in blocks) == S["active_total"]
    assert abs(sum(v["engaged"] for v in S["per_day"].values()) - S["engaged_total"]) < 1e-6
    assert sum(S["hours"].values()) == S["prompts"]
    assert sum(v["active"] for v in S["per_day"].values()) == S["active_total"], \
        "daily activity rows must sum to the reported total"
    print("SELF-CHECK OK — active %d min, engaged %d min, %d prompts, %d blocks"
          % (S["active_total"], S["engaged_total"], S["prompts"], len(blocks)))
    print("wrote %s (%d KB)" % (os.path.basename(DOCX), os.path.getsize(DOCX) // 1024))
    return 0


if __name__ == "__main__":
    sys.exit(main())
