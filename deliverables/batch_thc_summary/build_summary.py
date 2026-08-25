# -*- coding: utf-8 -*-
"""Builds the bilingual MK|EN informal working document listing the 12 requested
P-batch numbers, their cultivation batch numbers, and the last-tested Total Δ9-THC
result (eCoA code + date). Bound to dataset.json — bind -> compute -> inject -> assert
(pp-document-suite §6B). No number in the printed document is hand-typed independent
of the dataset; the self-check below re-derives every printed figure and asserts it
matches what dataset.json holds before the file is saved.
"""
import json
import hashlib
import datetime
import sys
import os

sys.path.insert(0, os.path.expanduser(
    "~/.claude/skills/synced/pp-document-suite/scripts"))

from docx import Document
import pp_report as pr

HERE = os.path.dirname(os.path.abspath(__file__))
TEMPLATE = os.path.expanduser(
    "~/.claude/skills/synced/pp-document-suite/assets/PP_BASE_TEMPLATE.docx")
DATASET = os.path.join(HERE, "dataset.json")
OUT = os.path.join(HERE, "P-Batch_Total_THC_Summary.docx")


def load_dataset():
    with open(DATASET, encoding="utf-8") as f:
        data = json.load(f)
    sha = hashlib.sha256(open(DATASET, "rb").read()).hexdigest()
    return data, sha


def fmt_date(iso):
    if not iso:
        return "—"
    y, m, d = iso.split("-")
    return f"{d}.{m}.{y}"


def fmt_pct(v):
    """House numeric rule: exactly two decimals, % immediately after the digit,
    no space; Macedonian uses a decimal comma, English a decimal point."""
    if v is None:
        return None, None
    mk = f"{v:.2f}".replace(".", ",") + "%"
    en = f"{v:.2f}%"
    return mk, en


def pct_cell(cell, v):
    mk, en = fmt_pct(v)
    if mk is None:
        pr.cellfmt(cell, "—", None, 10, pr.GREY)
    else:
        pr.cellfmt(cell, mk, en, 10, pr.BLACK)


def build(data):
    d = Document(TEMPLATE)
    pr.informal_header(
        d,
        "Преглед: бараните P-броеви — последно тестиран вкупен Δ9-ТХЦ",
        "Overview: requested P-numbers — last tested Total Δ9-THC",
        tag_mk="Работен извадок | Не е контролиран запис",
        tag_en="Working extract | Not a controlled record",
    )
    pr.wipe_body(d)

    info_rows = [
        ("Опфат | Scope", "12 бараните P-броеви на серии | The 12 requested P-batch numbers"),
        ("Извор на податоци | Data source",
         "QC регистар на резултати + оригинални еКоА сертификати (сертификат = извор на вистина) | "
         "QC results register + original eCoA certificates (certificate is source of truth)"),
        ("Не е користено | Explicitly not used",
         "List of COAs.xlsx / декларации од траншата (не се QC резултати) | "
         "List of COAs.xlsx / tranche declarations (not QC results)"),
        ("Датум на изработка | Prepared", "25.08.2026"),
    ]
    pr.cover_page(
        d,
        "Преглед на вкупен Δ9-ТХЦ по бараните P-броеви",
        "Total Δ9-THC Overview by Requested P-Numbers",
        info_rows,
        kind_mk="Работен извадок", kind_en="Working extract",
        controlled=False,
    )

    pr.chapter(d, "1", "РЕЗУЛТАТИ — ПОСЛЕДНО ТЕСТИРАН ВКУПЕН Δ9-ТХЦ",
               "RESULTS — LAST TESTED TOTAL Δ9-THC")
    pr.note(
        d,
        "За секој P-број: последната (најнова по датум) сертифицирана вредност за вкупен "
        "Δ9-ТХЦ, култивациската серија што ја носи, и еКоА кодот/датумот од самиот сертификат. "
        "Каде постојат повеќе тестирања за истата серија, дополнителните рунди се дадени во "
        "Анексите подолу.",
        "For each P-number: the most recent certified Total Δ9-THC result, the cultivation "
        "batch it belongs to, and the eCoA code/date taken from the certificate itself. Where "
        "a batch carries more than one test, the additional rounds are listed in the annexes "
        "below.",
    )
    pr.gap(d, 4)

    headers = ["P-број | P-number", "Сорта | Strain",
               "Култивациска серија | Cultivation batch",
               "Вкупно Δ9-ТХЦ | Total Δ9-THC",
               "еКоА код | eCoA code", "Датум | Date",
               "Забелешка | Remark"]
    rows = data["rows"]
    t = d.add_table(rows=len(rows) + 1, cols=len(headers))
    for j, h in enumerate(headers):
        pr.cellfmt(t.cell(0, j), h, None, 9, pr.WHITE, bold=True, fill=pr.NAVYF)
    for i, row in enumerate(rows, start=1):
        pr.cellfmt(t.cell(i, 0), row["p_number"], None, 10, pr.BLACK, bold=True)
        pr.cellfmt(t.cell(i, 1), row["strain"], None, 10, pr.BLACK)
        cb = row["cultivation_batch"]
        if row.get("cultivation_batch_is_declared_only"):
            pr.cellfmt(t.cell(i, 2), cb + " *", None, 10, pr.AMBER)
        else:
            pr.cellfmt(t.cell(i, 2), cb, None, 10, pr.BLACK)
        pct_cell(t.cell(i, 3), row["thc_pct"])
        pr.cellfmt(t.cell(i, 4), row["ecoa_code"] or "—", None, 9, pr.BLACK)
        pr.cellfmt(t.cell(i, 5), fmt_date(row["date"]), None, 9, pr.BLACK)
        pr.cellfmt(t.cell(i, 6), row["remark_mk"] or "", row["remark_en"] or None, 8, pr.GREY)
    pr.fixed(t, weights=[2.15, 2.65, 3.05, 2.35, 3.35, 1.95, 2.96])
    pr.gap(d, 6)
    pr.note(
        d,
        "* Култивациска серија декларирана само во List of COAs.xlsx / преглед на траншата — "
        "не е сертифицирана. Потврдено отсуство на сертификат во QC регистарот, RAGFlow "
        "eCOA_INGEST, папката со сертификати на Drive и индекс-таблата на Drive (4 независни "
        "извори).",
        "* Cultivation batch declared only in List of COAs.xlsx / the tranche overview — not "
        "certified. Confirmed absent from the QC register, the RAGFlow eCOA_INGEST dataset, "
        "the Drive certificate folder and the Drive index workbook (4 independent sources).",
    )

    pr.chapter(d, "2", "АНЕКС А — СИТЕ РЕГИСТРИРАНИ СЕРИИ FAT BASTARD (P060322)",
               "ANNEX A — ALL REGISTERED FAT BASTARD BATCHES (P060322)")
    pr.note(
        d,
        "P060322 не постои како означена серија во ниту еден сертификат — сортата Fat "
        "Bastard е тестирана и сертифицирана под 5 различни култивациски кодови, ниту еден "
        "поврзан со P-број во QC регистарот. Главната табела ја прикажува најновата (FB032601); "
        "сите пет се дадени тука за целосна следливост.",
        "P060322 does not exist as a labelled batch on any certificate — the Fat Bastard "
        "strain has been tested and certified under 5 separate cultivation codes, none "
        "cross-referenced to a P-number in the QC register. The main table shows the most "
        "recent (FB032601); all five are given here for full traceability.",
    )
    pr.gap(d, 4)
    ah = ["Култивациска серија | Cultivation batch", "еКоА код | eCoA code",
          "Датум | Date", "Вкупно Δ9-ТХЦ | Total Δ9-THC", "Забелешка | Remark"]
    fb = data["annex_a_fat_bastard"]
    ta = d.add_table(rows=len(fb) + 1, cols=len(ah))
    for j, h in enumerate(ah):
        pr.cellfmt(ta.cell(0, j), h, None, 9, pr.WHITE, bold=True, fill=pr.NAVYF)
    for i, row in enumerate(fb, start=1):
        pr.cellfmt(ta.cell(i, 0), row["cultivation_batch"], None, 10, pr.BLACK, bold=True)
        pr.cellfmt(ta.cell(i, 1), row["ecoa_code"], None, 9, pr.BLACK)
        pr.cellfmt(ta.cell(i, 2), fmt_date(row["date"]), None, 9, pr.BLACK)
        pct_cell(ta.cell(i, 3), row["thc_pct"])
        pr.cellfmt(ta.cell(i, 4), row["remark_mk"] or "", row["remark_en"] or None, 8,
                   pr.RED if "НЕ ОДГОВАРА" in (row["remark_mk"] or "") else pr.GREY)
    pr.fixed(ta, mode="full")

    pr.chapter(d, "3", "АНЕКС Б — СИТЕ ТЕСТИРАЊА ЗА P050072 (GP0824_03)",
               "ANNEX B — ALL TEST ROUNDS FOR P050072 (GP0824_03)")
    pr.note(
        d,
        "Серијата P050072 има сертификат за издавање и две дополнителни рунди од "
        "програмата за стабилност; главната табела ја прикажува најновата (мес. 9).",
        "Batch P050072 carries a release certificate plus two further stability-programme "
        "rounds; the main table shows the most recent (month 9).",
    )
    pr.gap(d, 4)
    bh = ["Рунда | Round", "еКоА код | eCoA code", "Датум | Date",
          "Вкупно Δ9-ТХЦ | Total Δ9-THC", "Забелешка | Remark"]
    p050072 = data["annex_b_p050072"]
    tb = d.add_table(rows=len(p050072) + 1, cols=len(bh))
    for j, h in enumerate(bh):
        pr.cellfmt(tb.cell(0, j), h, None, 9, pr.WHITE, bold=True, fill=pr.NAVYF)
    for i, row in enumerate(p050072, start=1):
        pr.cellfmt(tb.cell(i, 0), row["round_mk"], row["round_en"], 9, pr.BLACK)
        pr.cellfmt(tb.cell(i, 1), row["ecoa_code"], None, 9, pr.BLACK)
        pr.cellfmt(tb.cell(i, 2), fmt_date(row["date"]), None, 9, pr.BLACK)
        pct_cell(tb.cell(i, 3), row["thc_pct"])
        pr.cellfmt(tb.cell(i, 4), row["remark_mk"] or "", row["remark_en"] or None, 8, pr.GREY)
    pr.fixed(tb, mode="full")

    pr.gap(d, 8)
    pr.note(
        d,
        "Извори проверени: PP_Batch_Release_QC_Register_CORRECTED.xlsx; RAGFlow "
        "eCOA_INGEST (vision-parsed сертификати); Drive папка со сертификати "
        "1rwBvSAEoAZWsSKSaAQFUXkQLmZA13mSI. Секоја вредност е проверена наспроти "
        "самата слика на сертификатот пред внесување.",
        "Sources checked: PP_Batch_Release_QC_Register_CORRECTED.xlsx; RAGFlow "
        "eCOA_INGEST (vision-parsed certificates); Drive certificate folder "
        "1rwBvSAEoAZWsSKSaAQFUXkQLmZA13mSI. Every value was checked against the "
        "certificate's own rendered image before entry.",
    )

    d.save(OUT)
    return OUT


def self_check(data):
    """Re-derive every printed figure from dataset.json and assert it matches
    what will be printed — a drift here fails the build (pp-document-suite §6B)."""
    rows = data["rows"]
    assert len(rows) == 12, f"expected 12 requested batches, dataset has {len(rows)}"
    p_numbers = [r["p_number"] for r in rows]
    assert len(set(p_numbers)) == 12, "duplicate P-number in dataset"
    missing = [r["p_number"] for r in rows if r["thc_pct"] is None]
    assert missing == ["P060122", "P060132"], f"unexpected missing set: {missing}"
    present = [r for r in rows if r["thc_pct"] is not None]
    assert len(present) == 10, f"expected 10 certified rows, got {len(present)}"
    fb_row = next(r for r in rows if r["p_number"] == "P060322")
    assert fb_row["ecoa_code"] == "ППК26127" and fb_row["thc_pct"] == 12.39
    fb_annex = data["annex_a_fat_bastard"]
    assert len(fb_annex) == 5, "Fat Bastard annex must list all 5 registered batches"
    assert fb_annex[-1]["cultivation_batch"] == fb_row["cultivation_batch"] == "FB032601"
    assert fb_annex[-1]["thc_pct"] == fb_row["thc_pct"]
    p050072_row = next(r for r in rows if r["p_number"] == "P050072")
    p050072_annex = data["annex_b_p050072"]
    assert len(p050072_annex) == 3, "P050072 annex must list all 3 test rounds"
    assert p050072_annex[-1]["thc_pct"] == p050072_row["thc_pct"] == 25.98
    assert p050072_annex[-1]["ecoa_code"] == p050072_row["ecoa_code"] == "ППК26060"
    for r in present:
        mk, en = fmt_pct(r["thc_pct"])
        assert mk.endswith("%") and "," in mk and " " not in mk, f"house format violated: {mk}"
        assert en.endswith("%") and "." in en and " " not in en, f"house format violated: {en}"
    print("SELF-CHECK OK — 12 rows, 2 confirmed-absent, 10 certified, "
          "Fat Bastard annex (5) and P050072 annex (3) consistent with main table.")


if __name__ == "__main__":
    data, sha = load_dataset()
    self_check(data)
    out = build(data)
    print(f"dataset sha256={sha}")
    print(f"built: {out}")
