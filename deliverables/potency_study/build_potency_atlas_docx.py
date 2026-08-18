#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Potency Atlas — shareable Word edition (PP engine, informal/non-QMS).

A compact, presentable .docx mirror of the Atlas HTML's boards, for sharing
with colleagues: the grade rule, the three-column final board (original
specification | proposed grade ranges | names after renaming), the T1/T2/T3
stock table, and the sign-off block.

Single source of truth (§6B): every figure binds from potency_dataset.json;
the renamed-name group ladders are read from the already-built (and already
verified) Potency_Specs_and_Results.xlsx 'Atlas Grades — Renamed' sheet
rather than re-implementing the tier algorithm a fourth time.
"""
import json
import os
import re
import sys

sys.path.insert(0, "/root/.claude/skills/synced/pp-document-suite/scripts")
import pp_report as pr
import pp_assets
import openpyxl
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

D = os.path.dirname(os.path.abspath(__file__))
d = json.load(open(os.path.join(D, "potency_dataset.json"), encoding="utf-8"))
PM = json.load(open(os.path.join(D, "portfolio_master.json"), encoding="utf-8"))

NAVY = "2B547E"; INK = "16232B"; MUT = "4A5B6C"
GREEN = "14532B"; GFILL = "E9F4EC"; ROSE = "8A2E2E"; RFILL = "FDECEA"
GOLD = "8A6D14"; YFILL = "FFF8E7"; ZEBRA = "F4F7FB"

CANON = {"Cap Junkie": "Cap Junky", "CashCow": "Cash Cow", "GG4": "Gorilla Glue",
         "Grapes and Cream": "Graps & Creme", "Jelly Donuts": "Jelly Donutz",
         "Sleepy Joe": "Sleepy Joy", "Clemosa": "Clemosa a bud",
         "Wedding Crusher": "Wedding Crasher", "Appels & Bananas": "Apple and Banana",
         "Apples and Bananas": "Apple and Banana"}


def norm_b(b):
    b = (b or "").strip().upper().replace("OMP", "OPM")
    return re.sub(r"/0(\d)", r"/\1", b)


# renamed-name group ladders, from the verified workbook (no 4th algorithm copy)
wb = openpyxl.load_workbook(os.path.join(D, "Potency_Specs_and_Results.xlsx"), data_only=True)
ws = wb["Atlas Grades — Renamed"]
neu_tiers = {}
for row in ws.iter_rows(min_row=6):
    if not row[0].value:
        continue
    neu_tiers.setdefault(row[0].value, []).append(
        (float(row[2].value), float(row[3].value), float(row[4].value), float(row[5].value)))
for v in neu_tiers.values():
    v.sort()

doc = Document(pp_assets.TEMPLATE)
pr.wipe_body(doc)
pr.informal_header(doc, "АТЛАС НА ПОТЕНЦИЈА", "Potency Atlas",
                   "Неформален работен документ", "Informal working document")

pr.cover_page(
    doc, "Атлас на потенција", "Potency Atlas",
    [("Опфат | Scope",
      "Сите некогаш тестирани резултати за Вкупен Δ⁹-THC по сорта, предложените класи "
      "(номинала ± толеранција) за залихата Т1/Т2/Т3 и имињата по преименувањето | Every "
      "Total Δ⁹-THC result ever tested per strain, the proposed grade tiers (nominal ± "
      "tolerance) for the T1/T2/T3 stock, and the names after renaming"),
     ("Корпус | Corpus",
      "%d резултати од регистарот на сертификати · %d сорти · 78 серии на залиха | "
      "%d certificate-register results · %d strains · 78 stock batches"
      % (d["n_results"], d["n_strains"], d["n_results"], d["n_strains"])),
     ("Правило | Rule",
      "Класите се декларираат одозгора надолу: најсилната класа зема полни ±10,00% од "
      "номиналата, а секоја пониска се протега надолу до точно 0,01 под подот на класата "
      "над неа | Tiers are declared top-down: the strongest tier takes its full ±10.00%, "
      "each lower tier extends down to exactly 0.01 below the tier above"),
     ("Статус | Status",
      "Неформална работна анализа; меродавни остануваат лабораториските сертификати и "
      "издадените QCSP спецификации | Informal working analysis; the laboratory "
      "certificates and the issued QCSP specifications remain authoritative")],
    "Работен атлас", "Working atlas", "КК", "Quality Control",
    controlled=False)


def para(text, size=10, color=INK, italic=False, bold=False, before=2, after=2,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.06
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.font.italic = italic
    r.font.bold = bold
    r.font.color.rgb = RGBColor.from_string(color)
    return p


def shade(cell, hexv):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:color"), "auto")
    sh.set(qn("w:fill"), hexv)
    tcPr.append(sh)


def cell_lines(cell, lines, fill=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    """Write a multi-line cell: lines = [(text, size, color, bold, italic), ...]."""
    if fill:
        shade(cell, fill)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    first = True
    for text, size, color, bold, italic in lines:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.text = ""
        p.alignment = align
        p.paragraph_format.space_before = Pt(0.5)
        p.paragraph_format.space_after = Pt(0.5)
        r = p.add_run(text)
        r.font.name = "Calibri"
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.italic = italic
        r.font.color.rgb = RGBColor.from_string(color)


def hdr_cell(cell, text):
    cell_lines(cell, [(text, 8.2, "FFFFFF", True, False)], fill=NAVY,
               align=WD_ALIGN_PARAGRAPH.CENTER)


# ---------- 1. Grade rule ----------
pr.chapter(doc, "1", "ПРАВИЛО ЗА КЛАСИТЕ", "The Grade Rule")
para("Секоја класа е номинала ± толеранција: номиналата е секогаш на чекор од 0,50% "
     "(nn,00% или nn,50%), а толеранцијата никогаш не надминува 10,00% од номиналата. "
     "Класите се декларираат одозгора надолу, како што се одредуваат рачно: најсилната "
     "(највисоката) класа има приоритет и ја зема својата полна ширина ±10,00% — ништо не "
     "ја ограничува одозгора, па најсилната класа никогаш не се стеснува. Секоја пониска "
     "класа потоа се протега надолу од таа над неа: нејзината горна граница е точно 0,01 "
     "под долната граница на класата над неа (без празен простор), а зема онолку од "
     "своите ±10,00% колку што може додека допира до таа граница. Пониска класа е "
     "потесна од полните ±10% само кога точното допирање со класата над неа не остава "
     "повеќе простор — никогаш поради друга причина. | Each tier is a nominal ± "
     "tolerance: the nominal always sits on a 0.50% step (nn.00% or nn.50%) and the "
     "tolerance never exceeds 10.00% of the nominal. Tiers are declared top-down, the "
     "way they are set by hand: the strongest (highest) tier gets priority and takes its "
     "full ±10.00% — nothing constrains it from above, so the strongest grade is never "
     "squeezed. Each lower tier then extends downward from the one above: its ceiling "
     "sits exactly 0.01 below the floor of the tier above (no blind gap), taking as much "
     "of its own ±10.00% as it can while reaching that ceiling. A lower tier is narrower "
     "than its full ±10% only when meeting the tier above exactly leaves no more room — "
     "never for any other reason.", 9.5)
para("Сидро на серија = најновиот верификуван резултат (вредноста на тековниот CoQ за "
     "пуштање). Подот на ниту една класа не смее да е под 5,00% Вкупен THC. Само кога "
     "два тестирани резултати се толку далеку што ниту еден пар номинали не може да ги "
     "премости во рамки на 10,00%, останува вистински, означен јаз — не измислена "
     "преодна класа. | A batch's anchor = its most recent verified result (the value on "
     "the current release CoQ). No tier's floor may sit below 5.00% Total THC. Only "
     "where two tested results are too far apart for any nominal pair to bridge within "
     "the 10.00% cap does a genuine, flagged gap remain — never a fabricated bridge "
     "tier.", 9.5)
pr.databox(doc,
           [("Параметри: ±10,00% од номиналата (максимум) · номинала на чекор 0,50% · "
             "допир на класи точно 0,01 · под 5,00% · сидро = најнов резултат",
             "Parameters: ±10.00% of nominal (cap) · nominal on a 0.50% grid · tiers "
             "meet at exactly 0.01 · 5.00% floor · anchor = most recent result", 9, True)],
           GFILL, pr.GREEN)

# ---------- 2. Final board (3 columns) ----------
pr.chapter(doc, "2", "ФИНАЛНИ ОПСЕЗИ НА ПОТЕНЦИЈА ПО СОРТА",
           "Final Potency Grade Ranges per Strain")
para("Лева колона: оригиналната спецификациска сорта. Средна: предложените класи "
     "(номинала ± толеранција, опсег, серии). Десна: имињата што сериите ги носат по "
     "преименувањето (01_Portfolio_Master); каде сериите под новото име се групираат "
     "поинаку, неговите сопствени класи се дадени под името. | Left column: the "
     "original specification strain. Middle: the proposed tiers (nominal ± tolerance, "
     "range, batches). Right: the names the batches carry after renaming "
     "(01_Portfolio_Master); where a new name's batches cluster differently, its own "
     "tiers are given under that name.", 8.6, MUT, italic=True)

basis = {norm_b(b["batch"]): b["anchor"] is not None for b in d["stock"]}
by_origin = {}
neu_origins = {}
for r in PM:
    neu = (r.get("neu") or "").strip()
    if not neu:
        continue
    raw = (r.get("original") or "").strip()
    orig = CANON.get(raw, raw)
    by_origin.setdefault(orig, {}).setdefault(neu, []).append(r)
    neu_origins.setdefault(neu, set()).add(orig)

t = doc.add_table(rows=1, cols=3)
t.autofit = False
for c, hh in zip(t.rows[0].cells,
                 ["Оригинална спецификација | Original specification",
                  "Предложени опсези по класи | Proposed grade ranges",
                  "Имиња по преименување | Names after renaming"]):
    hdr_cell(c, hh)
pr._repeat_header(t)

n_def = n_prov = 0
for k, strain in enumerate(sorted(d["merged_ranges"])):
    tiers = d["merged_ranges"][strain]
    if not tiers:
        continue
    st = d["stats"].get(strain)
    sig_orig = tuple((x["nominal"], x["tol"]) for x in tiers)
    strain_prov = not any(basis.get(norm_b(b), False) for x in tiers for b in x["batches"])
    if strain_prov:
        n_prov += 1
    else:
        n_def += 1
    fill = ZEBRA if k % 2 else "FFFFFF"
    cells = t.add_row().cells

    c1 = [(strain, 10.5, NAVY, True, False)]
    c1.append((("Бр. тестирани резултати | results tested: %d" % st["n"]) if st
               else "без тестирања | no assays", 7.5, MUT, False, True))
    c1.append((("ПРОВИЗОРНО | PROVISIONAL" if strain_prov else
                "ДЕФИНИТИВНО | DEFINITIVE"), 7.5,
               ROSE if strain_prov else GOLD, True, False))
    cell_lines(cells[0], c1, fill=fill)

    c2 = []
    for i, x in enumerate(tiers):
        nb = len(x["batches"])
        c2.append(("Pot.-%d: %.2f%% ±%.2f%%  (%.2f%% – %.2f%%)"
                   % (i + 1, x["nominal"], x["tol"], x["range"][0], x["range"][1]),
                   9.5, GREEN, True, False))
        c2.append(("%s" % ", ".join(x["batches"]), 7.5, MUT, False, False))
        if x.get("gap_after") and i + 1 < len(tiers):
            c2.append(("Јаз — нема класа | Gap — no grade: %.2f%% – %.2f%%"
                       % (x["range"][1], tiers[i + 1]["range"][0]),
                       8, ROSE, True, False))
    cell_lines(cells[1], c2, fill=fill)

    c3 = []
    for neu, recs in (by_origin.get(strain) or {}).items():
        is_same = neu.strip().lower() == strain.strip().lower()
        brand = (recs[0].get("brand") or "").strip()
        line = neu + ((" · %s" % brand) if brand else "")
        line += " · %d %s" % (len(recs),
                              "серии | batches" if len(recs) != 1 else "серија | batch")
        if is_same:
            line += " · без промена | unchanged"
        c3.append((line, 8.6, INK if is_same else NAVY, not is_same, False))
        others = sorted(neu_origins.get(neu, set()) - {strain})
        if others:
            c3.append(("заедно со | with: %s" % ", ".join(others), 7.5, MUT, False, True))
        nt = neu_tiers.get(neu)
        if nt is not None:
            sig_neu = tuple((a, b) for a, b, _l, _h in nt)
            if others or sig_neu != sig_orig:
                for i, (nom, tol, lo, hi) in enumerate(nt):
                    c3.append(("   Pot.-%d: %.2f%% ±%.2f%% (%.2f%% – %.2f%%)"
                               % (i + 1, nom, tol, lo, hi), 7.5, GREEN, False, False))
    if not c3:
        c3 = [("— нема запис во мастерот | not in the Portfolio Master",
               7.5, MUT, False, True)]
    cell_lines(cells[2], c3, fill=fill)

for row in t.rows:
    for i, c in enumerate(row.cells):
        c.width = Cm([3.6, 8.9, 5.9][i])

para("%d сорти дефинитивно, %d провизорно (само декларирана основа). Формалното "
     "усвојување останува во спецификациите QCSP по редовна процедура. | %d strains "
     "definitive, %d provisional (declared basis only). Formal adoption remains with "
     "the QCSP specifications through the regular procedure."
     % (n_def, n_prov, n_def, n_prov), 8.6, MUT, italic=True)

# ---------- 3. Stock ----------
pr.chapter(doc, "3", "ЗАЛИХА Т1/Т2/Т3 — КЛАСА ПО СЕРИЈА",
           "T1/T2/T3 Stock — Grade per Batch")
para("Сидро = најнов верификуван резултат; каде нема ниту еден сертификат, основата е "
     "декларираната вредност (жолти редови — тестирајте пред формално декларирање). "
     "Простор надолу = растојание од сидрото до подот на класата, во процентни поени. | "
     "Anchor = most recent verified result; where no certificate exists the declared "
     "value is the basis (yellow rows — test before formal declaration). Downward "
     "headroom = distance from the anchor to the grade floor, in percentage points.",
     8.6, MUT, italic=True)

t2 = doc.add_table(rows=1, cols=7)
t2.autofit = False
for c, hh in zip(t2.rows[0].cells,
                 ["Т", "Серија | Batch", "Сорта | Strain", "Сидро | Anchor",
                  "Класа | Tier", "Номинала ± толер. | Nominal ± tol.",
                  "Простор ↓ (пп) | Headroom (pp)"]):
    hdr_cell(c, hh)
pr._repeat_header(t2)
k = 0
for b in sorted(d["stock"], key=lambda x: (x["tranche"], x["strain"], x["batch"])):
    if not b.get("proposed"):
        continue
    decl = b["anchor"] is None
    anc = ("%.2f%% (декл. | decl.)" % b["declared"]) if decl \
        else "%.2f%% (%s)" % (b["anchor"], b["anchor_date"])
    fill = YFILL if decl else (ZEBRA if k % 2 else "FFFFFF")
    k += 1
    vals = ["Т%s" % b["tranche"], b["batch"], b["strain"], anc,
            "Pot.-%s" % b.get("tier", "—"),
            "%.2f%% ± %.2f%%  (%.2f%% – %.2f%%)"
            % (b["nominal"], b["tol"], b["proposed"][0], b["proposed"][1]),
            "%.2f" % b["headroom_down"]]
    cells = t2.add_row().cells
    for i, v in enumerate(vals):
        cell_lines(cells[i], [(v, 8, INK, i == 5, False)], fill=fill,
                   align=WD_ALIGN_PARAGRAPH.CENTER if i != 2 else WD_ALIGN_PARAGRAPH.LEFT)
for row in t2.rows:
    for i, c in enumerate(row.cells):
        c.width = Cm([0.8, 2.9, 3.2, 3.2, 1.3, 4.5, 2.5][i])

# ---------- 4. Signatures ----------
pr.chapter(doc, "4", "ПОТПИСИ", "Signatures")
t3 = doc.add_table(rows=1, cols=2)
t3.autofit = False
roles = [("Изготвил и одобрил | Prepared and approved",
          "Менаџер за контрола на квалитет | Quality Control Manager"),
         ("Прегледал | Reviewed",
          "Менаџер за обезбедување квалитет | Quality Assurance Manager")]
for c, (role, title) in zip(t3.rows[0].cells, roles):
    lines = [(role, 10, NAVY, True, False), (title, 8.5, MUT, False, True)]
    for lbl in ("Име и презиме | Name", "Потпис | Signature", "Датум | Date"):
        lines.append((" ", 14, INK, False, False))
        lines.append(("________________________________", 10, INK, False, False))
        lines.append((lbl.upper(), 6.5, MUT, False, False))
    cell_lines(c, lines, fill="FFFFFF")
for row in t3.rows:
    for c in row.cells:
        c.width = Cm(9.2)
para("Неформален работен документ — потписите потврдуваат изработка, преглед и "
     "одобрување на предлогот; формалното усвојување останува во спецификациите QCSP "
     "по редовна процедура. | Informal working document — the signatures confirm "
     "preparation, review and approval of the proposal; formal adoption remains with "
     "the QCSP specifications through the regular procedure.", 8, MUT, italic=True)

# ---------- self-checks (§6B: printed values == dataset) ----------
tot_tiers = sum(len(v) for v in d["merged_ranges"].values())
assert tot_tiers == 43, tot_tiers
assert n_def + n_prov == len([s for s, v in d["merged_ranges"].items() if v]), (n_def, n_prov)
assert k == sum(1 for b in d["stock"] if b.get("proposed")), k
print("SELF-CHECK OK — %d tiers, %d strains (%d def / %d prov), %d stock rows"
      % (tot_tiers, n_def + n_prov, n_def, n_prov, k))

out = os.path.join(D, "Potency_Atlas.docx")
doc.save(out)

# run-spacing normalization (LibreOffice digit-spacing fix, as in the study builder)
doc2 = Document(out)
def fixrun(run):
    rPr = run._r.get_or_add_rPr()
    for tag in ("w:spacing", "w:kern"):
        for e in rPr.findall(qn(tag)):
            rPr.remove(e)
    for tag, val in (("w:spacing", "0"), ("w:kern", "0")):
        el = OxmlElement(tag)
        el.set(qn("w:val"), val)
        rPr.append(el)
def walk(cont):
    for p in cont.paragraphs:
        for r in p.runs:
            fixrun(r)
    for tb in cont.tables:
        for row in tb.rows:
            for c in row.cells:
                walk(c)
walk(doc2)
doc2.save(out)
print("wrote", out, os.path.getsize(out))
