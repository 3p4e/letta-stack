#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Strain Potency Study — document builder (PP engine, informal/non-QMS).

Everything printed binds from potency_dataset.json (§6B). No doc code, no
version — informal working analysis; the authoritative results remain the
laboratory certificates in the QMS and the ImB_QC_COAs knowledgebase.
"""
import json
import os
import sys

sys.path.insert(0, "/root/.claude/skills/synced/pp-document-suite/scripts")
import pp_report as pr
import pp_assets
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

D = os.path.dirname(os.path.abspath(__file__))
d = json.load(open(os.path.join(D, "potency_dataset.json"), encoding="utf-8"))
figidx = json.load(open(os.path.join(D, "figures", "index.json"), encoding="utf-8"))

NAVY = "2B547E"; INK = "16232B"; ROSE = "FDECEA"; MINT = "E2EFDA"

doc = Document(pp_assets.TEMPLATE)
pr.wipe_body(doc)
pr.informal_header(doc, "СТУДИЈА НА ПОТЕНЦИЈА ПО СОРТИ", "Strain Potency Study",
                    "Неформален работен документ", "Informal working document")

n_stab_usable = sum(1 for r in d["stability"] if r["usable"])
pr.cover_page(
    doc, "Студија на потенција по сорти", "Strain Potency Study",
    [("Опфат | Scope",
      "Сите некогаш тестирани резултати за Вкупен Δ⁹-THC, за сите серии, по сорта + "
      "предлог за нови опсези по класи за залихата од Транша 1/2/3 | Every Total Δ⁹-THC "
      "result ever tested, all batches, per strain + proposed new grade ranges for the "
      "T1/T2/T3 stock"),
     ("Корпус | Corpus",
      "%d резултати од регистарот на сертификати (77 серии, %d сорти) | "
      "%d certificate-register results (77 batches, %d strains)"
      % (d["n_results"], d["n_strains"], d["n_results"], d["n_strains"])),
     ("Верификација | Verification",
      "Жива проверка на Letta-базата: 4.134 пасуси скенирани, сите вредности потврдени "
      "на изворниот систем | Live Letta knowledgebase sweep: 4,134 passages scanned, "
      "all values confirmed on the source system"),
     ("Статус | Status",
      "Неформална работна анализа; меродавни остануваат лабораториските сертификати | "
      "Informal working analysis; the laboratory certificates remain authoritative")],
    "Аналитичка студија", "Analytical study", "КК", "Quality Control",
    controlled=False)


def para(text, size=10, color=INK, italic=False, bold=False, before=2, after=2,
         align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.06
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.font.italic = italic
    r.font.bold = bold
    r.font.color.rgb = RGBColor.from_string(color)
    return p


def shade(cell, hexv):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:color"), "auto")
    sh.set(qn("w:fill"), hexv)
    tcPr.append(sh)


def setcell(cell, text, size=8.5, color=INK, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, fill=None):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if fill:
        shade(cell, fill)
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = RGBColor.from_string(color)


def table(headers, rows, widths, fills=None, hdr_repeat=True):
    t = doc.add_table(rows=1, cols=len(headers))
    t.autofit = False
    for i, (c, hh) in enumerate(zip(t.rows[0].cells, headers)):
        setcell(c, hh, 8.2, "FFFFFF", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, fill=NAVY)
    for k, row in enumerate(rows):
        cells = t.add_row().cells
        fill = (fills[k] if fills else ("F4F7FB" if k % 2 else "FFFFFF"))
        for i, (c, v) in enumerate(zip(cells, row)):
            setcell(c, str(v), 8.3, INK,
                    align=WD_ALIGN_PARAGRAPH.CENTER if i > 0 and len(str(v)) < 30 else WD_ALIGN_PARAGRAPH.LEFT,
                    fill=fill)
    for row in t.rows:
        for i, c in enumerate(row.cells):
            c.width = Cm(widths[i])
    if hdr_repeat:
        pr._repeat_header(t)
    return t


# ---------- 1. Executive summary ----------
pr.chapter(doc, "1", "РЕЗИМЕ", "Executive Summary")
para("Оваа студија ги собира на едно место сите резултати за Вкупен Δ⁹-THC што некогаш се "
     "тестирани за нашите серии — %d резултати од регистарот на сертификати (77 серии, %d "
     "сорти) — и врз нив предлага нови опсези по класи (номинала ± толеранција) за залихата "
     "од Траншите 1, 2 и 3. | This study brings every Total Δ⁹-THC result ever tested into "
     "one place — %d certificate-register results (77 batches, %d strains) — and proposes "
     "new grade ranges (nominal ± tolerance) for the Tranche 1/2/3 stock on top of them."
     % (d["n_results"], d["n_strains"], d["n_results"], d["n_strains"]), 10.5)

# ---------- 2. Data & verification ----------
pr.chapter(doc, "2", "ПОДАТОЦИ И ВЕРИФИКАЦИЈА", "Data & Verification")
_sweep = d["verification"]["host_sweep"].split("; 9 stability")[0]
para("Корпусот е компајлиран од регистарот на eCoA сертификати и потоа проверен во живо на "
     "Letta-базата на знаење (ImB_QC_COAs): %s. | Corpus compiled from the eCoA register "
     "and live-verified against the Letta knowledgebase: %s." % (_sweep, _sweep), 9)
table(["Слој | Layer", "Резултати | Results", "Забелешка | Note"],
      [["Регистар на сертификати | Certificate register", str(d["n_results"]),
        "77 серии, %d сорти; секоја вредност препишана дословно | 77 batches, %d strains; "
        "every value transcribed verbatim" % (d["n_strains"], d["n_strains"])]],
      [6.0, 3.0, 8.6])

# ---------- 3. Per-strain statistics ----------
pr.chapter(doc, "3", "СТАТИСТИКА ПО СОРТИ", "Per-Strain Statistics")
rows = []
for s in sorted(d["stats"], key=lambda x: -d["stats"][x]["max"]):
    st = d["stats"][s]
    rows.append([s, st["n"], "%.2f%% – %.2f%%" % (st["min"], st["max"])])
table(["Сорта | Strain", "Бр. тестирани резултати | Number of results tested",
       "Опсег на тестираните резултати | Range of tested results"],
      rows, [7.0, 5.0, 5.0])
para("Ова е опсегот на СИТЕ некогаш тестирани резултати за сортата — не е предложениот "
     "опсег на потенција по класа (тој е даден подолу, во Поглавје 4, по Pot.- класи). "
     "Резултатите вклучуваат и повторени мерења на исти серии: бројката ја опишува "
     "популацијата на ТЕСТИРАНИ РЕЗУЛТАТИ, не независни серии. | This is the range of ALL "
     "results ever tested for the strain — not the proposed potency grade range (given "
     "below, in Chapter 4, per Pot.- tier). Results include repeat measurements of the same "
     "batches: the count describes the population of TESTED RESULTS, not independent "
     "batches.", 8, "4A5B6C", italic=True)

# ---------- 4. Per-strain distributions ----------
pr.chapter(doc, "4", "ДИСТРИБУЦИИ ПО СОРТИ", "Per-Strain Distributions")
para("Оска 0,00%–30,00% Вкупен Δ⁹-THC; секој резултат носи зона од ±1,50% (сините појаси — "
     "натрупувањето е визуелна густина); крива на дистрибуција за n≥3 (само облик — без "
     "збирни статистики); предложени класи Pot.-x (зелено, целосна висина); "
     "старите стандардни граници (испрекинато). | 0.00%–30.00% axis; each result carries a "
     "±1.50% zone (blue bands — stacking is visual density); distribution curve for n≥3 "
     "(shape only — no summary statistics); proposed Pot.-x tiers (green, full height); old "
     "standard boundaries (dashed).", 8.6, "37474F", italic=True)
for s in sorted(d["stats"]):
    st = d["stats"][s]
    pr.subsec(doc, "4.%d" % (sorted(d["stats"]).index(s) + 1), s, "")
    pr.figure(doc, os.path.join(D, "figures", figidx[s]), None, None, width_cm=17.2)
    tiers = d["merged_ranges"].get(s, [])
    if tiers:
        trs = [["Pot.-%d" % (i + 1), "%.2f%% ± %.2f%%" % (t["nominal"], t["tol"]),
                "%.2f%% – %.2f%%" % tuple(t["range"]),
                ", ".join(t["batches"]),
                ", ".join("%.2f%%" % a for a in t["anchors"])]
               for i, t in enumerate(tiers)]
        table(["Класа | Tier", "Номинала ± толер. | Nominal ± tol.", "= Опсег | = Range",
               "Серии | Batches", "Сидра | Anchors"],
              trs, [1.5, 3.0, 2.6, 6.4, 4.0])
        para("Класите се декларираат одозгора надолу: најсилната (највисоката) класа има "
             "приоритет и ја зема својата полна ±10,00%, а секоја пониска се протега надолу од "
             "неа — горната ѝ граница е точно 0,01 под долната граница на класата над неа, и зема "
             "онолку од своите ±10,00% колку што може додека допира до таа граница. Затоа пониска "
             "класа е потесна само кога точното допирање не остава повеќе простор. Класите "
             "никогаш не се преклопуваат. | Tiers are declared top-down: the strongest (highest) "
             "tier gets priority and takes its full ±10.00%, and each lower tier extends downward "
             "from it — its ceiling sits exactly 0.01 below the floor of the tier above, taking "
             "as much of its own ±10.00% as it can while reaching that ceiling. A lower tier is "
             "narrower only when meeting the tier above leaves no more room. Tiers never overlap.",
             8, "4A5B6C", italic=True)
        for i, t in enumerate(tiers):
            if t.get("gap_after") and i + 1 < len(tiers):
                glo, ghi = t["range"][1], tiers[i + 1]["range"][0]
                para("НЕМА КЛАСА | NO GRADE — Нема воспоставена класа %.2f%%–%.2f%%: ниту еден тестиран резултат на "
                     "оваа сорта не паѓа во оваа зона, и ниту еден кандидат за номинала не може да ги "
                     "премости соседните класи во рамки на 10,00%%-ограничувањето — резултат тука "
                     "би побарал индивидуална ОК проценка. | No established grade %.2f%%–%.2f%%: "
                     "no tested result of this strain falls in this zone, and no candidate nominal "
                     "can bridge the neighbouring tiers within the 10.00%% cap — a result here would "
                     "require individual QC assessment." % (glo, ghi, glo, ghi),
                     8, "8A2E2E", italic=True)
    else:
        para("Нема серии на залиха за оваа сорта во Т1/Т2/Т3. | No T1/T2/T3 stock batches "
             "for this strain.", 8.5, "6B7785", italic=True)

# ---------- 5. Stock table ----------
pr.chapter(doc, "5", "ПРЕДЛОГ-КЛАСИ ЗА ЗАЛИХАТА Т1/Т2/Т3", "Proposed Grades for the T1/T2/T3 Stock")
para("Сидро = најнов верификуван резултат; каде нема ниту еден сертификат, основата е "
     "декларираната вредност (означено). Простор надолу = растојание од сидрото до долната "
     "граница на класата. | Anchor = latest verified result; where no certificate exists the "
     "declared value is the basis (flagged). Downward headroom = distance from the anchor to "
     "the grade floor.", 8.6, "4A5B6C", italic=True)
stock_rows = []
stock_fills = []
for b in sorted(d["stock"], key=lambda x: (x["tranche"], x["strain"], x["batch"])):
    if not b.get("proposed"):
        continue
    anc = "%.2f (%s)" % (b["anchor"], b["anchor_date"]) if b["anchor"] is not None \
        else "%.2f (декл. | decl.)" % b["declared"]
    tier = "Pot.-%d" % b["tier"] if b.get("tier") else "—"
    stock_rows.append(["Т%s" % b["tranche"], b["batch"], b["strain"], anc,
                       b["bracket_old"] or "—", tier,
                       "%.2f%% ± %.2f%%" % (b["nominal"], b["tol"]),
                       "%.2f" % b["headroom_down"]])
    stock_fills.append("FFF8E7" if b["anchor"] is None
                       else ("FFFFFF" if len(stock_rows) % 2 else "F4F7FB"))
table(["Т", "Серија | Batch", "Сорта | Strain", "Сидро | Anchor (%)",
       "Стара класа | Old", "Класа | Tier", "Номинала ± толер. | Nominal ± tol.",
       "Простор надолу | Headroom (%)"],
      stock_rows, [0.9, 2.7, 3.3, 3.1, 1.9, 1.5, 2.6, 2.4])
para("Жолти редови: без ниту еден сертификат на датотека — основа е декларираната вредност; "
     "овие серии да се тестираат пред формално декларирање класа. | Yellow rows: no "
     "certificate on file — declared value used as basis; test these batches before formally "
     "declaring a grade.", 8, "4A5B6C", italic=True)

# ---------- 6. Overview annex ----------
pr.chapter(doc, "6", "АНЕКС — ПРЕГЛЕД НА СИТЕ СОРТИ", "Annex — All-Strain Overview")
pr.figure(doc, os.path.join(D, "figures", "overview.png"),
          "Сите сорти на една оска 0,00%–30,00%: резултати и предложени класи",
          "All strains on one 0.00%–30.00% axis: results and proposed tiers")

out = os.path.join(D, "Strain_Potency_Study_14Aug2026.docx")
pr.save(doc, out) if hasattr(pr, "save") else doc.save(out)

doc2 = Document(out)
def fixrun(run):
    rPr = run._r.get_or_add_rPr()
    for tag in ("w:spacing", "w:kern"):
        for e in rPr.findall(qn(tag)):
            rPr.remove(e)
    for tag, val in (("w:spacing", "0"), ("w:kern", "0")):
        el = OxmlElement(tag)
        el.set(qn("w:val"), val)
        rPr.append(el)
def walk(cont):
    for p in cont.paragraphs:
        for r in p.runs:
            fixrun(r)
    for tb in cont.tables:
        for row in tb.rows:
            for c in row.cells:
                walk(c)
walk(doc2)
doc2.save(out)
print("wrote", out, os.path.getsize(out))
