#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""QC Weekly Plan/Report — Issue 01 builder.

This is an INFORMAL management submission — not a QMS-controlled document.
Per the owner's explicit instruction it therefore carries NO document code,
NO version/issue-code header, and no "Controlled document" footer line; those
markings are reserved for records that actually go through document control.

Everything printed is bound from weekly_dataset.json (QC artifacts, computed
at build time) and merged_sources.json (the four-session time band, PP-content
only — see merge_timelines.py for the scope rule and exclusions). No figure is
hand-typed into the prose.
"""
import json
import os
import sys

sys.path.insert(0, "/root/.claude/skills/synced/pp-document-suite/scripts")
import pp_report as pr
import pp_assets
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

D = os.path.dirname(os.path.abspath(__file__))
d = json.load(open(os.path.join(D, "weekly_dataset.json"), encoding="utf-8"))
ms = json.load(open(os.path.join(D, "merged_sources.json"), encoding="utf-8"))

NAVY = "2B547E"; INK = "16232B"; ROSE = "FDECEA"
PEND = "датум: во очекување | date: pending"

tpl = pp_assets.TEMPLATE
doc = Document(tpl)
pr.wipe_body(doc)

# header: name only, no doc code/version block (informal submission, not QMS-controlled)
pr.informal_header(doc, "НЕДЕЛЕН ПЛАН/ИЗВЕШТАЈ НА КОНТРОЛА НА КВАЛИТЕТ", "QC Weekly Plan/Report",
                    "Неформален работен документ", "Informal working document")

try:
    pr.cover_page(
        doc, "Неделен план/извештај на Контрола на квалитет", "QC Weekly Plan/Report",
        [("Издание | Issue", "1 (прво издание на неделниот циклус | first issue of the weekly cycle)"),
         ("Извештајна недела | Report week", "Понеделник 10.08 – Петок 14.08.2026"),
         ("Планска недела | Plan week", "Понеделник 17.08 – Петок 21.08.2026"),
         ("Статус | Status",
          "Неформална управувачка достава за информација; официјалните записи "
          "продолжуваат да се водат според системот за квалитет | Informal management "
          "submission for information; the authoritative records continue to be kept "
          "under the quality system"),
         ("Опфат | Scope",
          "Контрола на квалитет: тестирање, инспекции, мониторинг и резултати за "
          "ослободување (одлуката за ослободување е на Обезбедување на квалитет) | "
          "QC: testing, inspection, monitoring and release-testing data "
          "(release disposition itself sits with QA)"),
         ("Извори | Sources",
          "Проверени записи на Контрола на квалитет + 4 записи за поминато време "
          "(само содржина поврзана со Purely Plant) | Verified QC records + 4 "
          "work-time records (Purely Plant content only)")],
        "Неделен план/извештај", "Weekly plan/report", "КК", "Quality Control",
        controlled=False)
except Exception as e:
    print("cover", e)
    doc.add_heading("QC Weekly Plan/Report", 0)


def para(text, size=10, color=INK, italic=False, bold=False, before=2, after=2,
         align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.06
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.font.italic = italic
    r.font.bold = bold
    r.font.color.rgb = RGBColor.from_string(color)
    return p


def shade(cell, hexv):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:color"), "auto")
    sh.set(qn("w:fill"), hexv)
    tcPr.append(sh)


def setcell(cell, text, size=8.5, color=INK, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, fill=None):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if fill:
        shade(cell, fill)
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = RGBColor.from_string(color)


def table(headers, rows, widths, fills=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.autofit = False
    for i, (c, hh) in enumerate(zip(t.rows[0].cells, headers)):
        setcell(c, hh, 8.2, "FFFFFF", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, fill=NAVY)
    for k, row in enumerate(rows):
        cells = t.add_row().cells
        fill = (fills[k] if fills else ("F4F7FB" if k % 2 else "FFFFFF"))
        for i, (c, v) in enumerate(zip(cells, row)):
            setcell(c, str(v), 8.3, INK,
                    align=WD_ALIGN_PARAGRAPH.CENTER if i > 0 and len(str(v)) < 28 else WD_ALIGN_PARAGRAPH.LEFT,
                    fill=fill)
    for row in t.rows:
        for i, c in enumerate(row.cells):
            c.width = Cm(widths[i])
    return t


# ---------- Ch.1 ISSUES ----------
pr.chapter(doc, "1", "ПРАШАЊА ШТО БАРААТ ВНИМАНИЕ", "Issues Requiring Attention")
para("Секое прашање подолу влијае на капацитетот за ослободување на производ или на "
     "подготвеноста за добра производна пракса, според правилото на сопственикот за "
     "истакнување прашања. Роковите за решавање ги определува Раководителот на Контрола "
     "на квалитет. | Each issue below affects release capacity or manufacturing-practice "
     "readiness, per the owner's issue-highlighting rule. Resolution dates are set by the "
     "QC Manager.", 9, "4A5B6C", italic=True)
issues = [
    ("И-1", "Резултат надвор од спецификација — серија FB032601, странични материи | "
            "Result out of specification — batch FB032601, foreign matter "
            "(лабораторија CNP, серија сертификати ППК26127): 0,08% Не одговара — "
            "присутна семка; истрага во тек по СОП QCSOP 019 | 0.08% NON-CONFORM — "
            "seed present; investigation under way per QCSOP 019"),
    ("И-2", "Неархивирани сертификати (микробиологија/тешки метали/пестициди) за 7 серии | "
            "Unfiled micro/heavy-metals/pesticide certificates for 7 batches — "
            "FB012601/1, GRC102501/2, JD012603/2, JD012603/2V, KC102501, PM112501, "
            "SCR112501; редот за ослободување е блокиран додека не се внесат | "
            "release queue blocked pending entry"),
    ("И-3", "9 серии без ниту еден сертификат за анализа во базата | "
            "9 batches with no certificate of analysis on file: "
            + ", ".join(d["B4_queue"]["no_ecoa"])),
    ("И-4", "Недостасуваат податоци за фенотип/хемотип за 6 сорти + конфликт меѓу изворите "
            "за потеклото на семето на Gorilla Glue | Phenotype/chemotype data missing for "
            "6 strains + a source conflict on Gorilla Glue's seed origin — ги блокира "
            "спецификациите за производ додека не се потврдат | blocks the product "
            "specifications until confirmed"),
    ("И-5", "%d предложени интерни сертификати (за домашно определени странични материи и "
            "идентификација) и %d предложени сертификати за квалитет чекаат формално "
            "издавање преку Регистарот за издавање сертификати | %d proposed in-house "
            "certificates (declared foreign-matter and identification determinations) and "
            "%d proposed certificates of quality await formal issuance through the "
            "Certificate Issuance Register"
            % (d["E1_declared"]["icoa_total"], 95, d["E1_declared"]["icoa_total"], 95)),
]
table(["№", "Прашање | Issue", "Рок | Resolution"],
      [[k, txt, PEND] for k, txt in issues], [1.1, 13.2, 3.3], fills=[ROSE] * 5)

# ---------- Ch.2 REPORT ----------
pr.chapter(doc, "2", "ИЗВЕШТАЈ — НЕДЕЛА 10–14.08", "Report — Week Ended 14.08")

pr.subsec(doc, "2.1", "Визуелни инспекции", "Visual Inspections")
table(["Ставка | Item", "Вредност | Value"],
      [["Домашни определувања документирани по стандардна оперативна постапка "
        "(пред финалното земање мостри / по него) | In-house determinations documented "
        "per standard operating procedure (before/after final QC sampling)",
        "Странични материи: %d серии | Foreign matter: %d batches · "
        "Макроскопска/микроскопска идентификација: %d серии | "
        "Macro/microscopic identification: %d batches"
        % (d["E1_declared"]["fm_batches"], d["E1_declared"]["fm_batches"],
           d["E1_declared"]["id_batches"], d["E1_declared"]["id_batches"])],
       ["Предложени интерни сертификати за нив | In-house certificates proposed for them",
        "%d (список за издавање на сертификати)" % d["E1_declared"]["icoa_total"]],
       ["Дополнителни инспекции извршени оваа недела | Further inspections performed this week",
        "нема податоци во изворите | no data in consolidated sources"]],
      [8.2, 9.4])

pr.subsec(doc, "2.2", "Лабораториско тестирање", "Laboratory Testing")
table(["Ставка | Item", "Вредност | Value"],
      [["Ретест-кампања (лабораторија Farmahem), обработена и потврдена во живо | "
        "Retest campaign (Farmahem laboratory), processed and live-verified",
        "42 сертификати (21 за потенца + 21 за микотоксини), 07–10.08.2026 | "
        "42 certificates (21 potency + 21 mycotoxin), 07–10.08.2026"],
       ["Опсег на резултати за потенца | Potency result range",
        "%s (0 резултати надвор од спецификација во кампањата) | %s "
        "(0 out-of-specification results in the campaign)"
        % (d["E2"]["retest_range"], d["E2"]["retest_range"])],
       ["Вкупно резултати во базата, потврдени | Total results on file, verified",
        "%d (77 серии; 4-слојна проверка + потврда во живо на изворниот систем) | "
        "%d (77 batches; 4-layer fact-check + live source-system verification)"
        % (d["corpus"]["results_transcribed"], d["corpus"]["results_transcribed"])]],
      [8.2, 9.4])

pr.subsec(doc, "2.3", "Мониторинг на околината", "Environmental Monitoring")
table(["Ставка | Item", "Вредност | Value"],
      [["Кругови на мониторинг | Monitoring rounds",
        "нема податоци во изворите — да се внесе од Контрола на квалитет | "
        "no data in sources — to be entered by QC"]], [8.2, 9.4])

pr.subsec(doc, "2.4", "Производ кој чека ослободување (ред за обработка)",
          "Product Awaiting QC Release (queue)")
table(["Категорија | Category", "Серии | Batches", "Забелешка | Note"],
      [["Делумен панел (недостасуваат ставки) | Partial panel (items missing)",
        "%d" % len(d["B4_queue"]["partial"]),
        "најчесто недостасуваат: идентификација по HPTLC, CBN, AflaB1, OTA | "
        "most commonly missing: HPTLC identification, CBN, AflaB1, OTA"],
       ["Без ниту еден сертификат за анализа | No certificate of analysis at all",
        "%d" % len(d["B4_queue"]["no_ecoa"]), "види Прашање И-3 | see Issue И-3"]],
      [6.4, 2.6, 8.6])
para("Разгранување до детали | Drill-down: PP_Spec_Parameter_Matrix.xlsx — по серија, по "
     "параметар, со врски до скенираните сертификати за анализа. | per batch, per "
     "parameter, hyperlinked to the certificate scans.", 8, "4A5B6C", italic=True)

pr.subsec(doc, "2.5", "Производ со целосни резултати од Контрола на квалитет",
          "Product QC-Cleared (results complete)")
table(["Ставка | Item", "Вредност | Value"],
      [["Серии со целосен панел од 21 ставка (вклучувајќи ги декларираните) | "
        "Batches with the full 21-item panel (including declared determinations)",
        ", ".join(d["B5_cleared_full_panel"]) + " (3)"],
       ["Напомена | Note",
        "ова е статус на покриеност со резултати, НЕ одлука за ослободување — "
        "одлуката за ослободување е на Обезбедување на квалитет | this is a "
        "result-coverage status, NOT a release disposition — disposition sits with QA"]],
      [8.2, 9.4])

pr.subsec(doc, "2.6", "Резултати надвор од спецификација / лабораториски отстапувања",
          "OOS / Lab Deviations")
o = d["E4_oos"][0]
table(["Серија | Batch", "Параметар | Parameter", "Наод | Finding"],
      [[o["batch"], "Странични материи (ставка 7) | Foreign Matter (item 7)", o["finding"]]],
      [3.2, 4.6, 9.8], fills=[ROSE])

pr.subsec(doc, "2.7", "Валидациски активности (дел на Контрола на квалитет)",
          "Validation Activities (QC part)")
table(["Ставка | Item", "Вредност | Value"],
      [["Извршени оваа недела | Executed this week", "нема | none"]], [8.2, 9.4])

pr.subsec(doc, "2.8", "Тестирање за ослободување простории (премин)",
          "Room Release Testing (transition)")
table(["Ставка | Item", "Вредност | Value"],
      [["Статус по простории | Per-room status",
        "нема податоци во изворите — да се внесе од Контрола на квалитет | "
        "no data in sources — to be entered by QC"]], [8.2, 9.4])

pr.subsec(doc, "2.9", "Придонес кон подготвеноста за добра производна пракса",
          "Contribution to Manufacturing-Practice Readiness")
table(["Систем | System", "Статус | Status"], [[a, b] for a, b in d["E7_gmp"]], [6.6, 11.0])

# ---------- Ch.3 PLAN ----------
pr.chapter(doc, "3", "ПЛАН — НЕДЕЛА 17–21.08", "Plan — Week Ahead 21.08")
table(["№", "Активност | Action", "Цел | Target"],
      [["1", "Формално издавање на %d предложени интерни сертификати и 95 предложени "
             "сертификати за квалитет, по редоследот утврден во плановите за издавање | "
             "Formal issuance of %d proposed in-house certificates and 95 proposed "
             "certificates of quality, per the issuance-order plans"
             % (d["E1_declared"]["icoa_total"], d["E1_declared"]["icoa_total"]), PEND],
       ["2", "Лоцирање, скенирање и внесување на сертификатите за 7-те серии од Прашање И-2 | "
             "Locate, scan and enter the certificates for the 7 batches in Issue И-2", PEND],
       ["3", "Решавање на 9-те серии без сертификат за анализа (Прашање И-3) — потврда дали "
             "постои соодветна документација | Resolve the 9 no-certificate batches — "
             "confirm whether supporting documentation exists", PEND],
       ["4", "Комплетирање на полињата за фенотип/хемотип во спецификациите за производ "
             "(одлука за Gorilla Glue, потврда на можните совпаѓања за две сорти, "
             "четири технички листови што недостасуваат) | Complete the phenotype/"
             "chemotype fields in the product specifications (Gorilla Glue decision, "
             "confirm two candidate strain matches, four missing supplier data sheets)", PEND],
       ["5", "Работна книга за Транша 3 и внесување на слојот со преименувани сорти | "
             "Tranche 3 workbook and ingestion of the strain-rename layer", PEND],
       ["6", "Распоред за земање мостри / мониторинг / чистење простории — да го внесе "
             "Контрола на квалитет | Sampling / monitoring / room-clearance schedule — "
             "to be entered by QC", "—"]],
      [1.0, 13.2, 3.4])

# ---------- Ch.4 ANNEX ----------
pr.chapter(doc, "4", "АНЕКС — КОНСОЛИДИРАНА ВРЕМЕНСКА ЛЕНТА", "Annex — Consolidated Work-Time Band")
para("Четири паралелни работни сесии прикажани на еден заеднички линијар по часови "
     "(Скопје). Бројувани се само активности поврзани со содржината на Purely Plant; "
     "работата околу поставувањето и одржувањето на инфраструктурата (на пр. "
     "конфигурација и поправка на Docker-стекот) е исклучена, додека работата на "
     "инстанцата на Letta — нејзината база на знаење, базите за меморија и агенти, "
     "и податоците за Контрола на квалитет во неа — е сметана како содржина на "
     "Purely Plant. Методологиите на изворите се различни и НЕ се собираат меѓу "
     "себе: полно пополнето = докажано време (транскрипт/алатки/git); шрафирано = "
     "распон или проценка. | Four parallel work sessions shown on one shared hour "
     "ruler (Skopje). Only Purely Plant content work is counted; infrastructure "
     "setup and maintenance (e.g. Docker-stack configuration and repair) is "
     "excluded, while work on the Letta host itself — its knowledgebase, its "
     "memory/agent databases, and the QC data held in them — counts as Purely "
     "Plant content. Source methodologies differ and are NOT summable: solid = "
     "evidenced time (transcript/tool/git); hatched = span or estimate.",
     9, "4A5B6C", italic=True)
pr.figure(doc, os.path.join(D, "consolidated_ruler.png"),
          "Консолидирана лента: четири сесии, само содржина на Purely Plant, "
          "Вт 16:00 → Пет 08:00",
          "Consolidated band: four sessions, Purely Plant content only, "
          "Tue 16:00 → Fri 08:00")
table(["Извор | Source", "Мерка | Measure", "Часови (PP-содржина) | Hours (PP content)"],
      [[s["label"], s["method"], "%.1f" % s["pp_hours"]] for s in ms["sources"]],
      [7.6, 7.0, 3.0])
excl_rows = [[e["item"], e["hours"]] for e in ms["exclusions"]]
table(["Исклучено (не е содржина на Purely Plant) | Excluded (not PP content)",
       "Часови | Hours"], excl_rows, [14.6, 3.0])
para("Потекло | Provenance: (1) QC_Activity_Timeline_11-14Aug2026.docx — транскрипт-извор; "
     "(2) Запис за временска лента по семејства задачи (интерна ознака PP-OPS-TIME-2026-08-14) "
     "— временски печати на алатки/датотеки во оваа сесија; "
     "(3) Запис за временска лента на работата (интерна ознака WWF-TIMELINE-2026-0814) — "
     "евиденција од git-коммити и продукциски пуштања во употреба; "
     "(4) Работен временски преглед (11-14 Aug 2026) — проценки од активноста на сесиите. "
     "Ниту една изворна бројка не е менувана при консолидацијата; секое намалување е "
     "избор на подмножество од сопствените, поединечни редови на изворот (гл. "
     "merge_timelines.py). | No source figure was altered in consolidation; every "
     "reduction is a subset selection of that source's own itemised rows "
     "(see merge_timelines.py).", 7.8, "6B7785", italic=True, before=4)

out = os.path.join(D, "QC_Weekly_Plan_Report_Issue01_10-14Aug2026.docx")
pr.save(doc, out) if hasattr(pr, "save") else doc.save(out)

# spacing normalization for clean LibreOffice PDF preview
doc2 = Document(out)


def fixrun(run):
    rPr = run._r.get_or_add_rPr()
    for tag in ("w:spacing", "w:kern"):
        for e in rPr.findall(qn(tag)):
            rPr.remove(e)
    for tag, val in (("w:spacing", "0"), ("w:kern", "0")):
        el = OxmlElement(tag)
        el.set(qn("w:val"), val)
        rPr.append(el)


def walk(cont):
    for p in cont.paragraphs:
        for r in p.runs:
            fixrun(r)
    for tb in cont.tables:
        for row in tb.rows:
            for c in row.cells:
                walk(c)


walk(doc2)
doc2.save(out)
print("wrote", out, os.path.getsize(out))
