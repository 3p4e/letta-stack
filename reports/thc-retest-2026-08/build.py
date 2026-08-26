# -*- coding: utf-8 -*-
"""PP-QC-THC-VER-001/2026 — Total Δ9-THC re-test vs previous release result, evaluated
   against the per-strain grade specification band. Data-driven build (§6B): bind → compute
   → inject → assert."""
import sys, json, hashlib, os
SK="/root/.claude/skills/synced/pp-document-suite"
sys.path.insert(0, SK+"/scripts")
import pp_report as pr
from docx import Document
import figs

# ---------- 1. BIND -------------------------------------------------------
DS="thc_dataset.json"
RAW=open(DS,'rb').read()
PROV=dict(file=os.path.basename(DS), sha256=hashlib.sha256(RAW).hexdigest(), bytes=len(RAW))
R=json.loads(RAW.decode())

# ---------- 2. COMPUTE ----------------------------------------------------
GRADES=list(range(8,29,2))
def band(N): return (0.9*N, 1.1*N)
def bandstr(N):
    lo,hi=band(N); return "%.1f – %.1f" % (lo,hi)
def bandstr_mk(N):
    lo,hi=band(N); return ("%.1f – %.1f"%(lo,hi)).replace(".",",")
def grade(v):
    if v is None: return None
    c=[N for N in GRADES if band(N)[0]-1e-9<=v<=band(N)[1]+1e-9]
    return min(c,key=lambda N:abs(v-N)) if c else None
def mk(x, dp=2):
    return ("%.*f"%(dp,x)).replace(".",",")

for r in R:
    r["g_prev"]=grade(r["prev"]); r["g_ret"]=grade(r["retest"])
    if r["prev"] is None:
        r["status"]="no_prior"; r["delta"]=None
    else:
        lo,hi=band(r["g_prev"])
        r["status"]="retained" if lo-1e-9<=r["retest"]<=hi+1e-9 else "changed"
        r["delta"]=round(r["retest"]-r["prev"],2)

COMP=[r for r in R if r["prev"] is not None]
NOPRIOR=[r for r in R if r["prev"] is None]
HELD=[r for r in COMP if r["status"]=="retained"]
MOVED=[r for r in COMP if r["status"]=="changed"]
UP=[r for r in MOVED if r["g_ret"]>r["g_prev"]]
DOWN=[r for r in MOVED if r["g_ret"]<r["g_prev"]]
N,NC,NH,NM,NU,ND,NP = len(R),len(COMP),len(HELD),len(MOVED),len(UP),len(DOWN),len(NOPRIOR)
deltas=[r["delta"] for r in COMP]
mean_d=sum(deltas)/len(deltas)
absd=sorted(abs(d) for d in deltas); med_abs=(absd[len(absd)//2] if len(absd)%2 else (absd[len(absd)//2-1]+absd[len(absd)//2])/2)
gt3=[r for r in COMP if abs(r["delta"])>3.0]
mx=max(COMP,key=lambda r:r["delta"]); mn=min(COMP,key=lambda r:r["delta"])

# ---------- 3. ASSERT (single source of truth) ---------------------------
assert N==48 and NC+NP==N and NH+NM==NC
assert NU+ND==NM
assert all(grade(r["retest"]) is not None for r in R), "value outside all grade bands"
SELF={"N":N,"NC":NC,"NH":NH,"NM":NM,"NU":NU,"ND":ND,"NP":NP,
      "mean_d":round(mean_d,2),"med_abs":round(med_abs,2),"gt3":len(gt3),
      "max":(mx["batch"],mx["delta"]),"min":(mn["batch"],mn["delta"])}
print("SELF-CHECK OK", SELF)

# ---------- 4. BUILD ------------------------------------------------------
CODE="PP-QC-THC-VER-001"; VER="1.0"
T_MK="Споредба на вкупен Δ⁹-ТХЦ: реанализа наспроти претходен резултат за пуштање во промет и класификација по спецификациски опсег"
T_EN="Total Δ9-THC Comparison: Re-test versus Previous Batch-Release Result and Classification against the Grade Specification Range"

d=Document(SK+"/assets/PP_BASE_TEMPLATE.docx")
pr.swap_header(d,"ИЗВЕШТАЈ — ", CODE, "Report — "+T_EN, CODE+"/", "2026", VER)
pr.wipe_body(d)

pr.cover_page(d, T_MK, T_EN,
  info_rows=[("Код на документ | Document code", CODE+"/2026"),
             ("Верзија | Version", VER),
             ("Тип | Type", "Компаративен извештај | Comparative report"),
             ("Опфат | Scope", "48 производствени серии (Транша 1 и 2) | 48 production batches (Tranche 1 and 2)"),
             ("Аналит | Analyte", "Вкупен Δ⁹-ТХЦ, % w/w | Total Δ9-THC, % w/w"),
             ("Извор на податоци | Data source", PROV["file"]+"  (SHA-256 "+PROV["sha256"][:16]+"…)")],
  kind_mk="КОМПАРАТИВЕН ИЗВЕШТАЈ НА КОНТРОЛА НА КВАЛИТЕТ", kind_en="Quality Control Comparative Report",
  study_mk="Реанализа 07.08.2026 (Транша 1) · 25–26.08.2026 (Транша 2)",
  study_en="Re-test 07.08.2026 (Tranche 1) · 25–26.08.2026 (Tranche 2)")
pr.toc_page(d)

# --- 1 Executive summary
pr.chapter(d,"1","РЕЗИМЕ","Executive Summary")
pr.body(d,
 f"Од {mk(N,0)} оценети серии, {mk(NC,0)} имаат претходен резултат и се компарабилни; {mk(NP,0)} немаат претходно испитување. "
 f"По реанализата, {mk(NH,0)} серии ({mk(100*NH/NC,1)} %) остануваат во истиот спецификациски опсег утврден со претходниот резултат, "
 f"а {mk(NM,0)} ({mk(100*NM/NC,1)} %) преминуваат во друг опсег — {mk(NU,0)} нагоре и {mk(ND,0)} надолу. "
 f"Средната промена е {'+' if mean_d>0 else ''}{mk(mean_d)} п.п., медијаната на апсолутната промена {mk(med_abs)} п.п.; "
 f"{mk(len(gt3),0)} серии се менуваат за повеќе од 3,0 п.п.",
 f"Of {N} batches assessed, {NC} carry a previous result and are comparable; {NP} have no prior test. "
 f"After re-test, {NH} batches ({100*NH/NC:.1f} %) remain within the same specification range established by the previous result, "
 f"and {NM} ({100*NM/NC:.1f} %) migrate to a different range — {NU} upward and {ND} downward. "
 f"Mean change is {'+' if mean_d>0 else ''}{mean_d:.2f} pp, median absolute change {med_abs:.2f} pp; "
 f"{len(gt3)} batches move by more than 3.0 pp.")
pr.gap(d,4)
pr.databox(d,[("Секоја серија мора да се прекласифицира според реанализата пред етикетирање и ослободување.",
               "Every batch must be re-classified on the re-test value before labelling and release.",11,True)],
           pr.LBL, pr.NAVY)

# --- 2 Specification basis
pr.chapter(d,"2","СПЕЦИФИКАЦИСКА ОСНОВА","Specification Basis")
pr.body(d,
 "Декларираната содржина се изразува преку номинален степен G ∈ {8, 10, …, 28} % w/w, со дозволено отстапување од ±10 % релативно околу номиналната вредност:",
 "Declared content is expressed as a nominal grade G ∈ {8, 10, …, 28} % w/w with a permitted deviation of ±10 % relative about the nominal value:")
if pr._EQN_OK:
    pr.eqn(d, r"\left[\,0.90\,G,\;1.10\,G\,\right],\qquad G\in\{8,10,\dots,28\}\ \%\,w/w")
else:   # MML2OMML.XSL (Word) unavailable on this host -> typeset the formula in Unicode
    from docx.enum.text import WD_ALIGN_PARAGRAPH as _AL
    _p=d.add_paragraph(); _p.alignment=_AL.CENTER; pr.sp(_p,8,8)
    pr.rin(_p,"[ 0,90\u00b7G ,  1,10\u00b7G ]",13,pr.NAVY,bold=True)
    pr.rin(_p,"        G \u2208 {8, 10, \u2026, 28} % w/w",11,pr.GREY,ital=True)
pr.body(d,
 "Серијата се доделува на оној степен чиј опсег ја содржи резултатот; при преклопување на соседни опсези се зема најблискиот номинален степен. Ова е единствениот критериум применет во овој извештај — тој е спецификацијата за декларација, не фармакопејска граница за пуштање во промет.",
 "A batch is assigned to the grade whose range contains the result; where adjacent ranges overlap, the nearest nominal grade is taken. This is the sole criterion applied here — it is the label-claim specification, not a pharmacopoeial release limit.")
pr.gap(d,4)
tb=d.add_table(rows=len(GRADES)+1,cols=4)
pr.cellfmt(tb.cell(0,0),"Степен | Grade",None,9,pr.WHITE,bold=True,fill=pr.NAVYF)
pr.cellfmt(tb.cell(0,1),"Опсег, % w/w | Range",None,9,pr.WHITE,bold=True,fill=pr.NAVYF)
pr.cellfmt(tb.cell(0,2),"Претходно | Previous, n",None,9,pr.WHITE,bold=True,fill=pr.NAVYF)
pr.cellfmt(tb.cell(0,3),"Реанализа | Re-test, n",None,9,pr.WHITE,bold=True,fill=pr.NAVYF)
for i,G in enumerate(GRADES,1):
    np_=sum(1 for r in R if r["g_prev"]==G); nr_=sum(1 for r in R if r["g_ret"]==G)
    pr.cellfmt(tb.cell(i,0),"G%d"%G,None,9,pr.NAVY,bold=True,fill=pr.LBL)
    pr.cellfmt(tb.cell(i,1),bandstr_mk(G)+" | "+bandstr(G),None,9)
    pr.cellfmt(tb.cell(i,2),str(np_),None,9)
    pr.cellfmt(tb.cell(i,3),str(nr_),None,9)
pr.fixed(tb); pr.borders(tb)

# --- 3 Outcome
pr.chapter(d,"3","ИСХОД ПО СЕРИЈА","Outcome by Batch")
pr.subsec(d,"3.1","Збирен исход","Aggregate outcome")
tb=d.add_table(rows=6,cols=3)
pr.cellfmt(tb.cell(0,0),"Исход | Outcome",None,9,pr.WHITE,bold=True,fill=pr.NAVYF)
pr.cellfmt(tb.cell(0,1),"n",None,9,pr.WHITE,bold=True,fill=pr.NAVYF)
pr.cellfmt(tb.cell(0,2),"% од компарабилни | of comparable",None,9,pr.WHITE,bold=True,fill=pr.NAVYF)
rowsA=[("Ист опсег задржан | Same range retained",NH),
       ("Премин во повисок опсег | Migrated upward",NU),
       ("Премин во понизок опсег | Migrated downward",ND),
       ("Вкупно компарабилни | Total comparable",NC),
       ("Без претходен резултат | No prior result",NP)]
for i,(lab,n) in enumerate(rowsA,1):
    bold = lab.startswith("Вкупно")
    pr.cellfmt(tb.cell(i,0),lab,None,9,pr.NAVY if bold else pr.BLACK,bold=bold,fill=pr.LBL if bold else None)
    pr.cellfmt(tb.cell(i,1),str(n),None,9,bold=bold)
    pr.cellfmt(tb.cell(i,2), "—" if lab.startswith("Без") else mk(100*n/NC,1)+" %", None,9,bold=bold)
pr.fixed(tb); pr.borders(tb)
pr.gap(d,6)
pr.note(d,
 f"Најголем пораст: {mx['batch']} ({mx['strain']}) {'+' if mx['delta']>0 else ''}{mk(mx['delta'])} п.п. Најголем пад: {mn['batch']} ({mn['strain']}) {mk(mn['delta'])} п.п.",
 f"Largest increase: {mx['batch']} ({mx['strain']}) {mx['delta']:+.2f} pp. Largest decrease: {mn['batch']} ({mn['strain']}) {mn['delta']:+.2f} pp.")

pr.subsec(d,"3.2","Серии кои го менуваат спецификацискиот опсег","Batches changing specification range")
tb=d.add_table(rows=len(MOVED)+1,cols=6)
H=["Серија | Batch","Сорта | Strain","Претх. | Prev, %","Реанализа | Re-test, %","Δ, п.п. | pp","Опсег: од → во | Range: from → to"]
for j,h in enumerate(H): pr.cellfmt(tb.cell(0,j),h,None,9,pr.WHITE,bold=True,fill=pr.NAVYF)
for i,r in enumerate(sorted(MOVED,key=lambda r:-abs(r["delta"])),1):
    pr.cellfmt(tb.cell(i,0),r["batch"],None,9,pr.NAVY,bold=True,fill=pr.LBL)
    pr.cellfmt(tb.cell(i,1),r["strain"],None,9)
    pr.cellfmt(tb.cell(i,2),mk(r["prev"]),None,9)
    pr.cellfmt(tb.cell(i,3),mk(r["retest"]),None,9)
    pr.cellfmt(tb.cell(i,4),("+" if r["delta"]>0 else "")+mk(r["delta"]),None,9,
               pr.GREEN if r["delta"]>0 else pr.RED, bold=True)
    pr.cellfmt(tb.cell(i,5),"G%d → G%d"%(r["g_prev"],r["g_ret"]),None,9,bold=True)
pr.fixed(tb); pr.borders(tb)

pr.subsec(d,"3.3","Серии кои го задржуваат опсегот","Batches retaining their range")
tb=d.add_table(rows=len(HELD)+1,cols=5)
H=["Серија | Batch","Сорта | Strain","Претх. | Prev, %","Реанализа | Re-test, %","Опсег | Range"]
for j,h in enumerate(H): pr.cellfmt(tb.cell(0,j),h,None,9,pr.WHITE,bold=True,fill=pr.NAVYF)
for i,r in enumerate(sorted(HELD,key=lambda r:r["batch"]),1):
    pr.cellfmt(tb.cell(i,0),r["batch"],None,9,pr.NAVY,bold=True,fill=pr.LBL)
    pr.cellfmt(tb.cell(i,1),r["strain"],None,9)
    pr.cellfmt(tb.cell(i,2),mk(r["prev"]),None,9)
    pr.cellfmt(tb.cell(i,3),mk(r["retest"]),None,9)
    pr.cellfmt(tb.cell(i,4),"G%d (%s)"%(r["g_prev"],bandstr(r["g_prev"])),None,9)
pr.fixed(tb); pr.borders(tb)

pr.subsec(d,"3.4","Серии без претходен резултат","Batches without a prior result")
tb=d.add_table(rows=len(NOPRIOR)+1,cols=4)
H=["Серија | Batch","Сорта | Strain","Реанализа | Re-test, %","Доделен опсег | Assigned range"]
for j,h in enumerate(H): pr.cellfmt(tb.cell(0,j),h,None,9,pr.WHITE,bold=True,fill=pr.NAVYF)
for i,r in enumerate(sorted(NOPRIOR,key=lambda r:r["batch"]),1):
    pr.cellfmt(tb.cell(i,0),r["batch"],None,9,pr.NAVY,bold=True,fill=pr.LBL)
    pr.cellfmt(tb.cell(i,1),r["strain"],None,9)
    pr.cellfmt(tb.cell(i,2),mk(r["retest"]),None,9)
    pr.cellfmt(tb.cell(i,3),"G%d (%s)"%(r["g_ret"],bandstr(r["g_ret"])),None,9)
pr.fixed(tb); pr.borders(tb)

# --- 4 Figures
pr.chapter(d,"4","ВИЗУЕЛНИ ПРЕТСТАВИ","Visual Representations")
pr.figure(d, figs.fig_scatter("f1.png"),
  "Реанализа наспроти претходен резултат; сенчените појаси се спецификациските опсези, испрекинатата линија е идентитет.",
  "Re-test versus previous result; shaded bands are the specification ranges, dashed line is identity.")
pr.figure(d, figs.fig_migration("f2.png"),
  "Распределба на серии по степен, претходно наспроти реанализа.",
  "Distribution of batches by grade, previous versus re-test.")

# --- 5 Conclusion
pr.chapter(d,"5","ЗАКЛУЧОК","Conclusion")
pr.body(d,
 f"Реанализата не ја потврдува претходната класификација за {mk(NM,0)} од {mk(NC,0)} компарабилни серии — половина од оценетиот сет. "
 f"Отстапувањето е претежно нагорно ({mk(NU,0)} наспроти {mk(ND,0)}), што исклучува систематска загуба на потентност при чување како единствено објаснување и укажува на разлика меѓу аналитичките серии. "
 "Декларацијата на секоја серија мора да се изведе од реанализата; серијата не смее да се етикетира според претходниот степен.",
 f"The re-test does not confirm the previous classification for {NM} of {NC} comparable batches — half the assessed set. "
 f"The shift is predominantly upward ({NU} versus {ND}), which excludes systematic potency loss on storage as the sole explanation and points to a between-run analytical difference. "
 "Each batch label claim must be derived from the re-test; a batch must not be labelled on its previous grade.")
pr.gap(d,4)
pr.body(d,
 f"Ограничување: {mk(NP,0)} серии немаат претходен резултат и не влегуваат во компарацијата; вредностите од Транша 2 (25–26.08.2026) се работни изводи без издаден сертификат за анализа и не се валидни како резултат за пуштање во промет додека не бидат сертифицирани.",
 f"Limitation: {NP} batches have no prior result and are excluded from the comparison; the Tranche 2 values (25–26.08.2026) are working extracts with no certificate of analysis issued and are not valid as release results until certificated.")
pr.gap(d,6)
pr.note(d,
 "Потекло на податоци | "+PROV["file"]+" · SHA-256 "+PROV["sha256"]+" · "+str(PROV["bytes"])+" B.",
 "Data provenance as stated. All statistics computed from the bound dataset at build time.")

pr.execution_signoff(d)
OUT="%s_Total_THC_Retest_vs_Previous_v%s.docx"%(CODE,VER)
d.save(OUT)
print("SAVED",OUT)
