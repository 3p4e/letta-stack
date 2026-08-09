# -*- coding: utf-8 -*-
# Purely Plant — REPORT / COMPUTATIONAL-RECORD engine (part of pp-document-suite).
# Bilingual MK|EN house style + native Word equations (OMML), worked calcs, forms, figures.
# House rules: body text JUSTIFIED; table text CENTERED (h+v); every table FITS the page;
# per-step two-role sign-off (Operator + QC Department Manager) for protocol record steps.
# Canonical tokens (pp_theme.py): one house navy #2B547E; 6 pt font floor.
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x2B,0x54,0x7E); GREY=RGBColor(0x59,0x59,0x59); BLACK=RGBColor(0,0,0)
WHITE=RGBColor(0xFF,0xFF,0xFF); GREEN=RGBColor(0x37,0x56,0x23); RED=RGBColor(0xC0,0,0); AMBER=RGBColor(0xC5,0x5A,0x11)
LBL="F2F2F2"; BORDER="2B547E"; NAVYF="2B547E"; GREENF="E2EFDA"; REDF="FCE4D6"; AMBERF="FFF2CC"
# fixed() (the single table-layout brain) and its exclusive private helpers now live in pp_format.py
# — imported here rather than redefined, so both engines share ONE implementation (no drift).
from pp_format import fixed, PAGE_W, _tctext, _gridspan, _settcw_dxa, _repeat_header, _ENTRY_TARGETS

def shade(cell,hexf):
    tcPr=cell._tc.get_or_add_tcPr()
    sh=OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:color'),'auto'); sh.set(qn('w:fill'),hexf); tcPr.append(sh)
def borders(tbl):
    tblPr=tbl._tbl.tblPr; b=OxmlElement('w:tblBorders')
    for e in ('top','left','bottom','right','insideH','insideV'):
        el=OxmlElement('w:'+e); el.set(qn('w:val'),'single'); el.set(qn('w:sz'),'4'); el.set(qn('w:space'),'0'); el.set(qn('w:color'),BORDER); b.append(el)
    tblPr.append(b)
def _settcw(tc, cm):
    tcPr=tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcW')): tcPr.remove(old)
    e=OxmlElement('w:tcW'); e.set(qn('w:type'),'pct'); e.set(qn('w:w'),str(max(1,int(cm/PAGE_W*5000)))); tcPr.append(e)

def rin(p,t,sz,col,bold=False,ital=False):
    r=p.add_run(t); f=r.font; f.name='Calibri'; f.size=Pt(sz); f.color.rgb=col; r.bold=bold; r.italic=ital; return r
def sp(p,bef=0,aft=0,line=None):
    pf=p.paragraph_format; pf.space_before=Pt(bef); pf.space_after=Pt(aft)
    if line: pf.line_spacing=line
def gap(d,pt=6):
    """A thin spacer paragraph — use ONLY at boundaries (between logical chunks, around tables/
       figures). Body paragraphs of the same style carry NO inter-paragraph space, so separation
       is added deliberately here, not after every line."""
    p=d.add_paragraph(); pf=p.paragraph_format
    pf.space_before=Pt(0); pf.space_after=Pt(0); pf.line_spacing=Pt(pt)   # exact thin height, no sub-6pt font run
    return p

# ---- house type scale ----
def chapter(d,num,mk,en):                       # MK20 | EN16 ; space BEFORE and AFTER ; outline lvl 0 -> TOC
    p=d.add_paragraph(); sp(p,12,7)
    rin(p,f"{num}.  {mk}",20,NAVY,bold=True); rin(p,"  |  ",18,GREY); rin(p,en,16,GREY)
    _outline(p,0); return p
def subsec(d,num,mk,en):                         # MK16 | EN12 ; space BEFORE and AFTER
    p=d.add_paragraph(); sp(p,8,4)
    rin(p,f"{num}  {mk}",16,NAVY,bold=True)
    if en: rin(p,"  |  ",14,GREY); rin(p,en,12,GREY)
    _outline(p,1); return p
def minilabel(d,mk,en):                          # unnumbered bold lead-in above a table
    p=d.add_paragraph(); sp(p,7,3)
    rin(p,mk,13,NAVY,bold=True)
    if en: rin(p,"  |  ",11,GREY); rin(p,en,11,GREY,ital=True)
    return p
def note(d,mk,en):                               # small note MK10 | EN8 ; justified
    p=d.add_paragraph(); sp(p,4.5,4.5); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    if mk: rin(p,mk,10,GREY,ital=True)
    if en:
        if mk: rin(p,"  |  ",8,GREY)
        rin(p,en,8,GREY,ital=True)
    return p
def body(d,mk,en):                               # prose MK11 | EN8 ; justified ; NO inter-paragraph space (same style)
    p=d.add_paragraph(); sp(p,0,0,line=1.04); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    if mk: rin(p,mk,11,BLACK)
    if en:
        if mk: rin(p,"  |  ",8,GREY)
        rin(p,en,8,GREY,ital=True)
    return p
def bullet(d,mk,en):                             # bullet MK11 | EN8 ; justified ; tight within the list
    p=d.add_paragraph(); sp(p,0,1); p.paragraph_format.left_indent=Cm(0.5); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    rin(p,"•  ",11,NAVY,bold=True); rin(p,mk,11,BLACK)
    if en: rin(p,"  |  ",8,GREY); rin(p,en,8,GREY,ital=True)
    return p

def titleblock(d,main,subtitle,annex,main_sz=18):
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; sp(p,2,2); rin(p,main,main_sz,NAVY,bold=True)
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; sp(p,0,1); rin(p,subtitle,12,GREY)
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; sp(p,0,8); rin(p,annex,11,GREY,ital=True)

def cellfmt(cell,mk,en,sz=10,col=BLACK,bold=False,enital=True,fill=None):   # table cells CENTERED (h+v); MK sz | EN sz-2 (>=7)
    if fill: shade(cell,fill)
    try: cell.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
    except Exception: pass
    p=cell.paragraphs[0]; p.text=''; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; sp(p,1,1)
    ensz=max(sz-2,7)
    if en is None and isinstance(mk,str) and " | " in mk:
        mk,en = mk.split(" | ",1)     # auto-split combined "MK | EN" so EN renders smaller
    if mk: rin(p,mk,sz,col,bold=bold)
    if en:
        if mk: rin(p," | ",ensz,col)
        rin(p,en,ensz,col,bold=bold and enital==False,ital=enital and not bold)
    return cell

def eqn_cell(cell, latex, fill=None, before=1, after=1):
    """Render a PURE formula inside a table cell as a NATIVE Word equation (OMML).
       Use on light/white cells only (OMML runs are black — unreadable on a navy fill)."""
    if fill: shade(cell, fill)
    try: cell.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
    except Exception: pass
    cell.text=''
    p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; sp(p,before,after)
    eqn_in(p, latex)
    return cell

def hmerge(tbl, row_idx, c0, c1, mk=None, en=None, sz=9, col=NAVY, bold=True, fill=LBL):
    """Merge cells [c0..c1] of a row into ONE centered cell with a single (optional) label.
       Use for summary/footer rows (e.g. a 'Mean …' row) so a spanning label is written ONCE,
       never repeated across columns. Pair with a single value cell in the remaining column(s)."""
    a=tbl.cell(row_idx,c0); b=tbl.cell(row_idx,c1); m=a.merge(b)
    for extra in m.paragraphs[1:]:        # collapse paragraphs left over from the merged cells
        extra._p.getparent().remove(extra._p)
    if mk is not None or en is not None:
        cellfmt(m, mk, en, sz, col, bold=bold, fill=fill)
    else:
        if fill: shade(m, fill)
        m.paragraphs[0].text=''
    return m

def _apply_widths(tbl, widths):
    """Apply EXACT column widths (cm) with AutoFit-to-Window + fixed layout (no content estimate)."""
    grid=tbl._tbl.tblGrid.gridCol_lst
    tbl.autofit=False; tbl.allow_autofit=False
    tblPr=tbl._tbl.tblPr
    for tag in ('w:tblLayout','w:tblW','w:tblInd'):
        for old in tblPr.findall(qn(tag)): tblPr.remove(old)
    lay=OxmlElement('w:tblLayout'); lay.set(qn('w:type'),'fixed'); tblPr.append(lay)
    tw=OxmlElement('w:tblW'); tw.set(qn('w:type'),'pct'); tw.set(qn('w:w'),'5000'); tblPr.append(tw)
    ti=OxmlElement('w:tblInd'); ti.set(qn('w:type'),'dxa'); ti.set(qn('w:w'),'0'); tblPr.append(ti)
    for j,gc in enumerate(grid):
        if j<len(widths): gc.set(qn('w:w'), str(max(1,int(widths[j]*567))))
    for tr in tbl._tbl.tr_lst:
        for j,tc in enumerate(tr.tc_lst):
            if j<len(widths): _settcw(tc, widths[j])
    borders(tbl); return tbl

def _noborder(cell):
    tcPr=cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcBorders')): tcPr.remove(old)
    b=OxmlElement('w:tcBorders')
    for e in ('top','left','bottom','right'):
        el=OxmlElement('w:'+e); el.set(qn('w:val'),'nil'); b.append(el)
    tcPr.append(b)

def status_grid(d, options, selected, ncols=2, sz=12, cb=1.1, gap=0.5):
    """Compact checkbox / option grid for a short mutually-exclusive list (e.g. the six method-status
       tiers). Lays options out column-major into `ncols` side-by-side blocks — each block a narrow
       checkbox column + a wide label column — separated by a thin borderless gap, so the list occupies
       minimal vertical space yet stays presentable. The selected option is ticked (☒) and highlighted."""
    n=len(options); nrows=(n+ncols-1)//ncols; tot=3*ncols-1
    t=d.add_table(rows=nrows,cols=tot); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    lblw=(PAGE_W-ncols*cb-(ncols-1)*gap)/ncols
    widths=[]
    for b in range(ncols):
        widths+=[cb,lblw]
        if b<ncols-1: widths+=[gap]
    spacers=[b*3+2 for b in range(ncols-1)]
    for b in range(ncols):
        cbc=b*3; lbc=b*3+1
        for r in range(nrows):
            idx=b*nrows+r
            if idx<n:
                opt=options[idx]; sel=(opt==selected)
                cellfmt(t.cell(r,cbc),("☒" if sel else "☐"),None,sz,(GREEN if sel else BLACK),bold=sel)
                cellfmt(t.cell(r,lbc),opt,None,sz,(GREEN if sel else BLACK),bold=sel)
            else:
                cellfmt(t.cell(r,cbc),"",None,sz,BLACK); cellfmt(t.cell(r,lbc),"",None,sz,BLACK)
    for r in range(nrows):
        for sc in spacers: cellfmt(t.cell(r,sc),"",None,sz,BLACK)
    _apply_widths(t,widths)
    for r in range(nrows):
        for sc in spacers: _noborder(t.cell(r,sc))
    return t

def databox(d,lines,fill,tcol):    # caveat/status box; lines=(mk,en,sz,bold)
    t=d.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    cc=t.cell(0,0); shade(cc,fill); first=True
    for (mk,en,sz,bold) in lines:
        p=cc.paragraphs[0] if first else cc.add_paragraph(); first=False
        p.alignment=WD_ALIGN_PARAGRAPH.CENTER; sp(p,2,2)
        if mk: rin(p,mk,sz,tcol,bold=bold)
        if en:
            if sz>=13:
                p2=cc.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER; sp(p2,0,2); rin(p2,en,max(sz-2,11),tcol,ital=True)
            else:
                rin(p,"  |  ",10,tcol); rin(p,en,10,tcol,ital=True)
    fixed(t,[PAGE_W]); borders(t); return t

def swap_header(d,title_mk_prefix,title_code,title_en,code_prefix,code_suffix_a,code_suffix_b):
    h=d.sections[0].header.tables[0]
    c1=h.cell(0,1).paragraphs[1].runs
    c1[0].text=title_mk_prefix; c1[2].text=title_code; c1[5].text=title_en
    c2=h.cell(0,2).paragraphs[2].runs
    c2[0].text=code_prefix; c2[2].text=code_suffix_a; c2[3].text=code_suffix_b

def wipe_body(d):
    bodyel=d.element.body; sectPr=bodyel.find(qn('w:sectPr'))
    for ch in list(bodyel):
        if ch is not sectPr: bodyel.remove(ch)
    return bodyel,sectPr

# =====================================================================
# Native Word math via LaTeX -> MathML (latex2mathml) -> OMML (Office MML2OMML.XSL).
# =====================================================================
import os as _os
_EQN_OK=False; _OMML_XSLT=None; _l2m=None
try:
    import latex2mathml.converter as _l2m
    from lxml import etree as _etree
    _XSL_CANDIDATES=[
        _os.environ.get("PP_MML2OMML_XSL",""),
        r"C:\Program Files (x86)\Microsoft Office\root\Office16\MML2OMML.XSL",
        r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL",
        r"C:\Program Files (x86)\Microsoft Office\Office16\MML2OMML.XSL",
        r"C:\Program Files\Microsoft Office\Office16\MML2OMML.XSL",
    ]
    _XSLP=next((p for p in _XSL_CANDIDATES if p and _os.path.exists(p)),None)
    if _XSLP:
        _OMML_XSLT=_etree.XSLT(_etree.parse(_XSLP)); _EQN_OK=True
except Exception:
    _EQN_OK=False

def _to_omml(latex):
    mml=_l2m.convert(latex)
    dom=_etree.fromstring(mml.encode("utf-8") if isinstance(mml,str) else mml)
    return _OMML_XSLT(dom).getroot()

def eqn(d, latex, center=True, before=3, after=3):
    """Display equation as native, editable Word math (OMML). Falls back to text if engine unavailable."""
    p=d.add_paragraph(); sp(p,before,after)
    if center: p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    if _EQN_OK:
        try: p._p.append(_to_omml(latex)); return p
        except Exception: pass
    rin(p,latex,11,BLACK); return p

_M_NS="http://schemas.openxmlformats.org/officeDocument/2006/math"
def _omml_fmt(root, color_hex=None, bold=False):
    """Apply run colour/bold to every math run (m:r) of an OMML tree — lets native equations be
       WHITE on a navy header cell, or NAVY+bold for an emphasised result."""
    for r in list(root.iter('{%s}r'%_M_NS)):
        mrpr=r.find('{%s}rPr'%_M_NS)
        wrpr=OxmlElement('w:rPr')
        if color_hex is not None:
            c=OxmlElement('w:color'); c.set(qn('w:val'),color_hex); wrpr.append(c)
        if bold:
            b=OxmlElement('w:b'); wrpr.append(b)
        r.insert((list(r).index(mrpr)+1) if mrpr is not None else 0, wrpr)
    return root

def eqn_in(p, latex, white=False, color_hex=None, bold=False):
    """Insert an inline native equation into an existing paragraph. white=True (or color_hex/bold)
       formats the math (e.g. white on a navy header)."""
    if _EQN_OK:
        try:
            root=_to_omml(latex)
            if white: color_hex="FFFFFF"
            if color_hex is not None or bold: _omml_fmt(root,color_hex,bold)
            p._p.append(root); return p
        except Exception: pass
    rin(p,latex,11,(WHITE if white else BLACK)); return p

def mathcell(cell, latex, fill=NAVYF, white=True, before=1, after=1):
    """Render a NATIVE equation as a table HEADER (or any) cell. On a navy fill the math is white
       and therefore legible — use for headers that are mathematical expressions (xᵢ, xᵢ−x̄, (xᵢ−x̄)²,
       u_c, RSD, U …). For a light cell pass fill=None, white=False (math renders black)."""
    if fill: shade(cell, fill)
    try: cell.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
    except Exception: pass
    cell.text=''
    p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; sp(p,before,after)
    eqn_in(p, latex, white=white)
    return cell

def eqn_result(d, latex, mk=None, en=None):
    """An EMPHASISED result/conclusion equation: a navy left-rule + light fill, the math in NAVY bold,
       optionally led by a short bilingual label. Use to mark the conclusion of a logical calc block."""
    p=d.add_paragraph(); sp(p,4,4); p.paragraph_format.left_indent=Cm(0.3)
    pPr=p._p.get_or_add_pPr()
    pbdr=OxmlElement('w:pBdr'); lb=OxmlElement('w:left')
    lb.set(qn('w:val'),'single'); lb.set(qn('w:sz'),'18'); lb.set(qn('w:space'),'8'); lb.set(qn('w:color'),'2B547E')
    pbdr.append(lb); pPr.append(pbdr)
    shd=OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:fill'),'EAF1F8'); pPr.append(shd)
    if mk or en:
        rin(p,(mk or ""),11,NAVY,bold=True)
        if en: rin(p,"  |  ",9,GREY); rin(p,en,9,GREY,ital=True)
        rin(p,"   ",11,NAVY)
    eqn_in(p, latex, color_hex="2B547E", bold=True)
    return p

def calc_step(d, mk=None, en=None, formula=None, substituted=None, result=None, tag=None):
    """Worked-calculation step: bilingual label → general formula (OMML) → substituted (OMML) →
       EMPHASISED concluding result as a NATIVE equation (navy left-rule + light fill, navy bold).
       `result` is LaTeX; `tag` is optional trailing text (e.g. a ✓ / verdict)."""
    if mk or en:
        p=d.add_paragraph(); sp(p,5,1); rin(p,(mk or ""),12,NAVY,bold=True)
        if en: rin(p,"  |  ",11,GREY); rin(p,en,11,GREY,ital=True)
    if formula: eqn(d, formula, before=2, after=2)
    if substituted: eqn(d, substituted, before=2, after=2)
    if result:
        p=d.add_paragraph(); sp(p,2,5); p.paragraph_format.left_indent=Cm(0.3)   # concluding result = EMPHASISED native eqn
        pPr=p._p.get_or_add_pPr()
        pbdr=OxmlElement('w:pBdr'); lb=OxmlElement('w:left')
        lb.set(qn('w:val'),'single'); lb.set(qn('w:sz'),'18'); lb.set(qn('w:space'),'8'); lb.set(qn('w:color'),'2B547E')
        pbdr.append(lb); pPr.append(pbdr)
        shd=OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:fill'),'EAF1F8'); pPr.append(shd)
        eqn_in(p, result, color_hex="2B547E", bold=True)
        if tag: rin(p,"   "+tag,11.5,NAVY,bold=True)

def step_signoff(d, mk=None, en=None, sz=9):
    """Per-step TWO-ROLE sign-off for protocol record steps that occur at different dates/times:
       (1) Operator/Analyst who executed the step and entered the raw data;
       (2) QC Department Manager who checked and approved this record-making part.
       Place ONE after every logical record-making step (table) in a validation/verification protocol."""
    if mk or en:
        p=d.add_paragraph(); sp(p,3,1)
        rin(p,(mk or "Потпис за овој запис"),9,NAVY,bold=True,ital=True)
        if en: rin(p,"  |  ",8,GREY); rin(p,en,8,GREY,ital=True)
    h=["Улога | Role","Име | Name","Датум | Date","Потпис | Signature"]
    roles=["Извршил и внел сурови податоци (оператор/аналитичар) | Executed & entered raw data (Operator/Analyst)",
           "Проверил и одобрил овој запис (Раководител на КК) | Checked & approved this record (QC Department Manager)"]
    t=d.add_table(rows=3,cols=4); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,x in enumerate(h): cellfmt(t.cell(0,j),x,None,sz,WHITE,bold=True,fill=NAVYF)
    for i,role in enumerate(roles,start=1):
        cellfmt(t.cell(i,0),role,None,sz,BLACK,fill=LBL)
        for j in (1,2,3): cellfmt(t.cell(i,j),"",None,sz,BLACK)
    fixed(t,[9.46,3.5,2.0,3.5]); borders(t); return t

def entry_table(d, headers, rows, widths, label_mk=None, label_en=None, sz=9, signoff=False, signoff_mk=None, signoff_en=None):
    """Blank data-ENTRY table. rows = int (all blank) or list of first-column labels.
       signoff=True appends a per-step two-role sign-off (Operator + QC Department Manager)."""
    if label_mk or label_en: minilabel(d, label_mk, label_en)
    rowlabels = rows if isinstance(rows,(list,tuple)) else [None]*int(rows)
    t=d.add_table(rows=len(rowlabels)+1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,h in enumerate(headers): cellfmt(t.cell(0,j),h,None,sz,WHITE,bold=True,fill=NAVYF)
    for i,rl in enumerate(rowlabels,start=1):
        cellfmt(t.cell(i,0),(rl if rl else ""),None,sz,(NAVY if rl else BLACK),bold=bool(rl),fill=(LBL if rl else None))
        for j in range(1,len(headers)): cellfmt(t.cell(i,j),"",None,sz,BLACK)
    fixed(t,widths); borders(t)
    if signoff: step_signoff(d, signoff_mk, signoff_en)
    return t

def execution_signoff(d, mk_exec="Извршил (КК) | Executed (QC)", en_exec=None, reviewer="J. Romevska", approver="B. Nikolov, M.Pharm. (Раководител КК | QC Manager)"):
    """Overall protocol-level sign-off cascade: Executed (analyst) -> Reviewed (QA) -> Approved (QC Manager)."""
    h=["Дејство | Action","Име | Name","Датум | Date","Потпис | Signature"]
    rws=[(mk_exec,"____________________"),("Прегледал (QA) | Reviewed (QA)",reviewer),("Одобрил (Раководител КК) | Approved (QC Manager)",approver)]
    t=d.add_table(rows=4,cols=4); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,x in enumerate(h): cellfmt(t.cell(0,j),x,None,10,WHITE,bold=True,fill=NAVYF)
    for i,(a,n) in enumerate(rws,start=1):
        cellfmt(t.cell(i,0),a,None,10,BLACK); cellfmt(t.cell(i,1),n,None,10,BLACK)
        cellfmt(t.cell(i,2),"",None,10,BLACK); cellfmt(t.cell(i,3),"",None,10,BLACK)
    fixed(t,[5.4,4.66,3.4,5.0]); borders(t); return t

def figure(d, png_path, mk=None, en=None, width_cm=15.5):
    """Embed a chart PNG (centered) with a bilingual MK | EN caption. Pair with pp_charts.py."""
    from docx.shared import Cm as _Cm
    d.add_picture(png_path, width=_Cm(width_cm))
    d.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    if mk or en:
        p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; sp(p,1,8)
        rin(p,"Слика | Figure — ",9,NAVY,bold=True)
        if mk: rin(p,mk,9,GREY,ital=True)
        if en: rin(p,"  |  ",8,GREY); rin(p,en,9,GREY,ital=True)
    return d

# =====================================================================
# Cover page (unnumbered) + Table of Contents (page 2) + TOC outline levels.
# Standard document opening for every report/protocol from this skill:
#   page 1 = cover (no number) ; page 2 = TOC ; page 3+ = content (Executive Summary first).
# =====================================================================
def _outline(p, lvl):
    pPr=p._p.get_or_add_pPr(); o=OxmlElement('w:outlineLvl'); o.set(qn('w:val'),str(lvl)); pPr.append(o)

def _info_table(d, rows):
    t=d.add_table(rows=len(rows),cols=2); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,(k,v) in enumerate(rows):
        cellfmt(t.cell(i,0),k,None,10,NAVY,bold=True,fill=LBL); cellfmt(t.cell(i,1),v,None,10,BLACK)
    fixed(t,[5.6,12.86]); borders(t); return t

def _cover_header_footer(sec):
    """Cover (page 1) keeps the STANDARD running header (logo + doc name + code) like every other page,
       but carries NO page number; numbering resumes on page 2 via the default footer (bottom-right)."""
    sectPr=sec._sectPr
    if sectPr.find(qn('w:titlePg')) is None: sectPr.append(OxmlElement('w:titlePg'))
    defid=None
    for hr in sectPr.findall(qn('w:headerReference')):
        if hr.get(qn('w:type'))=='default': defid=hr.get(qn('r:id'))
    for hr in list(sectPr.findall(qn('w:headerReference'))):
        if hr.get(qn('w:type'))=='first': sectPr.remove(hr)
    if defid:
        h1=OxmlElement('w:headerReference'); h1.set(qn('w:type'),'first'); h1.set(qn('r:id'),defid); sectPr.insert(0,h1)
    for fr in list(sectPr.findall(qn('w:footerReference'))):     # no first-page footer -> cover unnumbered
        if fr.get(qn('w:type'))=='first': sectPr.remove(fr)

def cover_page(d, title_mk, title_en, info_rows, kind_mk="", kind_en="", study_mk=None, study_en=None,
               approval_rows=None, approval_qp="Одобрил (Раководител КК) | Approved (QC Manager)",
               approval_qp_name="B. Nikolov, M.Pharm."):
    """Distinct, UNNUMBERED cover page (different-first-page): wordmark + big bilingual title +
       method INFORMATION block + APPROVAL block. Followed by a page break (TOC goes on page 2).
       Cover title is centered (NOT justified). Use once, first thing, in every report/protocol."""
    _cover_header_footer(d.sections[0])   # cover keeps the standard running header (logo+name+code); page 1 unnumbered, footer resumes p.2
    if kind_mk:
        p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; sp(p,40,2); rin(p,kind_mk,13,GREY,bold=True)
        if kind_en: rin(p,"  |  ",11,GREY); rin(p,kind_en,11,GREY)
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; sp(p,10,2); rin(p,title_mk,22,NAVY,bold=True)
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; sp(p,0,2); rin(p,title_en,13,NAVY)
    if study_mk:
        p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; sp(p,4,16); rin(p,study_mk,12,GREY,bold=True)
        if study_en: rin(p,"  |  ",10,GREY); rin(p,study_en,10,GREY,ital=True)
    else:
        p=d.add_paragraph(); sp(p,0,12)
    minilabel(d,"Информации за документот | Document information",None)
    _info_table(d, info_rows)
    minilabel(d,"Одобрување | Approval",None)
    rows=approval_rows or [("Изготвил | Prepared (QC Head)","B. Nikolov, M.Pharm."),
                           ("Прегледал | Reviewed (QA)","J. Romevska"),
                           (approval_qp, approval_qp_name)]
    h=["Дејство | Action","Име/Позиција | Name/Position","Датум | Date","Потпис | Signature"]
    t=d.add_table(rows=len(rows)+1,cols=4); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,x in enumerate(h): cellfmt(t.cell(0,j),x,None,10,WHITE,bold=True,fill=NAVYF)
    for i,(a,n) in enumerate(rows,start=1):
        cellfmt(t.cell(i,0),a,None,10,BLACK); cellfmt(t.cell(i,1),n,None,10,BLACK)
        cellfmt(t.cell(i,2),"",None,10,BLACK); cellfmt(t.cell(i,3),"",None,10,BLACK)
    fixed(t,[5.0,5.46,3.5,4.5]); borders(t)
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; sp(p,10,0); rin(p,"Контролиран документ | Controlled document",8,GREY,ital=True)
    d.add_page_break()

def _toc_field(d, placeholder="Десен-клик → Ажурирај поле (Update Field) | Right-click → Update Field"):
    p=d.add_paragraph()
    r=p.add_run()
    b=OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'),'begin'); r._r.append(b)
    it=OxmlElement('w:instrText'); it.set(qn('xml:space'),'preserve'); it.text=' TOC \\o "1-3" \\h \\z \\u '; r._r.append(it)
    s=OxmlElement('w:fldChar'); s.set(qn('w:fldCharType'),'separate'); r._r.append(s)
    r2=p.add_run(placeholder); r2.font.size=Pt(9); r2.font.italic=True; r2.font.color.rgb=GREY
    r3=p.add_run(); e=OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'),'end'); r3._r.append(e)
    return p

def toc_page(d, mk="СОДРЖИНА", en="Table of Contents"):
    """Table of contents on its own page (page 2). Native Word TOC field — user updates it on open
       (Right-click -> Update Field). Picks up chapter()/subsec() via their outline levels."""
    p=d.add_paragraph(); sp(p,6,8); rin(p,mk,16,NAVY,bold=True); rin(p,"  |  ",12,GREY); rin(p,en,12,GREY)
    _toc_field(d)
    d.add_page_break()
