"""
pp_format.py — Purely Plant Document Suite formatting engine (python-docx).

Builds bilingual Macedonian|English controlled .docx in the Purely Plant house style:
  - SOP   : two-column (MK left | EN right) with vertical divider + Word native TOC
  - Annex : full-page bilingual inline (MK 11pt | EN 7pt) with the PP v2 colour palette

Render to PDF with Microsoft Word COM (Windows, see render_pdf.ps1) or:
    soffice --headless --convert-to pdf <file>.docx --outdir .

Requires: python-docx  (pip install python-docx).  Run with PYTHONUTF8=1 on Windows.
NOTE: this is the shared engine — smoke-test on a one-page sample before first production run.

MANDATORY PP HEADER/FOOTER (required on EVERY Purely Plant document — SOP, annex, all):
the running header (leaf logo + wordmark | bilingual Document name | Code of document + Version)
and the "Page X of Y" footer live in the base template  assets/PP_BASE_TEMPLATE.docx . They are
loaded automatically by new_sop()/new_annex() (from_template=True, the default), the title/code/
version are stamped by apply_pp_header(), the body is cleared by wipe_body(), and save() re-seats
the section <w:sectPr> last so the header/footer/page-geometry apply document-wide. See formatting_specs.md.
"""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- palette / fonts ----------
NAVY  = RGBColor(0x2B, 0x54, 0x7E)   # annex section header / SOP accents
GREY  = RGBColor(0x59, 0x59, 0x59)
BLACK = RGBColor(0, 0, 0)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
SOP_GRAY = "E8E8E8"      # SOP section-header fill
SOP_BLUE = "D9E2F3"      # SOP approval-table header
A_SECTION = "2B547E"     # annex section header (white text)
A_LABEL   = "EDF2F7"     # annex label cell
A_ALT     = "F7FAFC"     # annex alternating row
A_WARM    = "FEF9E7"; A_ROSE = "FDEDEC"; A_MINT = "EAFAF1"
BORDER_GRAY = "B0BEC5"
FONT = "Calibri"

# ---------- mandatory PP base template (header + logo + footer + page geometry) ----------
import os
PP_TEMPLATE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "assets", "PP_BASE_TEMPLATE.docx")

# ---------- low-level helpers ----------
def run(p, t, sz, color=BLACK, bold=False, ital=False, font=FONT):
    r = p.add_run(t); f = r.font
    f.name = font; f.size = Pt(sz); f.color.rgb = color; r.bold = bold; r.italic = ital
    return r

def sp(p, before=0, after=0, line=None):
    pf = p.paragraph_format; pf.space_before = Pt(before); pf.space_after = Pt(after)
    if line: pf.line_spacing = line

def shade(cell, fill):
    """Cell fill — ALWAYS w:val='clear' (never 'solid')."""
    tcPr = cell._tc.get_or_add_tcPr(); sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:color'), 'auto'); sh.set(qn('w:fill'), fill)
    tcPr.append(sh)

def cell_borders(cell, sides, sz=4, color="000000"):
    """sides = subset of ('top','left','bottom','right') to draw; others nil."""
    tcPr = cell._tc.get_or_add_tcPr(); b = OxmlElement('w:tcBorders')
    for e in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement('w:' + e)
        if e in sides:
            el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), str(sz)); el.set(qn('w:color'), color)
        else:
            el.set(qn('w:val'), 'nil')
        b.append(el)
    tcPr.append(b)

def table_borders(tbl, sz=4, color=BORDER_GRAY):
    tblPr = tbl._tbl.tblPr; b = OxmlElement('w:tblBorders')
    for e in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement('w:' + e)
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), str(sz)); el.set(qn('w:space'), '0'); el.set(qn('w:color'), color)
        b.append(el)
    tblPr.append(b)

def fixed_widths(tbl, widths_cm):
    tbl.autofit = False; tbl.allow_autofit = False
    lay = OxmlElement('w:tblLayout'); lay.set(qn('w:type'), 'fixed'); tbl._tbl.tblPr.append(lay)
    for row in tbl.rows:
        for i, w in enumerate(widths_cm):
            if i < len(row.cells):       # guard for merged rows (e.g. annex section banners)
                row.cells[i].width = Cm(w)

def cell_margins(cell, top=29, bottom=29, left=58, right=58):
    tcPr = cell._tc.get_or_add_tcPr(); m = OxmlElement('w:tcMar')
    for side, val in (('top', top), ('left', left), ('bottom', bottom), ('right', right)):
        el = OxmlElement('w:' + side); el.set(qn('w:w'), str(val)); el.set(qn('w:type'), 'dxa'); m.append(el)
    tcPr.append(m)

# ==================== intelligent table-layout brain (shared with pp_report.py) ====================
# The single content-aware column-distribution engine (SKILL.md / formatting_specs.md: "Intelligent,
# use-aware column distribution — AUTOMATIC (`fixed(tbl)` is the single table-layout brain)"). Canonical
# home is HERE (pp_format.py) so both the SOP/Annex path and pp_report.py's report/record path share
# ONE implementation — pp_report.py imports `fixed` from this module rather than redefining it.
PAGE_W = 18.46   # A4 text width (cm) at 1.27 cm L/R margins — every table is fitted to this
_ENTRY_TARGETS = [(('потпис', 'signature'), 3.8), (('датум', 'date'), 2.6), (('име', 'name', 'позиц', 'position'), 5.2)]

def _tctext(tc):
    return "".join((n.text or "") for n in tc.iter(qn('w:t')))

def _settcw_dxa(tc, cm):
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcW')):
        tcPr.remove(old)
    w = OxmlElement('w:tcW'); w.set(qn('w:type'), 'dxa'); w.set(qn('w:w'), str(max(1, int(cm * 567)))); tcPr.append(w)

def _gridspan(tc):
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        return 1
    gs = tcPr.find(qn('w:gridSpan'))
    return int(gs.get(qn('w:val'))) if gs is not None else 1

def _repeat_header(tbl, n=1):
    """Mark the first n row(s) as a repeating header so they reprint atop every page the table spans."""
    for tr in tbl._tbl.tr_lst[:n]:
        trPr = tr.get_or_add_trPr()
        if trPr.find(qn('w:tblHeader')) is None:
            e = OxmlElement('w:tblHeader'); e.set(qn('w:val'), 'true'); trPr.append(e)

# --- fixed(): grafted from ACME_SOP repo pp_report.py @ 2026-07-14 (overflow compression + word-boundary entry matching) per DOCENGINE-CANON ---
def fixed(tbl, weights=None, min_cm=0.9, cap=34, mode=None, header_repeat=True):
    """INTELLIGENT, use-aware column distribution (the skill's single table-layout brain).
       It measures each column's real bilingual content and decides the table's *fit*:
         • narrow 'label | value' SUMMARY tables (≤3 columns whose content fits) → COMPACT &
           CENTERED at their natural width (no label stacked vertically, no value floating in a
           vast empty column);
         • wide DATA tables → fill the page width (never overflow);
       ordinal/number columns collapse to just-enough; hand-written ENTRY columns
       (Name/Date/Signature/Position) are purpose-sized even when blank; Role/Action take the
       remainder. The first row is set to repeat on page breaks. Pass `weights` to force ratios,
       or mode='compact'/'full' to override the auto fit."""
    grid=tbl._tbl.tblGrid.gridCol_lst; ncol=len(grid)
    if ncol==0: return tbl
    rows=[tr for tr in tbl._tbl.tr_lst if len(tr.tc_lst)==ncol]
    CHCM=0.176; PAD=0.42                              # ≈ cm per char (10 pt Calibri) + cell padding
    hdr=[_tctext(rows[0].tc_lst[j]).lower().strip() if rows else "" for j in range(ncol)]
    ent={}; forced=bool(weights and len(weights)==ncol)
    # Per-column content measures (for use-aware sizing of compressed data tables):
    #   data_cm = widest DATA cell — data must NEVER wrap, so this is a hard floor;
    #   hdr_cm  = header's full inline width — its demand for space;
    #   word_cm = header's longest single token — a header may wrap, but not below its longest word.
    data_cm=[0.8]*ncol; hdr_cm=[0.8]*ncol; word_cm=[0.6]*ncol
    if forced:
        ideal=[float(w) for w in weights]
    else:
        ideal=[0.8]*ncol
        for ri,tr in enumerate(rows):
            for j,tc in enumerate(tr.tc_lst):
                t=_tctext(tc).replace("\n"," ").strip()
                if not t: continue
                ps=t.split(" | ")
                L=(len(ps[0])+3+0.8*len(ps[1])) if len(ps)==2 else len(t)   # TRUE inline bilingual width (MK + sep + smaller EN)
                w=min(L,cap)*CHCM+PAD
                ideal[j]=max(ideal[j], w)
                if ri==0:
                    hdr_cm[j]=w
                    toks=t.replace(" | "," ").split()
                    word_cm[j]=(max((len(x) for x in toks), default=1))*CHCM+PAD
                else:
                    data_cm[j]=max(data_cm[j], w)
        for j,h in enumerate(hdr):                    # purpose-size hand-written entry columns
            for kws,cm in _ENTRY_TARGETS:             # match at a WORD boundary (prefix ok) so e.g. the
                if any(re.search(r'\b'+re.escape(k), h) for k in kws):  # 'име' in 'примерок' is NOT a hit
                    ent[j]=cm; break
        for j,cm in ent.items(): ideal[j]=cm
    ideal_sum=sum(ideal) or 1.0
    if   mode=="full":    compact=False
    elif mode=="compact": compact=True
    else: compact=(not forced) and ncol<=3 and ideal_sum<=0.86*PAGE_W and not ent
    if compact:
        widths=[max(w,min_cm) for w in ideal]
    else:
        if ent:
            targ=sum(ent.values())
            if targ>PAGE_W-2.0:
                sc=(PAGE_W-2.0)/targ; ent={j:c*sc for j,c in ent.items()}; targ=sum(ent.values())
            oth=[j for j in range(ncol) if j not in ent]; sw=sum(ideal[j] for j in oth) or 1.0; rem=PAGE_W-targ
            widths=[(ent[j] if j in ent else ideal[j]/sw*rem) for j in range(ncol)]
        else:
            even=PAGE_W/ncol
            if ncol>=4 and all(ideal[j]<=even+1e-6 for j in range(1,ncol)):  # labeled grid: data cols uniform & fit
                c0=max(ideal[0],min_cm); rest=(PAGE_W-c0)/(ncol-1)            # col0 sized to its labels (single-line header),
                widths=[c0]+[rest]*(ncol-1)                                   # columns 2..N split the remainder EVENLY
            elif (not forced) and ideal_sum>PAGE_W:                           # table would OVERFLOW → intelligent compression:
                base=[max(data_cm[j], word_cm[j]) for j in range(ncol)]       #   each column floored to its DATA (never wraps)
                sb=sum(base)                                                  #   and its header's longest word (headers may wrap);
                if sb>=PAGE_W:                                                #   if even the floors overflow, share proportionally,
                    widths=[b/sb*PAGE_W for b in base]
                else:                                                         #   else give the spare width to the LONGEST headers
                    slack=PAGE_W-sb; hw=sum(hdr_cm) or 1.0                    #   so headers wrap least where they are longest.
                    widths=[base[j]+slack*hdr_cm[j]/hw for j in range(ncol)]
            else:
                widths=[w/ideal_sum*PAGE_W for w in ideal]                    # else content-proportional
        for _ in range(6):                            # enforce a minimum, redistribute the rest
            below=[i for i,w in enumerate(widths) if w<min_cm-1e-6]
            if not below: break
            for i in below: widths[i]=min_cm
            ab=[i for i in range(ncol) if i not in below]; rem=PAGE_W-min_cm*len(below); sab=sum(widths[i] for i in ab) or 1.0
            for i in ab: widths[i]=widths[i]/sab*rem
    total=sum(widths)
    tbl.autofit=False; tbl.allow_autofit=False
    tblPr=tbl._tbl.tblPr
    for tag in ('w:tblLayout','w:tblW','w:tblInd','w:jc'):
        for old in tblPr.findall(qn(tag)): tblPr.remove(old)
    lay=OxmlElement('w:tblLayout'); lay.set(qn('w:type'),'fixed'); tblPr.append(lay)
    tw=OxmlElement('w:tblW'); tw.set(qn('w:type'),'dxa'); tw.set(qn('w:w'),str(int(total*567))); tblPr.append(tw)
    ti=OxmlElement('w:tblInd'); ti.set(qn('w:type'),'dxa'); ti.set(qn('w:w'),'0'); tblPr.append(ti)
    jc=OxmlElement('w:jc'); jc.set(qn('w:val'),'center'); tblPr.append(jc)   # centre on the page (matters for compact)
    for j,gc in enumerate(grid): gc.set(qn('w:w'), str(max(1,int(widths[j]*567))))
    for tr in tbl._tbl.tr_lst:
        tcs=tr.tc_lst
        if len(tcs)==ncol:
            for j,tc in enumerate(tcs): _settcw_dxa(tc, widths[j])
        else:
            gi=0
            for tc in tcs:
                spn=_gridspan(tc); _settcw_dxa(tc, sum(widths[gi:gi+spn]) or min_cm); gi+=spn
    if header_repeat: _repeat_header(tbl)
    return tbl
def bilingual(p, mk, en, mk_sz=11, en_sz=7, bold=False, color=BLACK):
    """Inline 'MK | EN' run group (annex + table style). Omits the separator when en is empty
    (e.g. a data-entry cell with no English counterpart) to avoid a stray trailing ' |'."""
    run(p, mk, mk_sz, color, bold=bold)
    if en:
        run(p, " | ", en_sz, color)
        run(p, en, en_sz, color, bold=bold)

def _page(doc, orient="portrait", margin_cm=1.27):
    s = doc.sections[0]
    if orient == "landscape":
        s.orientation = WD_ORIENT.LANDSCAPE; s.page_width = Cm(29.7); s.page_height = Cm(21.0)
    else:
        s.orientation = WD_ORIENT.PORTRAIT; s.page_width = Cm(21.0); s.page_height = Cm(29.7)
    s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Cm(margin_cm)
    return s

def _normal(doc):
    st = doc.styles['Normal']; st.font.name = FONT; st.font.size = Pt(11)

# ============================ mandatory header / base template ============================
def wipe_body(d):
    """Clear all body content but keep the section <w:sectPr> (header/footer/page refs).
    Stashes it on d._pp_sectpr; save() re-seats it last."""
    body = d.element.body
    sect = body.find(qn('w:sectPr'))
    for ch in list(body):
        if ch is not sect:
            body.remove(ch)
    d._pp_sectpr = sect
    return sect

def apply_pp_header(d, mk_name, code, en_name, version="1.0"):
    """Stamp the base-template running header: bilingual Document name, Code of document, Version.
    Header is a 2x3 table: [leaf logo+wordmark] | [Document name MK/EN] | [Code of document / Ver].
    Only the value runs are overwritten; the labels and the logo are preserved from the template.
    The document CODE is stamped ONLY in the right-hand 'Code of document' cell — the template's
    own centre-cell run layout has a middle slot (originally used for a secondary identifier) that
    must NOT be reused for the code, or the code appears twice on the same header. That middle
    slot (and its adjoining separator/newline runs) is cleared instead, so the centre cell shows
    just the bilingual title on two lines: MK title, then EN title."""
    try:
        h = d.sections[0].header.tables[0]
    except (IndexError, AttributeError):
        return False
    def setrun(runs, i, t):
        if i < len(runs):
            runs[i].text = t
    nm = h.cell(0, 1).paragraphs[1].runs
    setrun(nm, 0, mk_name + " ")
    for i in (1, 2, 3):          # clear the template's middle line + separator (was: code, duplicating the right cell)
        setrun(nm, i, "")
    setrun(nm, 5, en_name)
    cd = h.cell(0, 2).paragraphs[2].runs
    setrun(cd, 0, code); setrun(cd, 2, ""); setrun(cd, 3, "")
    vr = h.cell(1, 2).paragraphs[0].runs
    if vr:
        vr[-1].text = version
    return True

# ============================ SOP (two-column) ============================
def new_sop(margin_cm=1.27, from_template=True, code=None,
            mk_name="СТАНДАРДНА ОПЕРАТИВНА ПРОЦЕДУРА", en_name=None, version="1.0", template=None,
            mk_title=None, en_title=None):
    """SOP document. By DEFAULT starts FROM the mandatory PP base template so the running header
    (logo + bilingual doc name + code + version), the 'Page X of Y' footer and the A4 page geometry
    are present on every page.
    The running header's 'Document name' shows the ACTUAL SOP title — pass mk_title (Macedonian)
    and en_title (English). These take precedence over the generic 'Standard Operating Procedure'
    phrase (mk_name/en_name), which is used ONLY as a fallback when no title is supplied."""
    tpl = template or PP_TEMPLATE
    hdr_mk = mk_title or mk_name
    hdr_en = en_title or en_name or ("STANDARD OPERATING PROCEDURE — %s" % (code or ""))
    if from_template and os.path.exists(tpl):
        d = Document(tpl)
        apply_pp_header(d, hdr_mk, code or "", hdr_en, version)
        wipe_body(d)
        _normal(d)
        return d
    d = Document(); _page(d, "portrait", margin_cm); _normal(d); return d

def sop_titlepage(d, code, mk_title, en_title, roles=None, effective_date=True):
    """Title/approval page. `roles`, if given, is a list of (mk_role, en_role, name) triples —
    one signature row per approval role (e.g. Prepared/QC-Reviewed/QA-Reviewed/Approved), so any
    number of tiers is supported without hardcoding personnel names into the engine (personnel
    identity belongs to the document being built, never to the generator). Defaults to the
    3-tier Prepared/Reviewed/Approved convention with blank names for the user to fill in.
    Columns follow the house convention: Улога|Role · Име|Name · Потпис|Signature · Датум|Date.
    Set effective_date=False to omit the trailing 'Effective Date' row."""
    for txt, sz, col, bold in [("СТАНДАРДНА ОПЕРАТИВНА ПРОЦЕДУРА", 24, NAVY, True),
                               ("STANDARD OPERATING PROCEDURE", 14, GREY, False),
                               (code, 18, NAVY, True)]:
        p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; sp(p, 6, 4); run(p, txt, sz, col, bold=bold)
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; sp(p, 4, 12)
    run(p, mk_title, 12, BLACK, bold=True); run(p, " | ", 10, GREY); run(p, en_title, 12, GREY, bold=True)
    if roles is None:
        roles = [("Подготвено од", "Prepared by", ""),
                  ("Проверено од", "Reviewed by", ""),
                  ("Одобрено од", "Approved by", "")]
    nrows = 2 + len(roles) + (1 if effective_date else 0)
    t = d.add_table(rows=nrows, cols=4); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    band = t.rows[0].cells
    for c in band:
        shade(c, SOP_BLUE); c.paragraphs[0].text = ''
    bilingual(band[0].paragraphs[0], "ОДОБРУВАЊЕ НА ДОКУМЕНТ", "DOCUMENT APPROVAL", 11, 7, bold=True)
    hdr = [("Улога", "Role"), ("Име", "Name"), ("Потпис", "Signature"), ("Датум", "Date")]
    for j, (mk, en) in enumerate(hdr):
        c = t.cell(1, j); shade(c, SOP_BLUE); c.paragraphs[0].text = ''
        bilingual(c.paragraphs[0], mk, en, 11, 7, bold=True)
    for i, (mk_role, en_role, name) in enumerate(roles, 2):
        c0 = t.cell(i, 0); c0.paragraphs[0].text = ''; bilingual(c0.paragraphs[0], mk_role, en_role, 10, 7)
        if name:
            run(t.cell(i, 1).paragraphs[0], name, 10)
    if effective_date:
        r = nrows - 1
        c0 = t.cell(r, 0); c0.paragraphs[0].text = ''
        bilingual(c0.paragraphs[0], "Датум на ефективност", "Effective Date", 10, 7)
    fixed_widths(t, [4.7, 4.7, 4.7, 4.7]); table_borders(t, 4, "000000")
    d.add_page_break()

def sop_toc(d):
    p = d.add_paragraph(); run(p, "СОДРЖИНА | TABLE OF CONTENTS", 14, NAVY, bold=True)
    par = d.add_paragraph(); fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), r'TOC \o "1-3" \h \z \u'); par._p.append(fld)
    note = d.add_paragraph(); run(note, "(Update field after opening: right-click → Update Field)", 8, GREY, ital=True)
    d.add_page_break()

def sop_table(d):
    """Empty 2-column table to hold section/body rows."""
    t = d.add_table(rows=0, cols=2); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    return t

def _set_heading(p, level=1):
    """Apply a Heading style so the native TOC field populates. Run-level font is set afterwards,
    preserving the house appearance (the heading style only makes the paragraph TOC-eligible)."""
    try:
        p.style = "Heading %d" % max(1, min(3, level))
    except Exception:
        pass

def sop_section_row(t, num, mk, en, level=1):
    """Two-column section header. level controls the TOC depth (1=1.0, 2=1.1, 3=1.1.1)."""
    L, R = t.add_row().cells
    for c in (L, R):
        shade(c, SOP_GRAY)
    cell_borders(L, ('right',)); cell_borders(R, ('left',))
    L.paragraphs[0].text = ''; _set_heading(L.paragraphs[0], level); run(L.paragraphs[0], f"{num} {mk}", 12, BLACK, bold=True)
    R.paragraphs[0].text = ''; run(R.paragraphs[0], f"{num} {en}", 12, GREY, bold=True)

def sop_body_row(t, mk, en, bold=False):
    L, R = t.add_row().cells
    cell_borders(L, ('right',)); cell_borders(R, ('left',))
    for cell, txt, col in ((L, mk, BLACK), (R, en, GREY)):
        p = cell.paragraphs[0]; p.text = ''; p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; sp(p, 0, 6)
        run(p, txt, 11, col, bold=bold)

def sop_finalize(t):
    fixed_widths(t, [9.0, 9.0])   # two equal columns within ~18.46 cm usable

def sop_nested_table(t, headers, rows, widths_cm=None, mode=None, header_bg=SOP_BLUE):
    """Full-width table NESTED inside a merged SOP body row (SKILL.md / formatting_specs.md:
    'Tables nested in a full-width merged row; no divider; table text MK 11pt | EN 7pt').
    headers, and each row, are lists of (mk, en) pairs — one per nested-table column.
    Use for in-body tables (RACI matrices, abbreviation lists, colour-code legends, retention
    registers, revision history) that must live INSIDE the running two-column SOP body, not as
    a separate top-level table. Distinct from sop_block_table(), which builds a standalone
    top-level table for §7/§8/§9 (outside the two-column flow).
    Column widths are AUTO-DECIDED by fixed() — the skill's single use-aware table-layout brain
    (formatting_specs.md): a 2-column abbreviation list gets a COMPACT centred fit, a wide RACI
    matrix gets its label column sized to content with role columns split evenly, etc. Pass
    widths_cm to force explicit ratios, or mode='compact'|'full' to override the auto fit —
    do not hand-compute an even split; that is exactly the anti-pattern fixed() exists to avoid."""
    outer_row = t.add_row()
    L, R = outer_row.cells
    merged = L.merge(R)
    cell_borders(merged, ())  # no divider across a merged/nested-table row
    merged.paragraphs[0].text = ''
    ncol = len(headers)
    nt = merged.add_table(rows=1 + len(rows), cols=ncol)
    nt.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, (mk, en) in enumerate(headers):
        c = nt.cell(0, j); shade(c, header_bg); c.paragraphs[0].text = ''
        bilingual(c.paragraphs[0], mk, en, 11, 7, bold=True)
    for i, row in enumerate(rows, 1):
        for j, (mk, en) in enumerate(row):
            if j >= ncol:
                break
            c = nt.cell(i, j); c.paragraphs[0].text = ''; bilingual(c.paragraphs[0], mk, en, 11, 7)
    fixed(nt, weights=widths_cm, mode=mode)
    table_borders(nt, 4, BORDER_GRAY)
    return nt

def sop_block_table(d, headers, rows, widths_cm, header_bg=SOP_BLUE):
    """Full-width multi-column bilingual table for SOP §7 Records / §8 Related Documents / §9 Revision.
    headers, and each row, are lists of (mk, en) tuples (one per column)."""
    t = d.add_table(rows=1 + len(rows), cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, (mk, en) in enumerate(headers):
        c = t.cell(0, j); shade(c, header_bg); c.paragraphs[0].text = ''
        bilingual(c.paragraphs[0], mk, en, 11, 7, bold=True)
    for i, row in enumerate(rows, 1):
        for j, (mk, en) in enumerate(row):
            c = t.cell(i, j); c.paragraphs[0].text = ''; bilingual(c.paragraphs[0], mk, en, 11, 7)
    fixed_widths(t, widths_cm); table_borders(t, 4, "000000"); return t

# ============================ Annex (inline) ============================
def new_annex(orient="portrait", from_template=True, code=None,
              mk_name="АНЕКС", en_name=None, version="1.0", template=None,
              mk_title=None, en_title=None):
    """Annex document. Like new_sop, by DEFAULT starts FROM the PP base template so the mandatory
    header, logo and 'Page X of Y' footer are present (required on all annexes).
    The header 'Document name' shows the ACTUAL annex title (mk_title | en_title); the generic
    'ANNEX' phrase (mk_name/en_name) is used only as a fallback when no title is supplied."""
    tpl = template or PP_TEMPLATE
    hdr_mk = mk_title or mk_name
    hdr_en = en_title or en_name or ("ANNEX — %s" % (code or ""))
    if from_template and os.path.exists(tpl):
        d = Document(tpl)
        apply_pp_header(d, hdr_mk, code or "", hdr_en, version)
        wipe_body(d)
        _normal(d)
        if orient == "landscape":
            s = d.sections[0]; s.orientation = WD_ORIENT.LANDSCAPE
            s.page_width, s.page_height = Cm(29.7), Cm(21.0)
        return d
    d = Document(); _page(d, orient, 2.54 if orient == "portrait" else 1.27); _normal(d); return d

def annex_title_block(d, code, mk_title, en_title, parent_sop):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; sp(p, 2, 3); run(p, code, 12, NAVY, bold=True)
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; sp(p, 0, 3)
    run(p, mk_title, 14, BLACK, bold=True); run(p, " | ", 10, GREY); run(p, en_title, 10, GREY, bold=True)
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; sp(p, 0, 8); run(p, f"({parent_sop})", 10, GREY, ital=True)

def annex_table(d, widths_cm=None, ncol=None, mode=None):
    """widths_cm given → FORCED literal column widths (unchanged legacy behaviour, exact cm values,
    no rescaling — existing call sites are unaffected). widths_cm omitted → pass ncol instead and
    annex_finalize() will size columns AUTOMATICALLY via fixed() (formatting_specs.md's use-aware
    table-layout brain): compact-centred for a narrow label|value summary, content-proportional or
    evenly-split for a wide data grid. Do not hand-compute an even split; let fixed() decide the fit."""
    n = len(widths_cm) if widths_cm else ncol
    if not n:
        raise ValueError("annex_table requires either widths_cm or ncol")
    t = d.add_table(rows=0, cols=n); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t._pp_widths = widths_cm
    t._pp_mode = mode
    return t

def annex_cell(cell, mk, en, bold=False, bg=None):
    if bg:
        shade(cell, bg)
    color = WHITE if bg == A_SECTION else BLACK
    cell_margins(cell); cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]; p.text = ''; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; sp(p, 1, 1, line=0.85)
    bilingual(p, mk, en, 11, 7, bold=bold, color=color)

def annex_section_header(t, mk, en):
    """Full-width navy banner row (white text). Call after rows exist; merges the row."""
    row = t.add_row(); first = row.cells[0]
    if len(row.cells) > 1:
        first = first.merge(row.cells[-1])
    annex_cell(first, mk, en, bold=True, bg=A_SECTION)

def annex_row(t, values, bold_first=False, alt=False):
    """values = list of (mk,en) tuples, one per column."""
    row = t.add_row()
    for i, (mk, en) in enumerate(values):
        bg = A_LABEL if (bold_first and i == 0) else (A_ALT if alt else None)
        annex_cell(row.cells[i], mk, en, bold=(bold_first and i == 0), bg=bg)

def annex_finalize(t):
    """FORCED widths (t._pp_widths set) → legacy literal fixed_widths(), byte-for-byte unchanged
    for every existing caller. No forced widths (annex_table was given ncol only) → AUTO via
    fixed(), the same content-aware brain used by sop_nested_table and pp_report.py."""
    if t._pp_widths:
        fixed_widths(t, t._pp_widths)
    else:
        fixed(t, mode=getattr(t, '_pp_mode', None))
    table_borders(t, 1, BORDER_GRAY)

def annex_signoff(d, roles=(("Изготвил (КК)", "Prepared (QC)"), ("Проверил (QA)", "Reviewed (QA)"), ("Одобрил (QP)", "Approved (QP)"))):
    w = round(18.0 / len(roles), 2)
    t = annex_table(d, [w] * len(roles))
    annex_row(t, list(roles), bold_first=False)
    for c in t.rows[0].cells:
        shade(c, A_SECTION)
        for r in c.paragraphs[0].runs:
            r.font.color.rgb = WHITE; r.font.bold = True
    t.add_row()  # blank signature row
    annex_finalize(t)
    return t

# ============================ compact label|value layout (2026-06-30) ============================
# Intelligent compact layout for annex/form metadata blocks (Head-of-QC corrections 2026-06-30):
#   * VALUE cells sized to the EXPECTED HAND-ENTRY, never maximised (a Date/Batch needs ~2.6 cm).
#   * LABEL ("action") cells minimal width.
#   * Multiple label|value pairs PACK per row (default 6-col grid); long values SPAN.
#   * Line spacing 0.8, NO space before/after in every cell (incl. label cells).
#   * Related sections (e.g. Transfer information + Material description) merge into ONE table.
# These are first-class pf primitives; pp_format_layout_addons.py re-exports them. See SKILL.md §6D.

def cell08(cell, mk, en, mk_sz=9, en_sz=7, bold=False, bg=None, white=False, left=True):
    """House cell, COMPACT form variant: 0.8 line spacing, zero space before/after, bilingual MK | EN.
    Distinct from annex_cell() (centered, 11|7, 0.85 spacing) — use cell08 for packed metadata grids."""
    if bg:
        shade(cell, bg)
    cell_margins(cell); cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]; p.text = ''
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if left else WD_ALIGN_PARAGRAPH.CENTER
    fmt = p.paragraph_format; fmt.line_spacing = 0.8; fmt.space_before = Pt(0); fmt.space_after = Pt(0)
    col = WHITE if white else BLACK
    if mk:
        run(p, mk, mk_sz, col, bold=bold)
    if en:
        if mk:
            run(p, " | ", en_sz, col)
        run(p, en, en_sz, col, bold=False)

def _merge(row, a, b):
    """Merge cells a..b (inclusive) of a row; return the merged cell."""
    c = row.cells[a]
    for k in range(a + 1, b + 1):
        c = c.merge(row.cells[k])
    return c

def value_span(input_chars, ncol=6):
    """Map expected hand-entry length -> a sensible value-column count (1..ncol-1).
    <=12 chars -> 1 value col (input-sized); grows for longer/checkbox values."""
    if input_chars <= 12:
        return 1
    if input_chars <= 28:
        return 2
    if input_chars <= 48:
        return 3
    return ncol - 1  # full-width span (checkbox rows, long boilerplate)

def kv_block(d, sections, ncol=6, label_cm=1.9, value_cm=2.6):
    """Compact merged label|value block — the reusable metadata-layout primitive (SKILL.md §6D).

    sections: list of (banner_mk, banner_en, fields);
      fields: list of (label_mk, label_en, value_mk, value_en, input_chars).
    Builds ONE centred table; value cells are input-sized (via value_span); short fields PACK
    multiple pairs per row; long values SPAN; every cell uses line spacing 0.8 / no space
    before-after. Related sections are stacked into the SAME table. Returns the table."""
    widths = [(label_cm if i % 2 == 0 else value_cm) for i in range(ncol)]
    t = annex_table(d, widths)
    for bmk, ben, fields in sections:
        r = t.add_row()
        cell08(_merge(r, 0, ncol - 1), bmk, ben, 10, 8, bold=True, bg=A_SECTION, white=True, left=False)
        col, r = ncol, None  # force a new row on the first field
        for lmk, len_, vmk, ven, ichars in fields:
            span = value_span(ichars, ncol); need = 1 + span
            if r is None or col + need > ncol:
                r = t.add_row(); col = 0
            cell08(r.cells[col], lmk, len_, 9, 7, bold=True, bg=A_LABEL)
            v = _merge(r, col + 1, col + span) if span > 1 else r.cells[col + 1]
            cell08(v, vmk, ven, 9, 7)
            col += need
        if r is not None and col < ncol:  # blank-fill the tail of the last row
            cell08(_merge(r, col, ncol - 1), "", "")
    annex_finalize(t)
    return t

# ============================ save ============================
def save(d, path):
    sect = getattr(d, "_pp_sectpr", None)
    if sect is not None:
        body = d.element.body
        try:
            body.remove(sect); body.append(sect)
        except ValueError:
            pass
    d.save(path); return path

def figure(d, png_path, mk=None, en=None, width_cm=15.5):
    """Embed a chart PNG (centered) with a bilingual MK | EN caption. Generate charts with pp_charts.py.
    Charts are data visualisations (matplotlib PNGs) — distinct from formula notation, which stays as
    native Word equations. Use for statistical justification figures in reports."""
    d.add_picture(png_path, width=Cm(width_cm))
    d.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if mk or en:
        p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; sp(p, 1, 8)
        run(p, "Слика | Figure — ", 9, NAVY, bold=True)
        if mk: run(p, mk, 9, GREY, ital=True)
        if en: run(p, "  |  ", 8, GREY); run(p, en, 9, GREY, ital=True)
    return d
