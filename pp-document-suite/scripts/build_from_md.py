#!/usr/bin/env python3
"""
build_from_md.py — render bilingual Markdown into a Purely Plant controlled .docx
USING THE pp-document-suite ENGINE, with the skill's INTELLIGENT table layout.

  doctype SOP  -> pp_format two-column (MK|EN, divider, native TOC); §7/§8/§9 nested full-width
                  tables sized by pp_report.fixed() (the layout brain).
  doctype ANNEX/FORM/CHECKLIST/LOG -> pp_format base-template shell (logo header + Page X of Y),
                  navy #2B547E section banners, and ALL tables built via pp_report:
                    - label|value forms   -> 2-col, label #F2F2F2, value entry; fixed() purpose-sizes
                    - data tables         -> navy header row, zebra rows, fixed() use-aware widths +
                                             repeating header on page breaks
                    - single-select lists -> status_grid() compact checkbox grid
                  All widths are decided by pp_report.fixed() (compact-centred summaries vs full-width
                  data; Name/Date/Signature entry columns purpose-sized; ordinal columns minimal).

Markdown convention: PP_UNIFIED_DOCX_GUIDE.md. Usage: python3 scripts/build_from_md.py <src.md> <out.docx>
"""
import sys, os, re
# Locate the engine modules (pp_format/pp_report/...). Inside the packaged skill they are siblings
# of this file; in the ACME_SOP repo layout they live under ../pp-document-suite/scripts. Try both.
_here = os.path.dirname(os.path.abspath(__file__))
for _cand in (_here, os.path.join(_here, "..", "pp-document-suite", "scripts")):
    if os.path.exists(os.path.join(_cand, "pp_format.py")):
        sys.path.insert(0, _cand); break
sys.path.insert(0, _here)
import pp_format as pf
import pp_report as pr
from docx.enum.text import WD_ALIGN_PARAGRAPH

PAGE_W = pr.PAGE_W  # 18.46 cm text width (base-template 1.27 cm margins)

def parse(md):
    hd={}
    m=re.search(r'<!--HEADERDATA(.*?)-->', md, re.S)
    if m:
        for ln in m.group(1).strip().splitlines():
            if ':' in ln:
                k,v=ln.split(':',1); hd[k.strip()]=v.strip()
        md=md[m.end():]
    blocks=[]; lines=md.splitlines(); i=0
    while i<len(lines):
        s=lines[i].strip()
        if not s: i+=1; continue
        if s.startswith('[[PAGEBREAK]]') or s.startswith('[[NEWPAGE]]'):
            blocks.append(('pagebreak',)); i+=1; continue
        if s.startswith('[[TABLE') or s.startswith('[[FORM'):
            kind='form' if s.startswith('[[FORM') else 'table'; rows=[]
            mo=re.match(r'\[\[(?:FORM|TABLE):([^\]]+)\]\]', s); mode=mo.group(1).strip() if mo else None
            i+=1
            while i<len(lines):
                t=lines[i].strip()
                if t.startswith('[[/'): i+=1; break
                if t.startswith('[[TABLE') or t.startswith('[[FORM') or t.startswith('#'): break
                if t: rows.append([c.strip() for c in lines[i].split('|||')])
                i+=1
            blocks.append((kind, rows, mode)); continue
        hm=re.match(r'^(#+)\s+(.*)', s)
        if hm:
            lvl=len(hm.group(1)); mk,en=(hm.group(2).split('|',1)+[''])[:2]; num=''
            nm=re.match(r'^([0-9]+(?:\.[0-9]+)*\.?)\s+(.*)', mk.strip())
            if nm: num=nm.group(1); mk=nm.group(2)
            blocks.append(('h', lvl, num, mk.strip(), en.strip())); i+=1; continue
        bm=re.match(r'^[-*]\s+(.*)', s)
        if bm:
            mk,en=(bm.group(1).split('|||',1)+[''])[:2]
            blocks.append(('bullet', mk.strip(), en.strip())); i+=1; continue
        mk,en=(s.split('|||',1)+[''])[:2]
        blocks.append(('p', mk.strip(), en.strip())); i+=1
    return hd, blocks

def cs(raw):
    mk,en=(raw.split('~~',1)+[''])[:2]; return mk.strip(), en.strip()

def bil(mk,en):
    return (mk + (" | "+en if en else "")).strip()

# ---- single-select option list detection (-> status_grid) ----
OPT = re.compile(r'☐\s*([^☐]+?)\s*(?=☐|$)')
def as_options(rows):
    """If a FORM/TABLE block is really one 'select one' line of ☐ options, return the option list."""
    flat=[c for r in rows for c in r if c.strip()]
    joined=" ".join(flat)
    if joined.count('☐')>=2 and len(rows)<=2 and '~~' not in joined and '|||' not in joined:
        opts=[o.strip() for o in OPT.findall(joined) if o.strip()]
        if len(opts)>=2: return opts
    return None

# =========================== ANNEX ===========================
def emit_banner(d, num, mk, en):
    t=d.add_table(rows=1, cols=1); t.alignment=pr.WD_TABLE_ALIGNMENT.CENTER
    pr.cellfmt(t.cell(0,0), (num+" "+mk).strip(), en, sz=11, col=pr.WHITE, bold=True, fill=pr.NAVYF)
    pr.fixed(t, [PAGE_W], header_repeat=False); pr.borders(t)

def emit_subbar(d, num, mk, en):
    t=d.add_table(rows=1, cols=1); t.alignment=pr.WD_TABLE_ALIGNMENT.CENTER
    pr.cellfmt(t.cell(0,0), (num+" "+mk).strip(), en, sz=10, col=pr.NAVY, bold=True, fill=pr.LBL)
    pr.fixed(t, [PAGE_W], header_repeat=False); pr.borders(t)

def _form_bulk(cell, lmk, len_, val):
    """Render a bulk-text field as a full-width stacked block: bold label heading, then the value
       as a justified paragraph spanning the whole cell."""
    pr.shade(cell, pr.LBL)
    p=cell.paragraphs[0]; p.text=''; p.alignment=WD_ALIGN_PARAGRAPH.LEFT; pr.sp(p,2,2)
    if lmk: pr.rin(p, lmk, 10, pr.BLACK, bold=True)
    if len_:
        if lmk: pr.rin(p, " | ", 8, pr.BLACK)
        pr.rin(p, len_, 8, pr.BLACK, ital=True)
    pv=cell.add_paragraph(); pv.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY; pr.sp(pv,1,2)
    vmk,ven=val.split(" | ",1) if " | " in val else (val,'')   # form values carry MK/EN as "MK | EN"
    if vmk.strip(): pr.rin(pv, vmk.strip(), 10, pr.BLACK)
    if ven.strip():
        if vmk.strip(): pr.rin(pv, "  |  ", 8, pr.BLACK)
        pr.rin(pv, ven.strip(), 8, pr.BLACK, ital=True)

def _lbl_cell(cell, lmk, len_):
    c=pr.cellfmt(cell, lmk, len_, sz=10, col=pr.BLACK, bold=True, fill=pr.LBL)
    c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.LEFT

def emit_form(d, rows, mode=None):
    # single-select? -> status_grid
    opts=as_options(rows)
    if opts:
        pr.status_grid(d, opts, selected=None, ncols=min(2, len(opts))); return
    fields=[]
    for rd in rows:
        lmk=rd[0] if len(rd)>0 else ''; len_=rd[1] if len(rd)>1 else ''; val=rd[2] if len(rd)>2 else ''
        val=val if (val and val.strip()!='_') else ''
        vlen=len(val.strip())
        kind='bulk' if vlen>110 else ('span' if vlen>20 else 'short')   # short = fits a half-page pair
        fields.append({'lmk':lmk,'en':len_,'val':val,'kind':kind})

    if mode!='grid':
        # DEFAULT one-field-per-row: compact label col (see notes above), bulk values span full width.
        t=d.add_table(rows=len(fields), cols=2); t.alignment=pr.WD_TABLE_ALIGNMENT.CENTER
        lw=2.0
        for f in fields:
            if f['kind']!='bulk': lw=max(lw, min(max(len(f['lmk']),len(f['en'])),34)*0.176+0.5)
        lw=min(lw, PAGE_W*0.45)
        for ri,f in enumerate(fields):
            if f['kind']=='bulk':
                _form_bulk(t.cell(ri,0).merge(t.cell(ri,1)), f['lmk'], f['en'], f['val'])
            else:
                _lbl_cell(t.cell(ri,0), f['lmk'], f['en'])
                pr.cellfmt(t.cell(ri,1), f['val'], None, sz=10, col=pr.BLACK)
        pr.fixed(t, weights=[lw, PAGE_W-lw], header_repeat=False); pr.borders(t)
        return

    # GRID MODE: content-aware packing. Each field is measured on ITS OWN content — the label sized
    # to its own text, the value to what it holds (write-ins get a fixed minimum). Fields are then
    # packed greedily, as many per row as GENUINELY fit the page, and every column is sized to its own
    # content — no shared/global label or value width. Long values span; bulk values become a block.
    VMIN=2.6                                                   # minimum writable value width (cm)
    for f in fields:
        # Label sized to its OWN text: wide enough for the whole label (capped, may wrap at spaces),
        # but never below its longest single WORD — so a short label like "Датум" never breaks mid-word.
        toks=(f['lmk']+' '+f['en']).split()
        word=max((len(x) for x in toks), default=1); full=max(len(f['lmk']),len(f['en']))
        f['lw']=max(word*0.24+0.6, min(full*0.24+0.6, 5.8))
        vl=len(f['val'].strip())
        f['vw']=VMIN if vl==0 else max(vl*0.20+0.4, 1.4)       # value → its own text (write-in = VMIN)
        f['fw']=f['lw']+max(f['vw'],VMIN)+0.3                  # footprint if packed onto a shared row
        f['span']=(f['kind']!='bulk') and (f['vw']>4.6 or f['fw']>PAGE_W*0.62)  # too wide to pair

    plans=[]; cur=[]; curw=0.0                                 # greedy first-fit, order preserving
    def flush():
        nonlocal cur,curw
        if cur: plans.append(('pack',cur)); cur=[]; curw=0.0
    for f in fields:
        if f['kind']=='bulk': flush(); plans.append(('bulk',f)); continue
        if f['span']:         flush(); plans.append(('span',f)); continue
        if cur and curw+f['fw']>PAGE_W: flush()
        cur.append(f); curw+=f['fw']
    flush()

    for pl in plans:
        if pl[0]=='bulk':
            t=d.add_table(rows=1, cols=1); t.alignment=pr.WD_TABLE_ALIGNMENT.CENTER
            _form_bulk(t.cell(0,0), pl[1]['lmk'], pl[1]['en'], pl[1]['val'])
            pr._apply_widths(t, [PAGE_W])
        elif pl[0]=='span':
            f=pl[1]; t=d.add_table(rows=1, cols=2); t.alignment=pr.WD_TABLE_ALIGNMENT.CENTER
            _lbl_cell(t.cell(0,0), f['lmk'], f['en'])
            pr.cellfmt(t.cell(0,1), f['val'], None, sz=10, col=pr.BLACK)
            pr._apply_widths(t, [f['lw'], PAGE_W-f['lw']])
        else:  # pack: N fields side by side; labels snug to content, values share the remainder
            fs=pl[1]; n=len(fs); t=d.add_table(rows=1, cols=2*n); t.alignment=pr.WD_TABLE_ALIGNMENT.CENTER
            slack=(PAGE_W-sum(f['lw'] for f in fs))/n           # each value = equal share of leftover
            widths=[]
            for j,f in enumerate(fs):
                _lbl_cell(t.cell(0,2*j), f['lmk'], f['en'])
                pr.cellfmt(t.cell(0,2*j+1), f['val'], None, sz=10, col=pr.BLACK)
                widths+=[f['lw'], max(slack, VMIN)]
            pr._apply_widths(t, widths)

def emit_table(d, rows):
    if not rows: return
    opts=as_options(rows)
    if opts:
        pr.status_grid(d, opts, selected=None, ncols=min(2, len(opts))); return
    ncol=max(len(r) for r in rows)
    t=d.add_table(rows=len(rows), cols=ncol); t.alignment=pr.WD_TABLE_ALIGNMENT.CENTER
    for ri,rd in enumerate(rows):
        for ci in range(ncol):
            mk,en=cs(rd[ci]) if ci<len(rd) else ('','')
            if ri==0:
                pr.cellfmt(t.cell(ri,ci), mk, en, sz=9, col=pr.WHITE, bold=True, fill=pr.NAVYF)
            else:
                fill=pr.LBL if (ci==0 and mk) else ("F7FAFC" if ri%2==0 else None)
                pr.cellfmt(t.cell(ri,ci), mk, en, sz=9, col=pr.BLACK, fill=fill)
    pr.fixed(t)                # intelligent, use-aware distribution + repeat header row
    pr.borders(t)

def build_annex(hd, blocks, out):
    global PAGE_W
    land = hd.get('orient')=='landscape'
    # Fit tables/forms to the ACTUAL text width of the page: A4 portrait ≈ 18.46 cm, A4 landscape
    # ≈ 27.16 cm (both at the base-template 1.27 cm margins). The whole layout brain keys off this.
    PAGE_W = 27.16 if land else 18.46
    pr.PAGE_W = PAGE_W
    d=pf.new_annex(code=hd.get('code',''), version=hd.get('version','01'),
                   mk_title=hd.get('mk_title',''), en_title=hd.get('en_title',''),
                   orient='landscape' if land else 'portrait')
    pf.annex_title_block(d, hd.get('code',''), hd.get('mk_title',''), hd.get('en_title',''),
                         hd.get('parent', hd.get('code','')))
    if hd.get('supersedes'):
        pr.note(d, "Заменува: "+hd['supersedes'], "Supersedes: "+hd['supersedes'])
    for b in blocks:
        if b[0]=='h':
            (emit_banner if b[1]==1 else emit_subbar)(d, b[2], b[3], b[4])
        elif b[0]=='p': pr.body(d, b[1], b[2])
        elif b[0]=='bullet': pr.bullet(d, b[1], b[2])
        elif b[0]=='form': emit_form(d, b[1], mode=(b[2] if len(b)>2 else None))
        elif b[0]=='table': emit_table(d, b[1])
        elif b[0]=='pagebreak': d.add_page_break()
    PAGE_W = pr.PAGE_W = 18.46            # restore default for any subsequent build in-process
    pf.save(d, out)

# =========================== SOP ===========================
def build_sop(hd, blocks, out):
    d=pf.new_sop(code=hd.get('code',''), version=hd.get('version','01'),
                 mk_title=hd.get('mk_title',''), en_title=hd.get('en_title',''))
    pf.sop_titlepage(d, hd.get('code','')+" · v"+hd.get('version','01'), hd.get('mk_title',''), hd.get('en_title',''))
    pf.sop_toc(d)
    t=None
    def ensure():
        nonlocal t
        if t is None: t=pf.sop_table(d)
        return t
    for b in blocks:
        if b[0]=='h':
            pf.sop_section_row(ensure(), (b[2] or ''), b[3], b[4], level=b[1])
        elif b[0]=='p':
            pf.sop_body_row(ensure(), b[1], b[2])
        elif b[0]=='bullet':
            pf.sop_body_row(ensure(), "•  "+b[1], "•  "+b[2])
        elif b[0] in ('table','form'):
            if t is not None: pf.sop_finalize(t); t=None      # close two-col table; table goes full width
            rows=b[1]
            if not rows: continue
            headers=[cs(c) for c in rows[0]]; data=[[cs(c) for c in r] for r in rows[1:]]
            ncol=len(headers)
            bt=pf.sop_block_table(d, headers, data, [round(PAGE_W/ncol,2)]*ncol)
            pr.fixed(bt)                                       # intelligent widths + repeat header
    if t is not None: pf.sop_finalize(t)
    pf.save(d, out)

def main(src,out):
    hd,blocks=parse(open(src,encoding='utf-8').read())
    dt=hd.get('doctype','SOP').upper()
    (build_sop if dt=='SOP' else build_annex)(hd, blocks, out)
    print("WROTE", out, "("+dt+")")

if __name__=='__main__':
    main(sys.argv[1], sys.argv[2])
