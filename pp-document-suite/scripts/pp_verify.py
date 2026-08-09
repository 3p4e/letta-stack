# -*- coding: utf-8 -*-
"""
Purely Plant — pre-delivery VERIFICATION GATE (general purpose, all QMS teams).

Runs the SKILL.md §6 checks on any produced .docx and prints a PASS/FAIL report:
  * structure  : paragraphs, tables, native equations (oMath), figures (drawings)
  * typography : minimum font size across body + header + footer (house floor 6 pt)
  * bilingual  : presence of Cyrillic (MK) and Latin (EN) text
  * fidelity   : (optional) output word & character counts vs a --source .docx
                 (content-fidelity / no-impoverishment rule §5A: output >= source)

Usage:
  python pp_verify.py <document.docx> [--source <source.docx>] [--min-pt 6]

Exit code 0 = PASS, 1 = FAIL (font floor breach or fidelity shortfall).
Pure-Python (python-docx + stdlib); runs anywhere, no Office needed.
"""
import sys, re, zipfile, argparse
import docx

def _text(path):
    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            for c in row.cells:
                parts.append(c.text)
    return d, "\n".join(parts)

def _xml_parts(path):
    z = zipfile.ZipFile(path)
    return {n: z.read(n).decode("utf-8", "ignore") for n in z.namelist() if n.endswith(".xml")}

def _min_font(parts):
    mn = None
    for n, x in parts.items():
        if ("document" in n) or ("header" in n) or ("footer" in n):
            for m in re.findall(r'w:sz w:val="(\d+)"', x):
                v = int(m); mn = v if mn is None else min(mn, v)
    return mn

def analyse(path):
    d, txt = _text(path)
    parts = _xml_parts(path)
    doc = parts.get("word/document.xml", "")
    words = len(re.findall(r"\S+", txt)); chars = len(re.sub(r"\s+", "", txt))
    return dict(
        path=path, words=words, chars=chars, paras=len(d.paragraphs), tables=len(d.tables),
        omath=doc.count("<m:oMath"), figures=doc.count("<w:drawing"),
        min_half=_min_font(parts),
        cyr=bool(re.search(r"[Ѐ-ӿ]", txt)), lat=bool(re.search(r"[A-Za-z]", txt)),
    )

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("docx"); ap.add_argument("--source", action="append", default=[])
    ap.add_argument("--min-pt", type=float, default=6.0)
    a = ap.parse_args()
    r = analyse(a.docx); ok = True
    print(f"== VERIFY: {r['path']}")
    print(f"   paragraphs {r['paras']} · tables {r['tables']} · equations(oMath) {r['omath']} · figures {r['figures']}")
    print(f"   words {r['words']} · chars {r['chars']}")
    mnpt = (r['min_half']/2) if r['min_half'] else None
    floor_ok = (mnpt is None) or (mnpt >= a.min_pt)
    print(f"   min font {mnpt} pt  [{'OK' if floor_ok else 'FAIL <%.0f pt'%a.min_pt}]"); ok &= floor_ok
    biling = r['cyr'] and r['lat']
    print(f"   bilingual MK+EN {'OK' if biling else 'WARN (missing a language)'}")
    if a.source:
        sw = sc = 0
        for s in a.source:
            rs = analyse(s); sw += rs['words']; sc += rs['chars']
            print(f"   source {s}: words {rs['words']} · chars {rs['chars']}")
        fid = (r['words'] >= sw) and (r['chars'] >= sc)
        print(f"   FIDELITY (§5A) output>=source: words {r['words']}>={sw}, chars {r['chars']}>={sc}  [{'OK' if fid else 'FAIL'}]")
        ok &= fid
    print("RESULT:", "PASS" if ok else "FAIL")
    sys.exit(0 if ok else 1)

if __name__ == "__main__":
    main()
