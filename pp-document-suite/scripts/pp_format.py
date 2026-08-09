"""
pp_format.py — Purely Plant Document Suite formatting engine (python-docx).

Builds bilingual Macedonian|English controlled .docx in the Purely Plant house style:
  - SOP   : two-column (MK left | EN right) with vertical divider + Word native TOC
  - Annex : full-page bilingual inline (MK 11pt | EN 7pt) with the PP v2 colour palette

Render to PDF with Microsoft Word COM (Windows, see render_pdf.ps1) or:
    soffice --headless --convert-to pdf <file>.docx --outdir .

Requires: python-docx  (pip install python-docx).  Run with PYTHONUTF8=1 on Windows.
NOTE: this is the shared engine — smoke-test on a one-page sample before first production run.

MANDATORY PP HEADER/FOOTER (required on EVERY Purely Plant document — SOP, annex, all):
the running header (leaf logo + wordmark | bilingual Document name | Code of document + Version)
and the "Page X of Y" footer live in the base template  assets/PP_BASE_TEMPLATE.docx . They are
loaded automatically by new_sop()/new_annex() (from_template=True, the default), the title/code/
version are stamped by apply_pp_header(), the body is cleared by wipe_body(), and save() re-seats
the section <w:sectPr> last so the header/footer/page-geometry apply document-wide. See formatting_specs.md.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- palette / fonts ----------
NAVY  = RGBColor(0x2B, 0x54, 0x7E)   # annex section header / SOP accents
GREY  = RGBColor(0x59, 0x59, 0x59)
BLACK = RGBColor(0, 0, 0)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
SOP_GRAY = "E8E8E8"      # SOP section-header fill
SOP_BLUE = "D9E2F3"      # SOP approval-table header
A_SECTION = "2B547E"     # annex section header (white text)
A_LABEL   = "EDF2F7"     # annex label cell
A_ALT     = "F7FAFC"     # annex alternating row
A_WARM    = "FEF9E7"; A_ROSE = "FDEDEC"; A_MINT = "EAFAF1"
BORDER_GRAY = "B0BEC5"
FONT = "Calibri"

# ---------- mandatory PP base template (header + logo + footer + page geometry) ----------
import os
PP_TEMPLATE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "assets", "PP_BASE_TEMPLATE.docx")

# ---------- low-level helpers ----------
def run(p, t, sz, color=BLACK, bold=False, ital=False, font=FONT):
    r = p.add_run(t); f = r.font
    f.name = font; f.size = Pt(sz); f.color.rgb = color; r.bold = bold; r.italic = ital
    return r

def sp(p, before=0, after=0, line=None):
    pf = p.paragraph_format; pf.space_before = Pt(before); pf.space_after = Pt(after)
    if line: pf.line_spacing = line

def shade(cell, fill):
    """Cell fill — ALWAYS w:val='clear' (never 'solid')."""
    tcPr = cell._tc.get_or_add_tcPr(); sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:color'), 'auto'); sh.set(qn('w:fill'), fill)
    tcPr.append(sh)

def cell_borders(cell, sides, sz=4, color="000000"):
    """sides = subset of ('top','left','bottom','right') to draw; others nil."""
    tcPr = cell._tc.get_or_add_tcPr(); b = OxmlElement('w:tcBorders')
    for e in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement('w:' + e)
        if e in sides:
            el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), str(sz)); el.set(qn('w:color'), color)
        else:
            el.set(qn('w:val'), 'nil')
        b.append(el)
    tcPr.append(b)

def table_borders(tbl, sz=4, color=BORDER_GRAY):
    tblPr = tbl._tbl.tblPr; b = OxmlElement('w:tblBorders')
    for e in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement('w:' + e)
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), str(sz)); el.set(qn('w:space'), '0'); el.set(qn('w:color'), color)
        b.append(el)
    tblPr.append(b)

def fixed_widths(tbl, widths_cm):
    tbl.autofit = False; tbl.allow_autofit = False
    lay = OxmlElement('w:tblLayout'); lay.set(qn('w:type'), 'fixed'); tbl._tbl.tblPr.append(lay)
    for row in tbl.rows:
        for i, w in enumerate(widths_cm):
            if i < len(row.cells):       # guard for merged rows (e.g. annex section banners)
                row.cells[i].width = Cm(w)

def cell_margins(cell, top=29, bottom=29, left=58, right=58):
    tcPr = cell._tc.get_or_add_tcPr(); m = OxmlElement('w:tcMar')
    for side, val in (('top', top), ('left', left), ('bottom', bottom), ('right', right)):
        el = OxmlElement('w:' + side); el.set(qn('w:w'), str(val)); el.set(qn('w:type'), 'dxa'); m.append(el)
    tcPr.append(m)

def bilingual(p, mk, en, mk_sz=11, en_sz=7, bold=False, color=BLACK):
    """Inline 'MK | EN' run group (annex + table style)."""
    run(p, mk, mk_sz, color, bold=bold)
    run(p, " | ", en_sz, color)
    run(p, en, en_sz, color, bold=bold)

def _page(doc, orient="portrait", margin_cm=1.27):
    s = doc.sections[0]
    if orient == "landscape":
        s.orientation = WD_ORIENT.LANDSCAPE; s.page_width = Cm(29.7); s.page_height = Cm(21.0)
    else:
        s.orientation = WD_ORIENT.PORTRAIT; s.page_width = Cm(21.0); s.page_height = Cm(29.7)
    s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Cm(margin_cm)
    return s

def _normal(doc):
    st = doc.styles['Normal']; st.font.name = FONT; st.font.size = Pt(11)

# ============================ mandatory header / base template ============================
def wipe_body(d):
    """Clear all body content but keep the section <w:sectPr> (header/footer/page refs).
    Stashes it on d._pp_sectpr; save() re-seats it last."""
    body = d.element.body
    sect = body.find(qn('w:sectPr'))
    for ch in list(body):
        if ch is not sect:
            body.remove(ch)
    d._pp_sectpr = sect
    return sect

def apply_pp_header(d, mk_name, code, en_name, version="1.0"):
    """Stamp the base-template running header: bilingual Document name, Code of document, Version.
    Header is a 2x3 table: [leaf logo+wordmark] | [Document name MK/EN] | [Code of document / Ver].
    Only the value runs are overwritten; the labels and the logo are preserved from the template."""
    try:
        h = d.sections[0].header.tables[0]
    except (IndexError, AttributeError):
        return False
    def setrun(runs, i, t):
        if i < len(runs):
            runs[i].text = t
    nm = h.cell(0, 1).paragraphs[1].runs
    setrun(nm, 0, mk_name + " "); setrun(nm, 2, code); setrun(nm, 5, en_name)
    cd = h.cell(0, 2).paragraphs[2].runs
    setrun(cd, 0, code); setrun(cd, 2, ""); setrun(cd, 3, "")
    vr = h.cell(1, 2).paragraphs[0].runs
    if vr:
        vr[-1].text = version
    return True

# ============================ SOP (two-column) ============================
def new_sop(margin_cm=1.27, from_template=True, code=None,
            mk_name="СТАНДАРДНА ОПЕРАТИВНА ПРОЦЕДУРА", en_name=None, version="1.0", template=None,
            mk_title=None, en_title=None):
    """SOP document. By DEFAULT starts FROM the mandatory PP base template so the running header
    (logo + bilingual doc name + code + version), the 'Page X of Y' footer and the A4 page geometry
    are present on every page.
    The running header's 'Document name' shows the ACTUAL SOP title — pass mk_title (Macedonian)
    and en_title (English). These take precedence over the generic 'Standard Operating Procedure'
    phrase (mk_name/en_name), which is used ONLY as a fallback when no title is supplied."""
    tpl = template or PP_TEMPLATE
    hdr_mk = mk_title or mk_name
    hdr_en = en_title or en_name or ("STANDARD OPERATING PROCEDURE — %s" % (code or ""))
    if from_template and os.path.exists(tpl):
        d = Document(tpl)
        apply_pp_header(d, hdr_mk, code or "", hdr_en, version)
        wipe_body(d)
        _normal(d)
        return d
    d = Document(); _page(d, "portrait", margin_cm); _normal(d); return d

def sop_titlepage(d, code, mk_title, en_title):
    for txt, sz, col, bold in [("СТАНДАРДНА ОПЕРАТИВНА ПРОЦЕДУРА", 24, NAVY, True),
                               ("STANDARD OPERATING PROCEDURE", 14, GREY, False),
                               (code, 18, NAVY, True)]:
        p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; sp(p, 6, 4); run(p, txt, sz, col, bold=bold)
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; sp(p, 4, 12)
    run(p, mk_title, 12, BLACK, bold=True); run(p, " | ", 10, GREY); run(p, en_title, 12, GREY, bold=True)
    t = d.add_table(rows=4, cols=4); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = [("Дејство", "Action"), ("Позиција", "Position"), ("Име", "Name"), ("Датум/Потпис", "Date/Sign")]
    for j, (mk, en) in enumerate(hdr):
        c = t.cell(0, j); shade(c, SOP_BLUE); c.paragraphs[0].text = ''; bilingual(c.paragraphs[0], mk, en, 11, 7, bold=True)
    rows = [("Подготвил | Prepared by", "B. Nikolov, M.Pharm."),
            ("Проверил | Reviewed by", "J. Romevska"),
            ("Одобрил | Approved by", "S. Petrova")]
    for i, (act, name) in enumerate(rows, 1):
        t.cell(i, 0).paragraphs[0].text = ''; run(t.cell(i, 0).paragraphs[0], act, 10)
        run(t.cell(i, 2).paragraphs[0], name, 10)
    fixed_widths(t, [4.7, 4.7, 4.7, 4.7]); table_borders(t, 4, "000000")
    d.add_page_break()

def sop_toc(d):
    p = d.add_paragraph(); run(p, "СОДРЖИНА | TABLE OF CONTENTS", 14, NAVY, bold=True)
    par = d.add_paragraph(); fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), r'TOC \o "1-3" \h \z \u'); par._p.append(fld)
    note = d.add_paragraph(); run(note, "(Update field after opening: right-click → Update Field)", 8, GREY, ital=True)
    d.add_page_break()

def sop_table(d):
    """Empty 2-column table to hold section/body rows."""
    t = d.add_table(rows=0, cols=2); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    return t

def _set_heading(p, level=1):
    """Apply a Heading style so the native TOC field populates. Run-level font is set afterwards,
    preserving the house appearance (the heading style only makes the paragraph TOC-eligible)."""
    try:
        p.style = "Heading %d" % max(1, min(3, level))
    except Exception:
        pass

def sop_section_row(t, num, mk, en, level=1):
    """Two-column section header. level controls the TOC depth (1=1.0, 2=1.1, 3=1.1.1)."""
    L, R = t.add_row().cells
    for c in (L, R):
        shade(c, SOP_GRAY)
    cell_borders(L, ('right',)); cell_borders(R, ('left',))
    L.paragraphs[0].text = ''; _set_heading(L.paragraphs[0], level); run(L.paragraphs[0], f"{num} {mk}", 12, BLACK, bold=True)
    R.paragraphs[0].text = ''; run(R.paragraphs[0], f"{num} {en}", 12, GREY, bold=True)

def sop_body_row(t, mk, en, bold=False):
    L, R = t.add_row().cells
    cell_borders(L, ('right',)); cell_borders(R, ('left',))
    for cell, txt, col in ((L, mk, BLACK), (R, en, GREY)):
        p = cell.paragraphs[0]; p.text = ''; p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; sp(p, 0, 6)
        run(p, txt, 11, col, bold=bold)

def sop_finalize(t):
    fixed_widths(t, [9.0, 9.0])   # two equal columns within ~18.46 cm usable

def sop_block_table(d, headers, rows, widths_cm, header_bg=SOP_BLUE):
    """Full-width multi-column bilingual table for SOP §7 Records / §8 Related Documents / §9 Revision.
    headers, and each row, are lists of (mk, en) tuples (one per column)."""
    t = d.add_table(rows=1 + len(rows), cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, (mk, en) in enumerate(headers):
        c = t.cell(0, j); shade(c, header_bg); c.paragraphs[0].text = ''
        bilingual(c.paragraphs[0], mk, en, 11, 7, bold=True)
    for i, row in enumerate(rows, 1):
        for j, (mk, en) in enumerate(row):
            c = t.cell(i, j); c.paragraphs[0].text = ''; bilingual(c.paragraphs[0], mk, en, 11, 7)
    fixed_widths(t, widths_cm); table_borders(t, 4, "000000"); return t

# ============================ Annex (inline) ============================
def new_annex(orient="portrait", from_template=True, code=None,
              mk_name="АНЕКС", en_name=None, version="1.0", template=None,
              mk_title=None, en_title=None):
    """Annex document. Like new_sop, by DEFAULT starts FROM the PP base template so the mandatory
    header, logo and 'Page X of Y' footer are present (required on all annexes).
    The header 'Document name' shows the ACTUAL annex title (mk_title | en_title); the generic
    'ANNEX' phrase (mk_name/en_name) is used only as a fallback when no title is supplied."""
    tpl = template or PP_TEMPLATE
    hdr_mk = mk_title or mk_name
    hdr_en = en_title or en_name or ("ANNEX — %s" % (code or ""))
    if from_template and os.path.exists(tpl):
        d = Document(tpl)
        apply_pp_header(d, hdr_mk, code or "", hdr_en, version)
        wipe_body(d)
        _normal(d)
        if orient == "landscape":
            s = d.sections[0]; s.orientation = WD_ORIENT.LANDSCAPE
            s.page_width, s.page_height = Cm(29.7), Cm(21.0)
        return d
    d = Document(); _page(d, orient, 2.54 if orient == "portrait" else 1.27); _normal(d); return d

def annex_title_block(d, code, mk_title, en_title, parent_sop):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; sp(p, 2, 3); run(p, code, 12, NAVY, bold=True)
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; sp(p, 0, 3)
    run(p, mk_title, 14, BLACK, bold=True); run(p, " | ", 10, GREY); run(p, en_title, 10, GREY, bold=True)
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; sp(p, 0, 8); run(p, f"({parent_sop})", 10, GREY, ital=True)

def annex_table(d, widths_cm):
    t = d.add_table(rows=0, cols=len(widths_cm)); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t._pp_widths = widths_cm
    return t

def annex_cell(cell, mk, en, bold=False, bg=None):
    if bg:
        shade(cell, bg)
    color = WHITE if bg == A_SECTION else BLACK
    cell_margins(cell); cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]; p.text = ''; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; sp(p, 1, 1, line=0.85)
    bilingual(p, mk, en, 11, 7, bold=bold, color=color)

def annex_section_header(t, mk, en):
    """Full-width navy banner row (white text). Call after rows exist; merges the row."""
    row = t.add_row(); first = row.cells[0]
    if len(row.cells) > 1:
        first = first.merge(row.cells[-1])
    annex_cell(first, mk, en, bold=True, bg=A_SECTION)

def annex_row(t, values, bold_first=False, alt=False):
    """values = list of (mk,en) tuples, one per column."""
    row = t.add_row()
    for i, (mk, en) in enumerate(values):
        bg = A_LABEL if (bold_first and i == 0) else (A_ALT if alt else None)
        annex_cell(row.cells[i], mk, en, bold=(bold_first and i == 0), bg=bg)

def annex_finalize(t):
    fixed_widths(t, t._pp_widths); table_borders(t, 1, BORDER_GRAY)

def annex_signoff(d, roles=(("Изготвил (КК)", "Prepared (QC)"), ("Проверил (QA)", "Reviewed (QA)"), ("Одобрил (QP)", "Approved (QP)"))):
    w = round(18.0 / len(roles), 2)
    t = annex_table(d, [w] * len(roles))
    annex_row(t, list(roles), bold_first=False)
    for c in t.rows[0].cells:
        shade(c, A_SECTION)
        for r in c.paragraphs[0].runs:
            r.font.color.rgb = WHITE; r.font.bold = True
    t.add_row()  # blank signature row
    annex_finalize(t)
    return t

# ============================ save ============================
def save(d, path):
    sect = getattr(d, "_pp_sectpr", None)
    if sect is not None:
        body = d.element.body
        try:
            body.remove(sect); body.append(sect)
        except ValueError:
            pass
    d.save(path); return path

def figure(d, png_path, mk=None, en=None, width_cm=15.5):
    """Embed a chart PNG (centered) with a bilingual MK | EN caption. Generate charts with pp_charts.py.
    Charts are data visualisations (matplotlib PNGs) — distinct from formula notation, which stays as
    native Word equations. Use for statistical justification figures in reports."""
    d.add_picture(png_path, width=Cm(width_cm))
    d.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if mk or en:
        p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; sp(p, 1, 8)
        run(p, "Слика | Figure — ", 9, NAVY, bold=True)
        if mk: run(p, mk, 9, GREY, ital=True)
        if en: run(p, "  |  ", 8, GREY); run(p, en, 9, GREY, ital=True)
    return d
